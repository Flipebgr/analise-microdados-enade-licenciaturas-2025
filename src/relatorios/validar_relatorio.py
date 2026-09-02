from __future__ import annotations

import re
from collections.abc import Iterable
from pathlib import Path

from docx import Document

CAPITULOS_PADRAO = (
    "1 INTRODUÇÃO",
    "2 REFERENCIAL INSTITUCIONAL E METODOLÓGICO",
    "3 METODOLOGIA",
    "4 PANORAMA",
    "5 RESULTADOS",
    "6 DISCUSSÃO",
    "7 CONCLUSÃO",
    "REFERÊNCIAS",
)


def _contem_caminho_absoluto(texto: str) -> bool:
    """Detecta caminhos locais absolutos que não devem vazar para o relatório."""

    caminho_windows = re.search(r"(?i)\b[A-Z]:[\\/]", texto) is not None
    caminho_sandbox = "/mnt/data/" in texto
    return caminho_windows or caminho_sandbox


def validar_docx(
    caminho: Path,
    *,
    capitulos_esperados: Iterable[str] = CAPITULOS_PADRAO,
    minimo_figuras: int = 5,
    minimo_tabelas: int = 4,
) -> list[str]:
    """Valida requisitos estruturais mínimos de um relatório DOCX.

    O contrato é independente da área. O capítulo 4 é verificado pelo
    prefixo ``4 PANORAMA``, permitindo títulos como
    ``4 PANORAMA DA LICENCIATURA EM QUÍMICA``.
    """

    erros: list[str] = []

    if not caminho.exists() or caminho.stat().st_size == 0:
        return ["DOCX não gerado ou vazio"]

    doc = Document(caminho)
    texto = "\n".join(paragrafo.text for paragrafo in doc.paragraphs)

    for capitulo in capitulos_esperados:
        if capitulo not in texto:
            erros.append(f"Capítulo ausente: {capitulo}")

    if len(doc.inline_shapes) < minimo_figuras:
        erros.append(
            f"Relatório contém menos de {minimo_figuras} figuras incorporadas"
        )

    if len(doc.tables) < minimo_tabelas:
        erros.append(
            f"Relatório contém menos de {minimo_tabelas} tabelas"
        )

    if _contem_caminho_absoluto(texto):
        erros.append("Caminho absoluto encontrado no relatório")

    return erros
