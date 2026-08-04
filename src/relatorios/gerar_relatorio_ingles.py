from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from src.relatorios.conversao_pdf import converter_docx_para_pdf
from src.relatorios.figuras_relatorio import adicionar_figura
from src.relatorios.formatacao_abnt import configurar_cabecalho_rodape, configurar_documento
from src.relatorios.referencias import adicionar_referencias
from src.relatorios.resultado_relatorio import ResultadoRelatorio
from src.relatorios.tabelas_relatorio import adicionar_tabela

FONTE_DADOS = (
    "Elaboração própria com base nos microdados do Enade das Licenciaturas 2025 "
    "e na planilha de Conceito Enade."
)


def _ler_csv(caminho: Path, *, obrigatorio: bool = True) -> pd.DataFrame:
    if not caminho.exists():
        if obrigatorio:
            raise FileNotFoundError(f"Produto analítico ausente: {caminho}")
        return pd.DataFrame()
    return pd.read_csv(caminho)


def carregar_produtos(base: Path) -> dict[str, pd.DataFrame]:
    pasta = base / "dados_processados" / "ingles"
    arquivos = {
        "cursos": "cursos_ingles.csv",
        "base": "base_analitica_cursos.csv",
        "comparacoes": "comparacoes_regionais_validadas.csv",
        "benchmarks": "sensibilidade_benchmarks.csv",
        "associacoes": "associacoes_ecologicas.csv",
        "socio_ufpa": "tabela_socioeconomica_ufpa.csv",
        "processo_itens": "itens_processo_formativo.csv",
        "processo_diag": "diagnostico_consistencia_processo.csv",
        "recomendacao_dist": "distribuicao_recomendacao.csv",
        "auditoria": "auditoria_desempenho_sprint08.csv",
    }
    return {chave: _ler_csv(pasta / nome) for chave, nome in arquivos.items()}


def tabela_ofertas_ufpa(base: pd.DataFrame) -> pd.DataFrame:
    colunas = [
        "CO_CURSO",
        "ROTULO_OFERTA",
        "CONCEITO_ENADE_NUM",
        "INSCRITOS_NUM",
        "PARTICIPANTES_NUM",
        "TAXA_PARTICIPACAO_OFICIAL",
        "nt_ger_mean",
        "nt_obj_mean",
        "nt_dis_mean",
        "nt_ger_percentil_brasil",
    ]
    presentes = [c for c in colunas if c in base.columns]
    return (
        base.loc[base["CO_IES"].eq(569), presentes]
        .sort_values("ROTULO_OFERTA")
        .reset_index(drop=True)
    )


def tabela_regional_nt_ger(comparacoes: pd.DataFrame) -> pd.DataFrame:
    colunas = [
        "RECORTE",
        "N_CURSOS",
        "N_PARTICIPANTES",
        "MEDIA_CURSOS",
        "MEDIA_PONDERADA_PARTICIPANTES",
        "MEDIANA_CURSOS",
        "DP_CURSOS",
        "P25",
        "P75",
    ]
    base = comparacoes.loc[comparacoes["INDICADOR"].eq("nt_ger_mean")].copy()
    presentes = [c for c in colunas if c in base.columns]
    return base[presentes].reset_index(drop=True)


def tabela_perfil_ufpa(base: pd.DataFrame) -> pd.DataFrame:
    colunas = [
        "ROTULO_OFERTA",
        "CONCEITO_ENADE_NUM",
        "renda_ate_3sm_pct",
        "trabalha_pct",
        "acao_afirmativa_pct",
        "auxilio_permanencia_pct",
        "estudo_4h_ou_mais_pct",
        "pretende_magisterio_pct",
    ]
    presentes = [c for c in colunas if c in base.columns]
    return base.loc[base["CO_IES"].eq(569), presentes].sort_values("ROTULO_OFERTA")


def tabela_recomendacao_ufpa(base: pd.DataFrame) -> pd.DataFrame:
    colunas = [
        "ROTULO_OFERTA",
        "CONCEITO_ENADE_NUM",
        "qe_i68_n",
        "qe_i68_media",
        "qe_i68_mediana",
        "qe_i69_n",
        "qe_i69_media",
        "qe_i69_mediana",
        "qe_i70_n",
        "qe_i70_interesse_pct",
    ]
    presentes = [c for c in colunas if c in base.columns]
    return base.loc[base["CO_IES"].eq(569), presentes].sort_values("ROTULO_OFERTA")


def tabela_processo_ufpa(itens: pd.DataFrame, cursos: pd.DataFrame) -> pd.DataFrame:
    if itens.empty:
        return itens
    mapa = cursos.loc[cursos["CO_IES"].eq(569), ["CO_CURSO", "ROTULO_OFERTA"]]
    trabalho = itens.merge(mapa, on="CO_CURSO", how="inner", validate="many_to_one")
    if trabalho.empty:
        return trabalho
    resumo = (
        trabalho.groupby("ROTULO_OFERTA", observed=True)
        .agg(
            itens_com_resposta=("ITEM", "nunique"),
            n_valido_total=("n_valido", "sum"),
            media_itens=("media", "mean"),
            concordancia_media=("concordancia_pct", "mean"),
            ausencia_analitica_media=("ausencia_analitica_pct", "mean"),
        )
        .reset_index()
    )
    return resumo


def resumo_desempenho(base: pd.DataFrame) -> str:
    ufpa = base.loc[base["CO_IES"].eq(569)].copy()
    alvo = ufpa.loc[ufpa["CONCEITO_ENADE_NUM"].eq(1)]
    contraste = ufpa.loc[ufpa["CONCEITO_ENADE_NUM"].gt(1)]
    if alvo.empty:
        return "Não foram localizadas ofertas da UFPA com Conceito Enade 1 na base analítica."
    mediana_alvo = pd.to_numeric(alvo["nt_ger_mean"], errors="coerce").median()
    texto = (
        f"Foram localizadas {len(alvo)} ofertas da UFPA com Conceito Enade 1. "
        f"A mediana das médias de NT_GER dessas ofertas foi {mediana_alvo:.2f}."
    )
    if not contraste.empty:
        media_contraste = pd.to_numeric(contraste["nt_ger_mean"], errors="coerce").mean()
        texto += (
            " A oferta da UFPA com conceito superior apresentou média de NT_GER "
            f"de {media_contraste:.2f}. A diferença é descritiva e não implica causalidade."
        )
    return texto


def resumo_regional(comparacoes: pd.DataFrame) -> str:
    tab = tabela_regional_nt_ger(comparacoes).set_index("RECORTE")
    desejados = ["UFPA agregada", "Região Norte sem UFPA", "Brasil geral"]
    faltantes = [r for r in desejados if r not in tab.index]
    if faltantes:
        return "Nem todos os recortes regionais obrigatórios estavam disponíveis para a síntese textual."
    partes = []
    for recorte in desejados:
        linha = tab.loc[recorte]
        partes.append(
            f"{recorte}: média simples {linha['MEDIA_CURSOS']:.2f} e média ponderada por participantes "
            f"{linha['MEDIA_PONDERADA_PARTICIPANTES']:.2f}"
        )
    return "; ".join(partes) + ". As duas médias respondem a ponderações analíticas distintas."


def resumo_associacoes(associacoes: pd.DataFrame) -> str:
    validas = associacoes.dropna(subset=["SPEARMAN_RHO"]).copy()
    if validas.empty:
        return "Não houve associações ecológicas calculáveis com os indicadores disponíveis."
    idx = validas["SPEARMAN_RHO"].abs().idxmax()
    linha = validas.loc[idx]
    return (
        "A associação ecológica de maior magnitude absoluta entre as examinadas foi entre "
        f"{linha['INDICADOR_X']} e NT_GER médio (Spearman rho={linha['SPEARMAN_RHO']:.3f}; "
        f"N={int(linha['N_CURSOS'])} cursos). Trata-se de associação entre cursos, não entre estudantes, "
        "e o p-valor é apenas exploratório."
    )


def _p(doc: Document, texto: str, estilo: str | None = None):
    p = doc.add_paragraph(style=estilo)
    p.add_run(texto)
    return p


def _capa(doc: Document) -> None:
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = 0
    r = p.add_run("UNIVERSIDADE FEDERAL DO PARÁ")
    r.bold = True
    r.font.size = Pt(12)
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = 0
    r = p.add_run(
        "DESEMPENHO, COMPOSIÇÃO DISCENTE, TRAJETÓRIA ACADÊMICA E PROCESSO FORMATIVO "
        "NAS OFERTAS DE LETRAS–INGLÊS DA UFPA NO ENADE 2025"
    )
    r.bold = True
    r.font.size = Pt(14)
    for _ in range(7):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = 0
    p.add_run("Autor(a): [preencher]")
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = 0
    p.add_run("Belém\n2026")


def _figura(doc: Document, base: Path, nome: str, numero: int, titulo: str, leitura: str, limitacao: str) -> None:
    adicionar_figura(doc, base / "figuras" / "ingles" / nome, f"Figura {numero} – {titulo}", FONTE_DADOS)
    _p(doc, f"Descrição e interpretação: {leitura}")
    _p(doc, f"Limitação: {limitacao}")


def _markdown(dados: dict[str, pd.DataFrame]) -> str:
    base = dados["base"]
    ofertas = tabela_ofertas_ufpa(base)
    regional = tabela_regional_nt_ger(dados["comparacoes"])
    perfil = tabela_perfil_ufpa(base)
    recomendacao = tabela_recomendacao_ufpa(base)
    processo = tabela_processo_ufpa(dados["processo_itens"], dados["cursos"])
    associacoes = dados["associacoes"]
    linhas = [
        "# DESEMPENHO, COMPOSIÇÃO DISCENTE, TRAJETÓRIA ACADÊMICA E PROCESSO FORMATIVO NAS OFERTAS DE LETRAS–INGLÊS DA UFPA NO ENADE 2025",
        "",
        "## RESUMO",
        "",
        f"Este relatório analisa {len(base)} cursos de Letras–Inglês no Enade 2025, com ênfase nas {int(base['CO_IES'].eq(569).sum())} ofertas localizadas da UFPA. A unidade principal é CO_CURSO. Arquivos temáticos foram tratados separadamente, agregados por curso e somente então relacionados. São examinados desempenho, composição discente, trajetória, processo formativo, recomendação, benchmarks comparáveis e associações ecológicas. Os resultados descrevem padrões entre cursos e não sustentam inferência causal individual.",
        "",
        "**Palavras-chave:** Enade; Letras–Inglês; UFPA; formação de professores; microdados.",
        "",
        "## ABSTRACT",
        "",
        "This technical-scientific report analyzes English Language teacher education programs in the 2025 Enade, focusing on UFPA offers. The course (CO_CURSO) is the main unit of analysis. Thematic files were processed separately, aggregated at course level, and only then combined. Performance, student composition, academic trajectory, formative process, recommendation, comparable benchmarks, and ecological associations are examined. Results are descriptive at course level and do not support individual causal inference.",
        "",
        "**Keywords:** Enade; English Language teacher education; UFPA; teacher education; microdata.",
        "",
        "# 1 INTRODUÇÃO",
        "",
        "A pergunta central é: quais características de desempenho, composição discente, trajetória acadêmica e avaliação do processo formativo diferenciam as ofertas da UFPA com Conceito Enade 1 das demais ofertas da mesma área na UFPA, no Pará, na Região Norte e no Brasil?",
        "",
        "# 2 REFERENCIAL INSTITUCIONAL E METODOLÓGICO",
        "",
        "O Conceito Enade é tratado como classificação externa do curso. Ausência de conceito não é recodificada como Conceito 1. O desenho analítico respeita os limites relacionais dos arquivos temáticos do Enade.",
        "",
        "# 3 METODOLOGIA",
        "",
        "A unidade principal é CO_CURSO. Cada arquivo temático foi tratado separadamente, com tratamento de ausentes e agregação por curso antes das junções one-to-one. Não foi usada posição de linha nem identificador artificial. Os grupos independentes são: A) UFPA Conceito 1; B) demais ofertas da UFPA com conceito superior; C) outras IES do Pará; D) restante da Região Norte, excluindo Pará; E) restante do Brasil, excluindo Norte. Norte e Brasil completos aparecem apenas como benchmarks descritivos sobrepostos.",
        "",
        "# 4 PANORAMA DA LICENCIATURA EM LETRAS–INGLÊS",
        "",
        f"A base analítica contém {len(base)} cursos. Na UFPA, foram localizadas {len(ofertas)} ofertas; {int(ofertas['CONCEITO_ENADE_NUM'].eq(1).sum())} possuem Conceito Enade 1 e {int(ofertas['CONCEITO_ENADE_NUM'].gt(1).sum())} possui conceito superior.",
        "",
        "## Tabela 1 – Ofertas da UFPA",
        "",
        ofertas.to_markdown(index=False, floatfmt=".3f"),
        "",
        "# 5 RESULTADOS",
        "",
        "## 5.1 Desempenho",
        "",
        resumo_desempenho(base),
        "",
        "## 5.2 Perfil demográfico e socioeconômico",
        "",
        "Os indicadores foram calculados com denominadores válidos por item. Percentuais de renda, trabalho, ação afirmativa, auxílios, estudo e intenção de magistério são apresentados no nível do curso.",
        "",
        perfil.to_markdown(index=False, floatfmt=".3f"),
        "",
        "## 5.3 Trajetória e condições acadêmicas",
        "",
        "A trajetória acadêmica é examinada por indicadores agregados de turno e demais condições disponíveis nos arquivos temáticos. Relações com desempenho só podem ser interpretadas ecologicamente quando provenientes de arquivos diferentes.",
        "",
        "## 5.4 Processo formativo",
        "",
        "Os itens QE_I20–QE_I66 foram preservados em nível de item e curso. Não foi criado índice único. O alfa de Cronbach global é apenas diagnóstico e não substitui validação teórica de dimensionalidade, orientação da escala e itens potencialmente invertidos.",
        "",
        processo.to_markdown(index=False, floatfmt=".3f"),
        "",
        "## 5.5 Recomendação",
        "",
        "QE_I68 e QE_I69 são apresentados como recomendação, usando os rótulos funcionais dos itens e sem convertê-los automaticamente em satisfação geral. QE_I70 é utilizado quando pertinente.",
        "",
        recomendacao.to_markdown(index=False, floatfmt=".3f"),
        "",
        "## 5.6 Benchmark comparável",
        "",
        "O benchmark comparável restringe modalidade, categoria administrativa, organização acadêmica e porte, com análise de sensibilidade. Conjuntos pequenos são tratados como descritivos.",
        "",
        "## 5.7 Associações ecológicas",
        "",
        resumo_associacoes(associacoes),
        "",
        "## 5.8 Comparação regional e nacional",
        "",
        resumo_regional(dados["comparacoes"]),
        "",
        regional.to_markdown(index=False, floatfmt=".3f"),
        "",
        "# 6 DISCUSSÃO",
        "",
        "Os contrastes observados devem ser lidos como padrões de cursos. Diferenças de desempenho podem coexistir com diferenças de porte, composição discente, participação, condições acadêmicas e avaliação do processo formativo. O desenho não identifica mecanismos causais nem permite atribuir resultados a modalidade, interiorização, perfil socioeconômico ou qualquer indicador isolado.",
        "",
        "# 7 CONCLUSÃO",
        "",
        "As ofertas da UFPA com Conceito Enade 1 apresentam heterogeneidade interna e devem ser analisadas individualmente, mantendo como contraste a oferta da UFPA com conceito superior e referências territoriais e comparáveis. O conjunto de evidências serve para priorização de investigação institucional, não para inferência causal sobre estudantes ou cursos.",
        "",
        "# REFERÊNCIAS",
        "",
        "As referências bibliográficas e normativas são as mesmas cadastradas no módulo compartilhado de referências do projeto e são inseridas integralmente na versão DOCX.",
        "",
        "# APÊNDICES",
        "",
        "## APÊNDICE A – REGRAS DE INTEGRIDADE",
        "",
        "Não foram realizadas junções individuais entre arquivos temáticos. Todas as integrações ocorreram após agregação por CO_CURSO e validação one-to-one.",
        "",
        "## APÊNDICE B – APROFUNDAMENTOS SUGERIDOS",
        "",
        "1. Dimensionalidade de QE_I20–QE_I66: verificar agrupamentos teóricos e consistência interna por dimensão; método: análise de itens e consistência interna; limitação: escala ordinal e possíveis ausências estruturais.",
        "2. Sensibilidade ao porte: perguntar quanto a posição das ofertas muda sob faixas alternativas de participantes; método: benchmarks pareados/sensibilidade; limitação: poucos comparáveis em alguns estratos.",
        "3. Participação e desempenho: investigar se resultados se mantêm em cursos com taxas de participação semelhantes; método: estratificação e regressão ecológica; limitação: unidade curso.",
        "4. Perfil socioeconômico e desempenho: avaliar associações ecológicas com renda, trabalho, ação afirmativa e auxílios; método: Spearman, dispersão e análise de outliers; limitação: falácia ecológica.",
        "5. Recomendação e processo formativo: examinar convergência entre itens de recomendação e dimensões formativas validadas; método: indicadores agregados por curso; limitação: não representa relação individual entre respostas de arquivos distintos.",
    ]
    return "\n".join(linhas)


def gerar_relatorio(
    base_projeto: Path,
    saida_docx: Path,
    saida_md: Path,
) -> dict:
    dados = carregar_produtos(base_projeto)
    base = dados["base"]
    ofertas = tabela_ofertas_ufpa(base)
    regional = tabela_regional_nt_ger(dados["comparacoes"])
    perfil = tabela_perfil_ufpa(base)
    processo = tabela_processo_ufpa(dados["processo_itens"], dados["cursos"])
    recomendacao = tabela_recomendacao_ufpa(base)

    doc = Document()
    configurar_documento(doc)
    _capa(doc)
    doc.add_page_break()

    doc.add_heading("RESUMO", level=1)
    _p(
        doc,
        f"Este relatório analisa {len(base)} cursos de Letras–Inglês no Enade 2025, com ênfase nas "
        f"{len(ofertas)} ofertas localizadas da UFPA. A unidade principal é CO_CURSO. Arquivos temáticos "
        "foram processados separadamente, agregados por curso e somente então relacionados. São examinados "
        "desempenho, composição discente, trajetória acadêmica, processo formativo, recomendação, benchmarks "
        "comparáveis e associações ecológicas. Os resultados descrevem padrões entre cursos e não sustentam "
        "inferência causal individual.",
        "Resumo",
    )
    _p(doc, "Palavras-chave: Enade; Letras–Inglês; UFPA; formação de professores; microdados.")

    doc.add_heading("ABSTRACT", level=1)
    _p(
        doc,
        "This technical-scientific report analyzes English Language teacher education programs in the 2025 "
        "Enade, focusing on UFPA offers. The course (CO_CURSO) is the main unit of analysis. Thematic files "
        "were processed separately, aggregated at course level, and only then combined. Performance, student "
        "composition, academic trajectory, formative process, recommendation, comparable benchmarks, and "
        "ecological associations are examined. Results are descriptive at course level and do not support "
        "individual causal inference.",
        "Resumo",
    )
    _p(doc, "Keywords: Enade; English Language teacher education; UFPA; teacher education; microdata.")

    doc.add_heading("1 INTRODUÇÃO", level=1)
    _p(doc, "O relatório responde à seguinte pergunta: quais características de desempenho, composição discente, trajetória acadêmica e avaliação do processo formativo diferenciam as ofertas da UFPA com Conceito Enade 1 das demais ofertas da mesma área na UFPA, no Pará, na Região Norte e no Brasil?")

    doc.add_heading("2 REFERENCIAL INSTITUCIONAL E METODOLÓGICO", level=1)
    _p(doc, "O Conceito Enade é tratado como classificação externa do curso. A ausência de conceito permanece distinta de Conceito 1. O tratamento metodológico segue a estrutura dos microdados do Enade das Licenciaturas 2025 e preserva os limites de relacionamento entre arquivos temáticos.")

    doc.add_heading("3 METODOLOGIA", level=1)
    _p(doc, "A unidade de análise é CO_CURSO. Cada arquivo temático foi tratado separadamente, com tratamento de ausentes e agregação por curso antes de qualquer junção. Não foi usada posição de linha nem criado identificador artificial. As junções entre tabelas agregadas são validadas como one-to-one.")
    _p(doc, "Os grupos independentes são: A) UFPA com Conceito Enade 1; B) demais ofertas da UFPA da mesma área com conceito superior; C) outras IES do Pará; D) restante da Região Norte, excluindo Pará; E) restante do Brasil, excluindo Norte. Pará, Norte e Brasil completos são usados somente como benchmarks descritivos sobrepostos.")

    doc.add_heading("4 PANORAMA DA LICENCIATURA EM LETRAS–INGLÊS", level=1)
    _p(doc, f"A base analítica contém {len(base)} cursos de Letras–Inglês. Foram localizadas {len(ofertas)} ofertas da UFPA, das quais {int(ofertas['CONCEITO_ENADE_NUM'].eq(1).sum())} possuem Conceito Enade 1 e {int(ofertas['CONCEITO_ENADE_NUM'].gt(1).sum())} possui conceito superior.")
    adicionar_tabela(doc, "Tabela 1 – Ofertas da UFPA", ofertas, FONTE_DADOS)
    _figura(doc, base_projeto, "01_painel_ofertas_ufpa.png", 1, "Ofertas da UFPA", "O painel reúne as cinco ofertas localizadas e explicita conceito, município e porte observados.", "A visualização é descritiva e não controla diferenças estruturais entre ofertas.")

    doc.add_heading("5 RESULTADOS", level=1)
    doc.add_heading("5.1 Desempenho", level=2)
    _p(doc, resumo_desempenho(base))
    _figura(doc, base_projeto, "02_posicao_relativa_nt_ger.png", 2, "Posição relativa em NT_GER", "A posição relativa situa cada oferta da UFPA na distribuição da própria área.", "Percentis não eliminam diferenças de composição, porte e participação.")
    _figura(doc, base_projeto, "03_distribuicao_nt_ger.png", 3, "Distribuição de NT_GER", "A distribuição individual de NT_GER é apresentada apenas dentro do mesmo arquivo de desempenho.", "Cursos com N pequeno produzem distribuições mais instáveis.")
    _figura(doc, base_projeto, "04_distribuicao_nt_obj.png", 4, "Distribuição de NT_OBJ", "A componente objetiva complementa a leitura da nota geral.", "NT_OBJ e NT_GER possuem relação mecânica e não devem ser tratados como evidências independentes.")
    _figura(doc, base_projeto, "05_distribuicao_nt_dis.png", 5, "Distribuição de NT_DIS", "A componente discursiva é mostrada separadamente para preservar sua escala e distribuição.", "Diferenças de correção e N válido devem ser consideradas na interpretação.")

    doc.add_heading("5.2 Perfil demográfico e socioeconômico", level=2)
    _p(doc, "Os percentuais usam denominadores válidos por item e permanecem acompanhados das regras documentadas de classificação. A tabela abaixo resume indicadores socioeconômicos selecionados das ofertas da UFPA.")
    adicionar_tabela(doc, "Tabela 2 – Indicadores socioeconômicos selecionados da UFPA", perfil, FONTE_DADOS)
    _figura(doc, base_projeto, "06_perfil_socioeconomico.png", 6, "Perfil socioeconômico", "O gráfico compara indicadores agregados por curso, mantendo a oferta como unidade analítica.", "Não é possível inferir associação individual entre perfil socioeconômico e desempenho a partir de arquivos temáticos distintos.")

    doc.add_heading("5.3 Trajetória e condições acadêmicas", level=2)
    _p(doc, "Indicadores de turno, trabalho, bolsas, auxílios, horas de estudo e intenção de exercer o magistério foram agregados por curso. A interpretação deve considerar N válido, ausências e a impossibilidade de vincular registros individuais entre arquivos temáticos.")

    doc.add_heading("5.4 Processo formativo", level=2)
    _p(doc, "QE_I20–QE_I66 foram analisados item a item. O diagnóstico global de consistência interna não foi usado para criar índice único. A formação de dimensões requer leitura substantiva dos itens, verificação da orientação da escala e eventual tratamento de itens invertidos.")
    if not processo.empty:
        adicionar_tabela(doc, "Tabela 3 – Síntese descritiva dos itens de processo formativo nas ofertas da UFPA", processo, FONTE_DADOS)
    _figura(doc, base_projeto, "07_processo_formativo.png", 7, "Processo formativo", "A figura sintetiza respostas aos itens formativos sem assumir uma dimensão única.", "Médias entre itens heterogêneos são apenas descrições exploratórias; não equivalem a escala validada.")

    doc.add_heading("5.5 Recomendação", level=2)
    _p(doc, "QE_I68 e QE_I69 são apresentados como itens de recomendação e não são renomeados como satisfação. QE_I70 é incluído quando pertinente.")
    adicionar_tabela(doc, "Tabela 4 – Recomendação nas ofertas da UFPA", recomendacao, FONTE_DADOS)

    doc.add_heading("5.6 Benchmark comparável", level=2)
    _p(doc, "O benchmark comparável restringe modalidade, categoria administrativa, organização acadêmica e porte. A Sprint 08 também verificou sensibilidade a faixas de tamanho, evitando tratar um único conjunto de comparáveis como referência definitiva.")
    adicionar_tabela(doc, "Tabela 5 – Sensibilidade dos benchmarks comparáveis", dados["benchmarks"], FONTE_DADOS)
    _figura(doc, base_projeto, "validada_11_benchmark_comparavel.png", 8, "Benchmark comparável validado", "O contraste posiciona as ofertas Conceito 1 frente a cursos estruturalmente mais semelhantes.", "A comparabilidade é observacional e não constitui pareamento causal.")

    doc.add_heading("5.7 Associações ecológicas", level=2)
    _p(doc, resumo_associacoes(dados["associacoes"]))
    adicionar_tabela(doc, "Tabela 6 – Associações ecológicas exploratórias", dados["associacoes"], FONTE_DADOS)

    doc.add_heading("5.8 Comparação regional e nacional", level=2)
    _p(doc, resumo_regional(dados["comparacoes"]))
    adicionar_tabela(doc, "Tabela 7 – Comparação regional e nacional de NT_GER", regional, FONTE_DADOS)
    _figura(doc, base_projeto, "08_comparacao_regional_nacional.png", 9, "Comparação regional e nacional", "A figura compara UFPA, Norte e demais regiões na própria área de Letras–Inglês.", "Recortes territoriais completos se sobrepõem e por isso não são tratados como grupos independentes em testes.")
    _figura(doc, base_projeto, "validada_12_comparacao_regional.png", 10, "Comparação regional validada", "A validação confronta média simples, ponderação por participantes e dispersão entre cursos.", "Diferenças entre média simples e ponderada refletem o peso dos cursos maiores.")
    _figura(doc, base_projeto, "validada_09_participacao_ufpa.png", 11, "Participação nas ofertas da UFPA", "A participação é mostrada explicitamente para contextualizar a estabilidade das estimativas de desempenho.", "Taxas elevadas ou baixas de participação não são, isoladamente, explicações para o Conceito Enade.")
    _figura(doc, base_projeto, "validada_10_sensibilidade_desempenho.png", 12, "Sensibilidade do desempenho", "A análise verifica a estabilidade dos contrastes sob diferentes critérios analíticos.", "A sensibilidade reduz dependência de uma única especificação, mas não cria identificação causal.")

    doc.add_heading("6 DISCUSSÃO", level=1)
    _p(doc, "Os resultados indicam heterogeneidade entre as ofertas da UFPA com Conceito Enade 1 e reforçam a necessidade de evitar explicações únicas. Desempenho, participação, porte, composição discente, condições acadêmicas, avaliação do processo formativo e recomendação podem apresentar padrões concomitantes. Como os arquivos temáticos não compartilham identificador discente, relações entre temas distintos são ecológicas e devem ser interpretadas no nível do curso.")
    _p(doc, "A oferta de Belém, com conceito superior, funciona como contraste interno institucional, enquanto os benchmarks regionais, nacionais e comparáveis ampliam a referência. Ainda assim, diferenças de contexto e estrutura impedem interpretar os contrastes como efeito de campus, interiorização, modalidade ou perfil discente.")

    doc.add_heading("7 CONCLUSÃO", level=1)
    _p(doc, "A análise identifica padrões que distinguem as ofertas da UFPA, mas não estabelece causalidade. As quatro ofertas com Conceito Enade 1 devem ser acompanhadas individualmente, combinando desempenho, participação, perfil, trajetória, processo formativo, recomendação e comparação com cursos estruturalmente semelhantes. A síntese produzida fornece base empírica para investigação institucional orientada por evidências e para priorização de aprofundamentos.")

    doc.add_heading("REFERÊNCIAS", level=1)
    adicionar_referencias(doc)

    doc.add_heading("APÊNDICE A – REGRAS DE INTEGRIDADE", level=1)
    _p(doc, "A unidade principal é CO_CURSO. Não foi usada posição de linha, não foram criados identificadores artificiais e nenhuma junção individual foi realizada entre arquivos temáticos. As tabelas foram agregadas separadamente antes das junções one-to-one.")
    doc.add_heading("APÊNDICE B – APROFUNDAMENTOS SUGERIDOS", level=1)
    aprofundamentos = [
        ("Dimensionalidade do processo formativo", "Quais dimensões teóricas de QE_I20–QE_I66 são empiricamente sustentáveis?", "QE_I20–QE_I66", "leitura substantiva dos itens, consistência interna por dimensão e análise de sensibilidade", "escala ordinal, ausências e possíveis itens invertidos"),
        ("Sensibilidade ao porte", "A posição das ofertas Conceito 1 muda quando o porte dos comparáveis é restringido?", "participantes, modalidade, categoria e organização acadêmica", "benchmarks alternativos e pareamento observacional", "poucos comparáveis em alguns estratos"),
        ("Participação e desempenho", "Os contrastes de desempenho persistem entre cursos com participação semelhante?", "inscritos, participantes, taxa de participação, NT_GER", "estratificação e regressão ecológica", "unidade curso e possível instabilidade em N pequeno"),
        ("Perfil socioeconômico", "Quais indicadores agregados de perfil apresentam associação monotônica com desempenho médio?", "renda, trabalho, ação afirmativa, auxílios, NT_GER", "Spearman, dispersão, outliers e ponderação por participantes", "falácia ecológica e ausência de vínculo individual"),
        ("Recomendação e formação", "A recomendação varia de forma coerente com dimensões formativas validadas?", "QE_I68, QE_I69, QE_I70 e dimensões validadas de QE_I20–QE_I66", "associações agregadas por curso", "não representa relação entre respostas dos mesmos estudantes em arquivos distintos"),
    ]
    for titulo, pergunta, variaveis, metodo, limitacao in aprofundamentos:
        _p(doc, f"{titulo}. Pergunta: {pergunta} Variáveis: {variaveis}. Método: {metodo}. Limitação: {limitacao}.")

    configurar_cabecalho_rodape(doc)
    saida_docx.parent.mkdir(parents=True, exist_ok=True)
    saida_md.parent.mkdir(parents=True, exist_ok=True)
    doc.save(saida_docx)
    saida_md.write_text(_markdown(dados), encoding="utf-8")
    conversao = converter_docx_para_pdf(saida_docx, saida_docx.parent)
    return ResultadoRelatorio(saida_docx, saida_md, conversao).como_dict()
