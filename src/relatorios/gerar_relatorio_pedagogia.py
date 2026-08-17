from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from src.pedagogia import CO_CURSO_CASTANHAL, PEDAGOGIA
from src.pedagogia.rotulos_questionario import ROTULOS_QE, rotulo_item
from src.relatorios.conversao_pdf import converter_docx_para_pdf
from src.relatorios.figuras_relatorio import adicionar_figura
from src.relatorios.formatacao_abnt import configurar_cabecalho_rodape, configurar_documento
from src.relatorios.referencias import adicionar_referencias
from src.relatorios.resultado_relatorio import ResultadoRelatorio
from src.relatorios.tabelas_relatorio import adicionar_tabela

FONTE_DADOS = (
    "Elaboração própria com base nos microdados do Enade das Licenciaturas 2025 "
    "e na planilha de Conceito Enade."
)
FONTE_QUESTIONARIO = (
    "Inep, Dicionário de Variáveis e Questionário do Estudante – "
    "Enade das Licenciaturas 2025."
)


def _ler_csv(caminho: Path) -> pd.DataFrame:
    if not caminho.exists():
        raise FileNotFoundError(
            f"Produto analítico ausente: {caminho}. "
            "Execute as Sprints 13 e 14 antes da Sprint 15."
        )
    return pd.read_csv(caminho, low_memory=False)


def carregar_produtos(base_projeto: Path) -> dict[str, pd.DataFrame]:
    pasta = base_projeto / "dados_processados" / "pedagogia"
    arquivos = {
        "cursos": "cursos_pedagogia.csv",
        "base": "base_analitica_cursos.csv",
        "comparacoes": "comparacoes_regionais_validadas_sprint14.csv",
        "sensibilidade": "sensibilidade_benchmarks_sprint14.csv",
        "membros_benchmark": "membros_benchmarks_sprint14.csv",
        "contraste": "contraste_interno_ufpa_validado.csv",
        "perfil": "perfil_recortes_validado.csv",
        "processo": "processo_formativo_castanhal_validado.csv",
        "associacoes": "associacoes_ecologicas_sprint14.csv",
        "outliers": "diagnostico_outliers_sprint14.csv",
        "auditoria": "auditoria_desempenho_sprint14.csv",
        "recomendacao_dist": "distribuicao_recomendacao.csv",
        "processo_bruto": "itens_processo_formativo.csv",
    }
    return {chave: _ler_csv(pasta / nome) for chave, nome in arquivos.items()}


def _num(valor) -> float:
    return pd.to_numeric(pd.Series([valor]), errors="coerce").iloc[0]


def _fmt(valor, casas: int = 2) -> str:
    valor = _num(valor)
    if pd.isna(valor):
        return "não disponível"
    return f"{valor:.{casas}f}".replace(".", ",")


def obter_castanhal(base: pd.DataFrame) -> pd.Series:
    mask = pd.to_numeric(base["CO_CURSO"], errors="coerce").eq(CO_CURSO_CASTANHAL)
    alvo = base.loc[mask]
    if len(alvo) != 1:
        raise ValueError(
            f"Esperada exatamente uma oferta de Castanhal ({CO_CURSO_CASTANHAL}); "
            f"encontradas {len(alvo)}."
        )
    return alvo.iloc[0]


def tabela_ofertas_ufpa(base: pd.DataFrame) -> pd.DataFrame:
    colunas = [
        "CO_CURSO",
        "ROTULO_OFERTA",
        "CONCEITO_ENADE_NUM",
        "INSCRITOS_NUM",
        "PARTICIPANTES_NUM",
        "TAXA_PARTICIPACAO_OFICIAL",
        "taxa_presenca_microdados",
        "nt_ger_mean",
        "nt_obj_mean",
        "nt_dis_mean",
        "nt_ger_percentil_brasil",
        "nt_ger_percentil_norte",
        "nt_ger_percentil_para",
    ]
    presentes = [c for c in colunas if c in base.columns]
    return (
        base.loc[base["CO_IES"].eq(PEDAGOGIA.co_ies_focal), presentes]
        .sort_values("ROTULO_OFERTA")
        .reset_index(drop=True)
    )


def tabela_regional_nt_ger(comparacoes: pd.DataFrame) -> pd.DataFrame:
    trabalho = comparacoes.loc[comparacoes["INDICADOR"].eq("nt_ger_mean")].copy()
    colunas = [
        "RECORTE",
        "N_CURSOS",
        "N_PARTICIPANTES",
        "MEDIA_CURSOS",
        "MEDIA_PONDERADA_PARTICIPANTES",
        "MEDIANA_CURSOS",
        "DP_CURSOS",
        "P25",
        "P75",
        "AMPLITUDE_IQR",
    ]
    presentes = [c for c in colunas if c in trabalho.columns]
    return trabalho[presentes].reset_index(drop=True)


def tabela_benchmark_principal(sensibilidade: pd.DataFrame) -> pd.DataFrame:
    principal = sensibilidade.loc[
        sensibilidade["CENARIO"].eq("estrutura_porte_0_5_2_0")
    ].copy()
    colunas = [
        "CO_CURSO_ALVO",
        "ROTULO_ALVO",
        "CONCEITO_ALVO",
        "N_CURSOS",
        "nt_ger_mean_ALVO",
        "nt_ger_mean_MEDIA_BENCHMARK",
        "nt_ger_mean_DIFERENCA",
        "nt_ger_mean_Z",
        "nt_obj_mean_DIFERENCA",
        "nt_dis_mean_DIFERENCA",
        "taxa_presenca_microdados_DIFERENCA",
    ]
    presentes = [c for c in colunas if c in principal.columns]
    return principal[presentes].sort_values("ROTULO_ALVO").reset_index(drop=True)


def tabela_perfil_interno(perfil: pd.DataFrame) -> pd.DataFrame:
    indicadores = [
        "sexo_feminino_pct",
        "idade_media",
        "mae_superior_pct",
        "pai_superior_pct",
        "renda_ate_3sm_pct",
        "trabalha_pct",
        "acao_afirmativa_pct",
        "auxilio_permanencia_pct",
        "bolsa_academica_pct",
        "estudo_4h_ou_mais_pct",
        "turno_noturno_pct",
        "anos_desde_ingresso_media",
        "qe_i68_media",
        "qe_i69_media",
        "qe_i70_interesse_pct",
    ]
    trabalho = perfil.loc[
        perfil["RECORTE_PEDAGOGIA"].isin(["UFPA — Conceito 5", "UFPA — Conceito 4"])
        & perfil["INDICADOR"].isin(indicadores)
    ].copy()
    return trabalho[
        [
            "RECORTE_PEDAGOGIA",
            "INDICADOR",
            "N_CURSOS",
            "MEDIA_CURSOS",
            "MEDIANA_CURSOS",
            "DP_CURSOS",
        ]
    ].reset_index(drop=True)


def tabela_processo_castanhal(processo: pd.DataFrame, n: int = 12) -> pd.DataFrame:
    trabalho = processo.loc[processo["REFERENCIA"].eq("UFPA — Conceito 4")].copy()
    if trabalho.empty:
        return trabalho
    trabalho["ABS_DIF"] = pd.to_numeric(
        trabalho["DIFERENCA_CASTANHAL_REFERENCIA"], errors="coerce"
    ).abs()
    trabalho["ROTULO_OFICIAL"] = trabalho["ITEM"].map(rotulo_item)
    colunas = [
        "ITEM",
        "ROTULO_OFICIAL",
        "MEDIA_CASTANHAL",
        "N_VALIDO_CASTANHAL",
        "MEDIA_REFERENCIA",
        "N_CURSOS_REFERENCIA",
        "DIFERENCA_CASTANHAL_REFERENCIA",
    ]
    return (
        trabalho.sort_values("ABS_DIF", ascending=False)
        .head(n)[colunas]
        .reset_index(drop=True)
    )


def resumo_panorama(base: pd.DataFrame) -> str:
    ufpa = tabela_ofertas_ufpa(base)
    conceitos = sorted(
        pd.to_numeric(ufpa["CONCEITO_ENADE_NUM"], errors="coerce")
        .dropna()
        .astype(int)
        .unique()
        .tolist()
    )
    return (
        f"O universo analítico reúne {len(base)} cursos de Pedagogia identificados por CO_CURSO. "
        f"Foram localizadas {len(ufpa)} ofertas da UFPA, com Conceitos Enade {conceitos}. "
        "Não existe oferta da UFPA com Conceito Enade 1 em Pedagogia. Por isso, o Grupo A "
        "permanece vazio e não é reconstruído por aproximação. Castanhal, Conceito 5, é usada "
        "como referência interna descritiva diante das seis ofertas Conceito 4, sem interpretar "
        "Conceito 4 como insuficiência e sem atribuir causalidade ao conceito."
    )


def resumo_castanhal(base: pd.DataFrame, contraste: pd.DataFrame) -> str:
    c = obter_castanhal(base)
    texto = (
        f"Castanhal (CO_CURSO={CO_CURSO_CASTANHAL}) apresenta NT_GER médio {_fmt(c.get('nt_ger_mean'))}, "
        f"NT_OBJ {_fmt(c.get('nt_obj_mean'))}, NT_DIS {_fmt(c.get('nt_dis_mean'))}, "
        f"taxa oficial de participação {_fmt(c.get('TAXA_PARTICIPACAO_OFICIAL'))}% "
        f"e presença observada nos microdados {_fmt(c.get('taxa_presenca_microdados'))}%."
    )
    linha = contraste.loc[contraste["INDICADOR"].eq("nt_ger_mean")]
    if not linha.empty:
        r = linha.iloc[0]
        texto += (
            " No contraste Castanhal menos média das seis ofertas Conceito 4 da UFPA, "
            f"a diferença de NT_GER é {_fmt(r.get('DIFERENCA'))} ponto(s), "
            f"com z descritivo {_fmt(r.get('Z_DESCRITIVO'))}."
        )
    texto += (
        " A comparação é entre cursos e não constitui teste causal nem demonstra que "
        "o Conceito 5 produza diferenças de desempenho."
    )
    return texto


def resumo_regional(comparacoes: pd.DataFrame) -> str:
    tab = tabela_regional_nt_ger(comparacoes)
    if tab.empty:
        return "As comparações regionais não estavam disponíveis para síntese textual."
    recortes = ["UFPA agregada", "Região Norte sem UFPA", "Brasil geral"]
    partes: list[str] = []
    for recorte in recortes:
        sub = tab.loc[tab["RECORTE"].eq(recorte)]
        if sub.empty:
            continue
        linha = sub.iloc[0]
        partes.append(
            f"{recorte}: média simples {_fmt(linha.get('MEDIA_CURSOS'))} e média "
            f"ponderada por participantes {_fmt(linha.get('MEDIA_PONDERADA_PARTICIPANTES'))}"
        )
    if not partes:
        return "Os recortes de síntese regional não foram localizados."
    return "; ".join(partes) + "."


def resumo_benchmarks(sensibilidade: pd.DataFrame) -> str:
    principal = tabela_benchmark_principal(sensibilidade)
    if principal.empty:
        return "O benchmark estrutural principal não estava disponível."
    dif = pd.to_numeric(principal["nt_ger_mean_DIFERENCA"], errors="coerce")
    acima = int((dif > 0).sum())
    abaixo = int((dif < 0).sum())
    return (
        "No cenário estrutural principal — mesma modalidade, categoria administrativa, "
        "organização acadêmica e porte entre 0,5x e 2x participantes — "
        f"{acima} oferta(s) da UFPA apresentam NT_GER médio acima do respectivo benchmark "
        f"e {abaixo} abaixo. A Sprint 14 calculou cinco cenários por oferta, totalizando "
        "35 combinações oferta-cenário; a estabilidade entre cenários deve ser considerada "
        "antes de qualquer leitura substantiva."
    )


def resumo_associacoes(associacoes: pd.DataFrame) -> str:
    validas = associacoes.dropna(subset=["SPEARMAN_RHO"]).copy()
    if validas.empty:
        return "Não houve associações ecológicas calculáveis."
    linha = validas.loc[validas["SPEARMAN_RHO"].abs().idxmax()]
    return (
        "Entre as associações ecológicas examinadas, a maior magnitude absoluta foi entre "
        f"{linha['INDICADOR_X']} e NT_GER médio "
        f"(Spearman rho={linha['SPEARMAN_RHO']:.3f}; N={int(linha['N_CURSOS'])} cursos). "
        "O coeficiente usa cursos como unidades, não estudantes; portanto, não representa "
        "associação individual e não sustenta inferência causal."
    )


def resumo_processo(processo: pd.DataFrame) -> str:
    tabela = tabela_processo_castanhal(processo, n=6)
    if tabela.empty:
        return "Não houve comparação item a item disponível para o processo formativo."
    neg = tabela.sort_values("DIFERENCA_CASTANHAL_REFERENCIA").head(2)
    pos = tabela.sort_values(
        "DIFERENCA_CASTANHAL_REFERENCIA", ascending=False
    ).head(2)
    texto_neg = "; ".join(
        f"{r.ITEM} ({r.ROTULO_OFICIAL}): {r.DIFERENCA_CASTANHAL_REFERENCIA:+.2f}"
        for r in neg.itertuples()
    )
    texto_pos = "; ".join(
        f"{r.ITEM} ({r.ROTULO_OFICIAL}): {r.DIFERENCA_CASTANHAL_REFERENCIA:+.2f}"
        for r in pos.itertuples()
    )
    return (
        "Na comparação de Castanhal com as ofertas UFPA Conceito 4, entre os itens com maiores "
        f"diferenças absolutas, aparecem diferenças negativas em {texto_neg} e positivas em "
        f"{texto_pos}. Os textos oficiais do Inep são usados para interpretação, mas QE_I20–QE_I66 "
        "permanecem item a item; não é criado índice único sem validação teórica adicional."
    )


def _p(doc: Document, texto: str, estilo: str | None = None):
    par = doc.add_paragraph(style=estilo)
    par.add_run(texto)
    return par


def _capa(doc: Document) -> None:
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = 0
    r = p.add_run("UNIVERSIDADE FEDERAL DO PARÁ")
    r.bold = True
    r.font.size = Pt(12)
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = 0
    r = p.add_run(
        "PEDAGOGIA NO ENADE 2025: PANORAMA DA UFPA, BENCHMARKS E "
        "CONTRASTE INTERNO DAS OFERTAS"
    )
    r.bold = True
    r.font.size = Pt(14)
    for _ in range(7):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = 0
    p.add_run("Autor(a): [preencher]")
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = 0
    p.add_run("Belém\n2026")


def _figura(
    doc: Document,
    base_projeto: Path,
    nome: str,
    numero: int,
    titulo: str,
    leitura: str,
    hipotese: str,
    limitacao: str,
) -> None:
    adicionar_figura(
        doc,
        base_projeto / "figuras" / "pedagogia" / nome,
        f"Figura {numero} – {titulo}",
        FONTE_DADOS,
    )
    _p(doc, f"Descrição e interpretação: {leitura}")
    _p(doc, f"Hipótese analítica: {hipotese}")
    _p(doc, f"Limitação: {limitacao}")


def _markdown(dados: dict[str, pd.DataFrame]) -> str:
    base = dados["base"]
    ofertas = tabela_ofertas_ufpa(base)
    regional = tabela_regional_nt_ger(dados["comparacoes"])
    benchmark = tabela_benchmark_principal(dados["sensibilidade"])
    perfil = tabela_perfil_interno(dados["perfil"])
    processo = tabela_processo_castanhal(dados["processo"])

    linhas = [
        "# PEDAGOGIA NO ENADE 2025: PANORAMA DA UFPA, BENCHMARKS E CONTRASTE INTERNO DAS OFERTAS",
        "",
        "## RESUMO",
        "",
        f"Este relatório técnico-científico analisa {len(base)} cursos de Pedagogia presentes nos "
        f"microdados do Enade das Licenciaturas 2025, com {len(ofertas)} ofertas localizadas da UFPA. "
        "A unidade principal é CO_CURSO. Os arquivos temáticos foram processados separadamente, "
        "tratados quanto a ausências, agregados por curso e somente então relacionados por junções "
        "one-to-one. Não existe oferta da UFPA com Conceito Enade 1 em Pedagogia; o Grupo A permanece "
        "vazio. O contraste institucional examina Castanhal, Conceito 5, frente às seis ofertas "
        "Conceito 4, sem tratar Conceito 4 como insuficiência. São analisados desempenho, participação, "
        "perfil discente, trajetória, processo formativo, recomendação, benchmarks estruturais, "
        "comparações regionais e associações ecológicas. As interpretações são descritivas e não causais.",
        "",
        "**Palavras-chave:** Enade; Pedagogia; UFPA; formação de professores; microdados; benchmark.",
        "",
        "## ABSTRACT",
        "",
        "This technical-scientific report analyzes Pedagogy teacher education programs in the 2025 "
        "Enade, focusing on seven UFPA offers. CO_CURSO is the main unit of analysis. Thematic files "
        "were processed separately, aggregated at course level and only then combined one-to-one. "
        "UFPA has no Pedagogy offer with Enade Concept 1; therefore, no artificial Concept-1 group is "
        "created. Castanhal (Concept 5) is compared descriptively with six Concept-4 UFPA offers. "
        "Performance, participation, student profile, academic trajectory, formative process, "
        "recommendation, structural benchmarks, regional comparisons and ecological associations "
        "are examined. Results are descriptive and non-causal.",
        "",
        "**Keywords:** Enade; Pedagogy; UFPA; teacher education; microdata; benchmark.",
        "",
        "# 1 INTRODUÇÃO",
        "",
        "A pergunta central é: quais características de desempenho, participação, composição discente, "
        "trajetória acadêmica, condições socioeconômicas e avaliação do processo formativo diferenciam "
        "as ofertas de Pedagogia da UFPA entre si e em relação a cursos comparáveis no Pará, na Região "
        "Norte e no Brasil?",
        "",
        "# 2 REFERENCIAL INSTITUCIONAL E METODOLÓGICO",
        "",
        "O Conceito Enade é tratado como classificação externa do curso. Ausência de conceito não é "
        "recodificada. Em Pedagogia, não existe oferta UFPA com Conceito Enade 1. O contraste interno "
        "Castanhal Conceito 5 versus seis ofertas Conceito 4 é descritivo e não representa um desenho "
        "causal nem transforma Conceito 4 em categoria de insuficiência.",
        "",
        "# 3 METODOLOGIA",
        "",
        "A unidade principal de análise é CO_CURSO. Não se usa posição de linha como chave, não se cria "
        "identificador artificial e não são realizadas junções individuais entre arquivos temáticos. "
        "O fluxo é arquivo temático → tratamento de ausentes → agregação por CO_CURSO → uma linha por "
        "curso → junções one-to-one → comparação entre cursos. Relações entre temas distintos são "
        "ecológicas e recebem ressalva explícita de falácia ecológica.",
        "",
        "Os grupos exclusivos são: UFPA Conceito 5; UFPA Conceito 4; outras IES do Pará; Norte sem Pará; "
        "e Brasil sem Norte. Comparações completas de Pará, Norte e Brasil são usadas como benchmarks "
        "descritivos, não como grupos independentes em testes. Para cada oferta da UFPA foram calculados "
        "cinco cenários de benchmark. O principal exige mesma modalidade, categoria administrativa e "
        "organização acadêmica, com porte entre 0,5x e 2x participantes.",
        "",
        "# 4 PANORAMA DA LICENCIATURA EM PEDAGOGIA",
        "",
        resumo_panorama(base),
        "",
        "## Tabela 1 – Ofertas de Pedagogia da UFPA",
        "",
        ofertas.to_markdown(index=False, floatfmt=".3f"),
        "",
        "# 5 RESULTADOS",
        "",
        "## 5.1 Desempenho",
        "",
        resumo_castanhal(base, dados["contraste"]),
        "",
        "## 5.2 Perfil demográfico e socioeconômico",
        "",
        "Os indicadores são calculados por curso sobre respostas válidas, com N e ausências preservados. "
        "Relações entre perfil e desempenho são examinadas apenas ecologicamente.",
        "",
        perfil.to_markdown(index=False, floatfmt=".3f"),
        "",
        "## 5.3 Trajetória e condições acadêmicas",
        "",
        "A trajetória inclui turno, tempo desde o ingresso, trabalho, horas de estudo, bolsas e auxílios, "
        "sempre agregados por CO_CURSO. Diferenças entre ofertas podem refletir composição e contexto "
        "institucional não observados.",
        "",
        "## 5.4 Processo formativo",
        "",
        resumo_processo(dados["processo"]),
        "",
        processo.to_markdown(index=False, floatfmt=".3f"),
        "",
        "## 5.5 Recomendação",
        "",
        "QE_I68 representa recomendação do curso, QE_I69 recomendação da IES e QE_I70 interesse em "
        "participar da Prova Nacional Docente de 2025. Esses itens são apresentados com seus rótulos "
        "oficiais e não são automaticamente denominados satisfação.",
        "",
        "## 5.6 Benchmark comparável",
        "",
        resumo_benchmarks(dados["sensibilidade"]),
        "",
        benchmark.to_markdown(index=False, floatfmt=".3f"),
        "",
        "## 5.7 Associações ecológicas",
        "",
        resumo_associacoes(dados["associacoes"]),
        "",
        dados["associacoes"].to_markdown(index=False, floatfmt=".3f"),
        "",
        "## 5.8 Comparações regionais e nacionais",
        "",
        resumo_regional(dados["comparacoes"]),
        "",
        regional.to_markdown(index=False, floatfmt=".3f"),
        "",
        "## 5.9 Contraste interno das ofertas da UFPA",
        "",
        "### 5.9.1 Participação e desempenho",
        "",
        resumo_castanhal(base, dados["contraste"]),
        "",
        "### 5.9.2 Perfil discente",
        "",
        "O contraste interno apresenta Castanhal e o conjunto das seis ofertas Conceito 4 sem assumir "
        "equivalência estrutural perfeita entre os cursos.",
        "",
        "### 5.9.3 Trajetória acadêmica",
        "",
        "Turno, tempo desde o ingresso e dedicação acadêmica são interpretados como características "
        "agregadas das ofertas, não como atributos causais individuais.",
        "",
        "### 5.9.4 Processo formativo",
        "",
        resumo_processo(dados["processo"]),
        "",
        "### 5.9.5 Recomendação",
        "",
        "A recomendação do curso e da IES é analisada separadamente. Diferenças entre Castanhal e as "
        "demais ofertas não são interpretadas como satisfação global nem como mecanismo causal.",
        "",
        "### 5.9.6 Benchmark comparável",
        "",
        resumo_benchmarks(dados["sensibilidade"]),
        "",
        "### 5.9.7 Síntese dos pontos distintivos",
        "",
        "O contraste interno deve ser lido em conjunto com participação, desempenho, perfil, processo "
        "formativo, recomendação e sensibilidade dos benchmarks. Nenhuma dimensão isolada explica o "
        "Conceito Enade observado.",
        "",
        "# 6 DISCUSSÃO",
        "",
        "Os resultados mostram heterogeneidade entre as sete ofertas da UFPA e entre as referências "
        "territoriais e estruturais. Castanhal, única oferta Conceito 5, constitui um caso interno de "
        "contraste, mas a existência de apenas uma unidade nesse estrato impede tratá-la como população "
        "independente para inferência. As diferenças observadas devem ser interpretadas como padrões e "
        "hipóteses para investigação institucional. Benchmarks reduzem parte da heterogeneidade observável, "
        "mas não controlam todos os fatores relevantes.",
        "",
        "# 7 CONCLUSÃO",
        "",
        "A análise de Pedagogia preserva a ausência de ofertas UFPA Conceito 1 e desloca a pergunta para "
        "a heterogeneidade interna e a posição relativa das sete ofertas. O conjunto de evidências permite "
        "identificar dimensões em que Castanhal e as ofertas Conceito 4 se aproximam ou se distinguem, "
        "sempre com N, dispersão, ausências e sensibilidade de benchmark. Não há base para atribuir "
        "causalidade individual, territorial ou institucional aos padrões observados.",
        "",
        "# REFERÊNCIAS",
        "",
        "As referências bibliográficas e normativas são inseridas integralmente na versão DOCX pelo "
        "módulo compartilhado de referências do projeto.",
        "",
        "# APÊNDICES",
        "",
        "## APÊNDICE A – REGRAS DE INTEGRIDADE",
        "",
        "A unidade principal é CO_CURSO; não há join individual entre arquivos temáticos; as integrações "
        "ocorrem somente após agregação por curso e validação one-to-one; ausência de conceito não é "
        "Conceito 1; comparações entre temas são ecológicas; QE_I20–QE_I66 não são condensados em índice "
        "único sem validação; notas brutas não são comparadas entre áreas.",
        "",
        "## APÊNDICE B – APROFUNDAMENTOS SUGERIDOS",
        "",
        "1. Robustez do contraste Castanhal versus UFPA Conceito 4 — pergunta: a posição relativa de "
        "Castanhal se mantém após reponderação ou pareamento mais estrito? Variáveis: modalidade, categoria "
        "administrativa, organização acadêmica, participantes, NT_GER, NT_OBJ e NT_DIS. Método: matching "
        "ecológico e análise de sensibilidade. Limitação: pequeno número de ofertas UFPA.",
        "2. Componentes do desempenho — pergunta: quais dimensões de NT_OBJ, NT_DIS, QT_ACERTOS e PROFICIENCIA "
        "concentram diferenças entre ofertas? Método: análise individual exclusivamente no arquivo de desempenho "
        "e comparação de distribuições por curso. Limitação: relações mecânicas entre indicadores.",
        "3. Processo formativo item a item — pergunta: quais itens QE_I20–QE_I66 apresentam contrastes robustos "
        "após considerar N válido e múltiplas referências? Método: diferenças por item, intervalos e sensibilidade. "
        "Limitação: escala ordinal e respostas autorreferidas.",
        "4. Perfil, permanência e trajetória — pergunta: quais combinações agregadas de renda, trabalho, bolsas, "
        "auxílios, horas de estudo e trajetória distinguem ofertas? Método: perfis padronizados e associações "
        "ecológicas. Limitação: falácia ecológica e confundimento residual.",
        "5. Recomendação e processo formativo — pergunta: no universo nacional de Pedagogia, recomendação do curso "
        "e da IES se associa ecologicamente a dimensões formativas específicas? Variáveis: QE_I68, QE_I69 e "
        "QE_I20–QE_I66. Método: Spearman por curso, outliers e ponderação opcional por participantes. Limitação: "
        "não representa relação individual entre respostas.",
        "",
        "## APÊNDICE C – RÓTULOS OFICIAIS DOS ITENS QE_I20–QE_I66",
        "",
    ]
    for codigo in [f"QE_I{i}" for i in range(20, 67)]:
        linhas.append(f"- **{codigo}** — {ROTULOS_QE[codigo]}")
    linhas.extend(["", f"Fonte dos rótulos: {FONTE_QUESTIONARIO}"])
    return "\n".join(linhas)


def gerar_relatorio(
    base_projeto: Path,
    saida_docx: Path,
    saida_md: Path,
) -> dict:
    dados = carregar_produtos(base_projeto)
    base = dados["base"]
    ofertas = tabela_ofertas_ufpa(base)
    regional = tabela_regional_nt_ger(dados["comparacoes"])
    benchmark = tabela_benchmark_principal(dados["sensibilidade"])
    perfil = tabela_perfil_interno(dados["perfil"])
    processo = tabela_processo_castanhal(dados["processo"])

    doc = Document()
    configurar_documento(doc)
    _capa(doc)
    doc.add_page_break()

    doc.add_heading("RESUMO", level=1)
    _p(
        doc,
        f"Este relatório analisa {len(base)} cursos de Pedagogia presentes nos microdados do Enade "
        f"das Licenciaturas 2025, com {len(ofertas)} ofertas da UFPA. A unidade principal é CO_CURSO. "
        "Os arquivos temáticos foram processados separadamente, agregados por curso e somente então "
        "relacionados por junções one-to-one. Não existe oferta da UFPA com Conceito Enade 1 em "
        "Pedagogia. Castanhal, Conceito 5, é comparada descritivamente às seis ofertas Conceito 4. "
        "São examinados desempenho, participação, perfil discente, trajetória, processo formativo, "
        "recomendação, benchmarks, comparações regionais e associações ecológicas. A interpretação "
        "é descritiva e não causal.",
        "Resumo",
    )
    _p(
        doc,
        "Palavras-chave: Enade; Pedagogia; UFPA; formação de professores; microdados; benchmark.",
    )

    doc.add_heading("ABSTRACT", level=1)
    _p(
        doc,
        "This technical-scientific report analyzes Pedagogy teacher education programs in the 2025 "
        "Enade, focusing on seven UFPA offers. CO_CURSO is the main unit of analysis. Thematic files "
        "were processed separately, aggregated at course level and combined one-to-one. UFPA has no "
        "Pedagogy offer with Enade Concept 1. Castanhal (Concept 5) is descriptively compared with "
        "six Concept-4 UFPA offers. Results are descriptive and non-causal.",
        "Resumo",
    )
    _p(doc, "Keywords: Enade; Pedagogy; UFPA; teacher education; microdata; benchmark.")

    doc.add_heading("1 INTRODUÇÃO", level=1)
    _p(
        doc,
        "Este relatório responde à pergunta: quais características de desempenho, participação, "
        "composição discente, trajetória acadêmica, condições socioeconômicas e avaliação do processo "
        "formativo diferenciam as ofertas de Pedagogia da UFPA entre si e em relação a cursos "
        "comparáveis no Pará, na Região Norte e no Brasil?",
    )

    doc.add_heading("2 REFERENCIAL INSTITUCIONAL E METODOLÓGICO", level=1)
    _p(
        doc,
        "O Conceito Enade é tratado como classificação externa do curso e não como variável causal. "
        "Ausência de conceito permanece distinta de Conceito 1. Em Pedagogia não existe oferta UFPA "
        "com Conceito Enade 1. O contraste entre Castanhal, Conceito 5, e seis ofertas Conceito 4 é "
        "descritivo; Conceito 4 não é interpretado como insuficiência.",
    )

    doc.add_heading("3 METODOLOGIA", level=1)
    _p(
        doc,
        "A unidade principal é CO_CURSO. Cada arquivo temático foi tratado separadamente, com tratamento "
        "de ausências e agregação por curso antes de qualquer junção. Não foi usada posição de linha, "
        "não foi criado identificador artificial e não foram realizadas junções individuais entre temas. "
        "As integrações das tabelas agregadas são one-to-one. Relações entre arquivos distintos são "
        "ecológicas e recebem ressalva de falácia ecológica.",
    )
    _p(
        doc,
        "Os grupos exclusivos usados na análise são UFPA Conceito 5, UFPA Conceito 4, outras IES do "
        "Pará, Norte sem Pará e Brasil sem Norte. Para cada oferta UFPA são calculados cinco cenários "
        "de benchmark. O cenário principal preserva modalidade, categoria administrativa, organização "
        "acadêmica e porte de participantes entre 0,5x e 2x. Médias simples e ponderadas por participantes "
        "são apresentadas nos recortes territoriais.",
    )

    doc.add_heading("4 PANORAMA DA LICENCIATURA EM PEDAGOGIA", level=1)
    _p(doc, resumo_panorama(base))
    adicionar_tabela(doc, "Tabela 1 – Ofertas de Pedagogia da UFPA", ofertas, FONTE_DADOS)
    _figura(
        doc, base_projeto, "01_painel_ofertas_ufpa.png", 1,
        "Ofertas de Pedagogia da UFPA",
        "O painel reúne as sete ofertas, conceitos, porte e indicadores iniciais.",
        "Diferenças de porte, participação e composição podem contribuir para a heterogeneidade observada.",
        "O painel é descritivo e não controla diferenças estruturais entre ofertas.",
    )

    doc.add_heading("5 RESULTADOS", level=1)

    doc.add_heading("5.1 Desempenho", level=2)
    _p(doc, resumo_castanhal(base, dados["contraste"]))
    _figura(
        doc, base_projeto, "02_posicao_relativa_nt_ger.png", 2,
        "Posição relativa das ofertas em NT_GER",
        "As ofertas são situadas na distribuição nacional da própria área.",
        "A posição relativa pode ajudar a separar diferenças absolutas de padrões próprios da área.",
        "Percentis não eliminam diferenças de participação, porte ou composição.",
    )
    _figura(
        doc, base_projeto, "03_distribuicao_nt_ger.png", 3,
        "Distribuição de NT_GER",
        "A distribuição individual é construída apenas com variáveis do arquivo de desempenho.",
        "Diferenças de centro e dispersão entre ofertas podem indicar perfis de desempenho distintos.",
        "Cursos com N menor apresentam maior instabilidade amostral.",
    )
    _figura(
        doc, base_projeto, "04_distribuicao_nt_obj.png", 4,
        "Distribuição de NT_OBJ",
        "O componente objetivo é analisado separadamente da nota geral.",
        "Uma diferença mais intensa em NT_OBJ que em NT_DIS pode indicar contraste concentrado no componente objetivo.",
        "NT_OBJ e NT_GER têm relação mecânica e não são evidências independentes.",
    )
    _figura(
        doc, base_projeto, "05_distribuicao_nt_dis.png", 5,
        "Distribuição de NT_DIS",
        "O componente discursivo é preservado em sua escala própria.",
        "A comparação com NT_OBJ permite localizar em qual componente se concentram as diferenças.",
        "N válido e particularidades de correção devem ser considerados.",
    )
    _figura(
        doc, base_projeto, "09_desempenho_ofertas_ufpa.png", 6,
        "Desempenho das ofertas da UFPA",
        "NT_GER, NT_OBJ e NT_DIS são comparados entre as sete ofertas.",
        "A heterogeneidade interna pode não seguir estritamente a ordenação dos Conceitos Enade.",
        "O número reduzido de ofertas impede inferência estatística robusta entre estratos de conceito.",
    )
    _figura(
        doc, base_projeto, "10_percentis_ofertas_ufpa.png", 7,
        "Percentis das ofertas da UFPA",
        "Os percentis posicionam cada curso na distribuição de Pedagogia.",
        "Cursos de conceitos distintos podem apresentar sobreposição parcial nas posições relativas.",
        "Percentis são descritivos e dependem do universo de cursos válidos.",
    )

    doc.add_heading("5.2 Perfil demográfico e socioeconômico", level=2)
    _p(
        doc,
        "Sexo, idade, escolaridade parental, renda, trabalho, ação afirmativa, bolsas, auxílios e "
        "horas de estudo são agregados por CO_CURSO. Percentuais são calculados sobre respostas válidas; "
        "ausências permanecem documentadas nos produtos analíticos.",
    )
    adicionar_tabela(doc, "Tabela 2 – Perfil interno da UFPA por estrato de conceito", perfil, FONTE_DADOS)
    _figura(
        doc, base_projeto, "06_perfil_socioeconomico.png", 8,
        "Perfil socioeconômico das ofertas",
        "A figura compara indicadores selecionados de composição discente.",
        "Diferenças de composição podem coexistir com diferenças de desempenho sem relação causal individual.",
        "Os indicadores são agregados por curso e estão sujeitos à falácia ecológica.",
    )
    _figura(
        doc, base_projeto, "validada_18_perfil_interno_ufpa.png", 9,
        "Perfil interno validado da UFPA",
        "O contraste resume Castanhal e o conjunto das seis ofertas Conceito 4.",
        "Características de composição podem orientar aprofundamentos sobre permanência e trajetória.",
        "Castanhal representa uma única oferta no estrato Conceito 5.",
    )

    doc.add_heading("5.3 Trajetória e condições acadêmicas", level=2)
    _p(
        doc,
        "Turno, anos desde o ingresso, trabalho, bolsas, auxílios e dedicação aos estudos são "
        "interpretados como características agregadas das ofertas. Relações com desempenho, quando "
        "envolvem arquivos distintos, são examinadas somente no nível ecológico do curso.",
    )

    doc.add_heading("5.4 Processo formativo", level=2)
    _p(doc, resumo_processo(dados["processo"]))
    adicionar_tabela(
        doc,
        "Tabela 3 – Itens QE_I20–QE_I66 de maior diferença entre Castanhal e UFPA Conceito 4",
        processo,
        FONTE_QUESTIONARIO,
    )
    _figura(
        doc, base_projeto, "07_processo_formativo.png", 10,
        "Itens do processo formativo",
        "QE_I20–QE_I66 são examinados item a item, preservando a escala oficial.",
        "Padrões por item podem sugerir dimensões para validação teórica posterior.",
        "Não é criado índice único sem validação de escala, dimensionalidade e consistência.",
    )
    _figura(
        doc, base_projeto, "validada_17_processo_castanhal.png", 11,
        "Processo formativo: Castanhal versus UFPA Conceito 4",
        "A figura destaca itens com maiores diferenças absolutas no contraste interno.",
        "Itens com diferenças persistentes em múltiplas referências merecem aprofundamento institucional.",
        "A escala é ordinal e as respostas são autorreferidas.",
    )

    doc.add_heading("5.5 Recomendação", level=2)
    _p(
        doc,
        f"{rotulo_item('QE_I68')} {rotulo_item('QE_I69')} {rotulo_item('QE_I70')} "
        "Os três indicadores são tratados de acordo com seus rótulos oficiais e não são reunidos "
        "automaticamente sob o termo satisfação.",
    )
    _figura(
        doc, base_projeto, "12_recomendacao.png", 12,
        "Recomendação do curso, da IES e interesse na PND",
        "A figura preserva a distinção entre QE_I68, QE_I69 e QE_I70.",
        "Recomendação pode acompanhar dimensões do processo formativo, hipótese examinável apenas com desenho adequado.",
        "Não se deve interpretar recomendação como medida geral de satisfação.",
    )

    doc.add_heading("5.6 Benchmark comparável", level=2)
    _p(doc, resumo_benchmarks(dados["sensibilidade"]))
    adicionar_tabela(
        doc,
        "Tabela 4 – Benchmark estrutural principal das sete ofertas da UFPA",
        benchmark,
        FONTE_DADOS,
    )
    _figura(
        doc, base_projeto, "11_benchmarks_ofertas_ufpa.png", 13,
        "Benchmarks comparáveis por oferta",
        "Cada oferta UFPA é comparada a cursos externos com características estruturais selecionadas.",
        "A redução da heterogeneidade observável pode alterar a magnitude das diferenças brutas.",
        "O benchmark não controla confundidores não observados e não constitui pareamento causal.",
    )
    _figura(
        doc, base_projeto, "validada_16_benchmarks_por_oferta.png", 14,
        "Sensibilidade dos benchmarks",
        "A figura compara resultados entre cenários de benchmark.",
        "Diferenças estáveis entre cenários são mais robustas descritivamente que diferenças dependentes de uma única definição.",
        "Alguns cenários podem conter poucos cursos, aumentando a incerteza.",
    )

    doc.add_heading("5.7 Associações ecológicas", level=2)
    _p(doc, resumo_associacoes(dados["associacoes"]))
    adicionar_tabela(
        doc,
        "Tabela 5 – Associações ecológicas com NT_GER médio",
        dados["associacoes"],
        FONTE_DADOS,
    )

    doc.add_heading("5.8 Comparações regionais e nacionais", level=2)
    _p(doc, resumo_regional(dados["comparacoes"]))
    adicionar_tabela(
        doc,
        "Tabela 6 – NT_GER por recortes regionais e nacionais",
        regional,
        FONTE_DADOS,
    )
    _figura(
        doc, base_projeto, "08_comparacao_regional_nacional.png", 15,
        "Comparação regional e nacional",
        "As médias das ofertas e as médias ponderadas por participantes são apresentadas por recorte.",
        "Diferenças entre média simples e ponderada podem indicar influência do porte dos cursos.",
        "Recortes territoriais completos se sobrepõem e não são tratados como grupos independentes em testes.",
    )
    _figura(
        doc, base_projeto, "validada_19_comparacao_regional.png", 16,
        "Comparação regional validada",
        "A versão validada acrescenta dispersão e auditoria dos recortes obrigatórios.",
        "A posição da UFPA deve ser interpretada no contexto da dispersão entre cursos.",
        "Médias agregadas podem ocultar heterogeneidade interna.",
    )

    doc.add_heading("5.9 Contraste interno das ofertas da UFPA", level=2)

    doc.add_heading("5.9.1 Participação e desempenho", level=3)
    _p(doc, resumo_castanhal(base, dados["contraste"]))
    _figura(
        doc, base_projeto, "validada_14_participacao_ufpa.png", 17,
        "Participação das ofertas da UFPA",
        "Inscritos, participantes e presença são comparados entre as sete ofertas.",
        "Diferenças de participação podem afetar estabilidade e representatividade das estimativas por curso.",
        "Participação não deve ser interpretada isoladamente como explicação do desempenho.",
    )
    _figura(
        doc, base_projeto, "validada_15_contraste_interno_ufpa.png", 18,
        "Contraste interno Castanhal versus UFPA Conceito 4",
        "A figura reúne indicadores do contraste descritivo interno.",
        "Castanhal pode apresentar vantagens em algumas dimensões e proximidade ou desvantagens em outras.",
        "Há apenas uma oferta no estrato Conceito 5, impedindo inferência entre grupos.",
    )
    _figura(
        doc, base_projeto, "13_contraste_interno_ufpa.png", 19,
        "Síntese do contraste interno das ofertas",
        "A figura integra indicadores selecionados das ofertas da UFPA.",
        "A leitura multidimensional evita reduzir a comparação ao Conceito Enade.",
        "Indicadores de arquivos distintos são comparáveis somente no nível agregado do curso.",
    )

    doc.add_heading("5.9.2 Perfil discente", level=3)
    _p(
        doc,
        "Castanhal é comparada ao conjunto das seis ofertas Conceito 4 em indicadores de perfil. "
        "As diferenças representam composição agregada e não permitem inferir como características "
        "individuais se relacionam ao desempenho de estudantes específicos.",
    )

    doc.add_heading("5.9.3 Trajetória acadêmica", level=3)
    _p(
        doc,
        "Indicadores de turno, tempo desde o ingresso, trabalho, bolsas, auxílios e horas de estudo "
        "são usados para caracterizar as ofertas. Eles podem motivar hipóteses institucionais, mas não "
        "estabelecem mecanismos causais para o Conceito Enade.",
    )

    doc.add_heading("5.9.4 Processo formativo", level=3)
    _p(doc, resumo_processo(dados["processo"]))

    doc.add_heading("5.9.5 Recomendação", level=3)
    _p(
        doc,
        "QE_I68 e QE_I69 são apresentados separadamente como recomendação do curso e da IES. "
        "QE_I70 representa interesse na Prova Nacional Docente. Nenhum deles é convertido "
        "automaticamente em índice de satisfação.",
    )

    doc.add_heading("5.9.6 Benchmark comparável", level=3)
    _p(doc, resumo_benchmarks(dados["sensibilidade"]))

    doc.add_heading("5.9.7 Síntese dos pontos distintivos", level=3)
    _p(
        doc,
        "O contraste interno mostra que a interpretação de Castanhal exige combinar participação, "
        "desempenho, composição discente, trajetória, processo formativo, recomendação e posição "
        "diante dos benchmarks. O Conceito Enade é um marcador externo do curso, não uma explicação "
        "para as diferenças observadas.",
    )

    doc.add_heading("6 DISCUSSÃO", level=1)
    _p(
        doc,
        "A principal característica do caso de Pedagogia na UFPA é a ausência de ofertas Conceito 1. "
        "Isso impede reproduzir o contraste originalmente previsto para outras licenciaturas e exige "
        "uma leitura da heterogeneidade interna. Castanhal, única oferta Conceito 5, oferece uma "
        "referência institucional útil, porém não constitui sozinha um grupo inferencial. A comparação "
        "com seis ofertas Conceito 4, combinada a referências territoriais e benchmarks estruturais, "
        "permite formular hipóteses sobre dimensões de desempenho, participação, composição e formação "
        "que merecem aprofundamento. Tais hipóteses não autorizam causalidade individual, de campus ou "
        "de modalidade.",
    )

    doc.add_heading("7 CONCLUSÃO", level=1)
    _p(
        doc,
        "O relatório caracteriza as sete ofertas de Pedagogia da UFPA e sua posição no Pará, na "
        "Região Norte e no Brasil. A estratégia preserva a unidade CO_CURSO, a separação dos arquivos "
        "temáticos e a natureza ecológica das relações entre dimensões. Os resultados devem ser usados "
        "para priorizar investigação institucional e orientar análises posteriores, não para concluir "
        "que o Conceito 5 de Castanhal seja causado por uma característica isolada ou que as ofertas "
        "Conceito 4 sejam insuficientes.",
    )

    doc.add_heading("REFERÊNCIAS", level=1)
    adicionar_referencias(doc)

    doc.add_heading("APÊNDICE A – REGRAS DE INTEGRIDADE", level=1)
    _p(
        doc,
        "Não se usa posição de linha como chave; não se cria identificador artificial; não há join "
        "individual entre desempenho, perfil e percepção; todos os arquivos temáticos são agregados "
        "por CO_CURSO antes das integrações; as junções agregadas são one-to-one; ausência de conceito "
        "não equivale a Conceito 1; relações entre temas distintos são ecológicas; QE_I20–QE_I66 não "
        "formam índice único sem validação; notas brutas não são comparadas entre áreas.",
    )

    doc.add_heading("APÊNDICE B – APROFUNDAMENTOS SUGERIDOS", level=1)
    aprofundamentos = [
        "1. Robustez do contraste Castanhal versus UFPA Conceito 4. Pergunta: a posição relativa de "
        "Castanhal permanece após pareamento mais estrito? Variáveis: modalidade, categoria, organização "
        "acadêmica, participantes e notas. Método: matching ecológico e análise de sensibilidade. "
        "Limitação: pequeno número de ofertas UFPA.",
        "2. Componentes do desempenho. Pergunta: em quais componentes se concentram as diferenças entre "
        "ofertas? Variáveis: NT_GER, NT_OBJ, NT_DIS, QT_ACERTOS e PROFICIENCIA. Método: análise individual "
        "somente no arquivo de desempenho, com distribuições e tamanhos de efeito. Limitação: relações "
        "mecânicas entre indicadores.",
        "3. Processo formativo item a item. Pergunta: quais itens QE_I20–QE_I66 apresentam diferenças "
        "robustas em múltiplas referências? Método: diferenças por item, N válido, intervalos e sensibilidade. "
        "Limitação: escala ordinal, autorrelato e multiplicidade.",
        "4. Perfil, permanência e trajetória. Pergunta: quais combinações agregadas de renda, trabalho, "
        "ações afirmativas, auxílios, bolsas e horas de estudo distinguem as ofertas? Método: perfis "
        "padronizados e associações ecológicas. Limitação: falácia ecológica e confundimento residual.",
        "5. Recomendação e formação. Pergunta: QE_I68 e QE_I69 se associam ecologicamente a itens formativos "
        "no universo nacional de Pedagogia? Método: Spearman por curso, análise de outliers e ponderação "
        "opcional por participantes. Limitação: não representa relação individual.",
    ]
    for texto in aprofundamentos:
        _p(doc, texto)

    doc.add_heading("APÊNDICE C – RÓTULOS OFICIAIS DOS ITENS QE_I20–QE_I66", level=1)
    for codigo in [f"QE_I{i}" for i in range(20, 67)]:
        _p(doc, f"{codigo} — {ROTULOS_QE[codigo]}")
    _p(doc, f"Fonte: {FONTE_QUESTIONARIO}")

    configurar_cabecalho_rodape(doc)
    saida_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(saida_docx)

    saida_md.parent.mkdir(parents=True, exist_ok=True)
    saida_md.write_text(_markdown(dados), encoding="utf-8")

    conversao = converter_docx_para_pdf(saida_docx, saida_docx.parent)
    return ResultadoRelatorio(saida_docx, saida_md, conversao).como_dict()
