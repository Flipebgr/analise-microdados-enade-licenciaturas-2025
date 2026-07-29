from __future__ import annotations

import shutil
import subprocess
from datetime import datetime
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from src.relatorios.figuras_relatorio import adicionar_figura
from src.relatorios.formatacao_abnt import configurar_cabecalho_rodape, configurar_documento, nova_secao
from src.relatorios.gerar_apendices_fisica import exportar_apendices
from src.relatorios.gerar_tabelas_fisica import (
    associacoes_selecionadas,
    benchmark_sensibilidade,
    comparacao_territorial,
    desempenho_ufpa,
    dimensoes_processo,
    ofertas_ufpa,
    participacao,
    recomendacao_dificuldade,
    socioeconomico,
)
from src.relatorios.gerar_texto_fisica import (
    aprofundamentos,
    associacoes_texto,
    desempenho_texto,
    panorama_texto,
    resumo_executivo,
    socio_texto,
    tucurui_texto,
)
from src.relatorios.referencias import adicionar_referencias
from src.relatorios.tabelas_relatorio import adicionar_tabela
from src.relatorios.validar_relatorio_fisica import validar_relatorio


def _p(doc: Document, texto: str, estilo: str | None = None):
    p = doc.add_paragraph(style=estilo)
    p.add_run(texto)
    return p


def _h(doc: Document, texto: str, nivel: int = 1):
    return doc.add_heading(texto, level=nivel)


def _ler(base: Path) -> dict[str, pd.DataFrame]:
    p = base / "dados_processados" / "fisica"
    arquivos = {
        "cursos": "cursos_fisica.csv",
        "base": "base_analitica_cursos.csv",
        "presenca": "auditoria_presenca_validada.csv",
        "territorio": "comparacao_territorial_validada.csv",
        "sensibilidade": "sensibilidade_benchmarks.csv",
        "dimensoes": "diagnostico_dimensoes_processo.csv",
        "associacoes": "associacoes_ecologicas.csv",
        "socio": "tabela_socioeconomica_ufpa.csv",
    }
    dados: dict[str, pd.DataFrame] = {}
    for chave, nome in arquivos.items():
        caminho = p / nome
        if not caminho.exists():
            raise FileNotFoundError(f"Produto da Sprint 5 ausente: {caminho}")
        dados[chave] = pd.read_csv(caminho)
    return dados


def _capa(doc: Document, metadados: dict[str, str]) -> None:
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = 0
    r = p.add_run(metadados.get("instituicao", "UNIVERSIDADE FEDERAL DO PARÁ"))
    r.bold = True
    r.font.size = Pt(12)
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = 0
    r = p.add_run(
        "DESEMPENHO, COMPOSIÇÃO DISCENTE E PROCESSO FORMATIVO NAS OFERTAS "
        "DE LICENCIATURA EM FÍSICA DA UFPA NO ENADE 2025"
    )
    r.bold = True
    r.font.size = Pt(14)
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = 0
    p.add_run(metadados.get("autor", "Autor(a): [preencher]"))
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = 0
    p.add_run(f"{metadados.get('cidade', 'Belém')}\n2026")


def _adicionar_figura_com_leitura(
    doc: Document,
    caminho: Path,
    numero: int,
    titulo: str,
    leitura: str,
    limitacao: str,
) -> None:
    adicionar_figura(
        doc,
        caminho,
        f"Figura {numero} – {titulo}",
        "Elaboração própria com base nos microdados do Enade das Licenciaturas 2025 e na planilha de Conceito Enade.",
    )
    _p(doc, f"Descrição e interpretação: {leitura}")
    _p(doc, f"Limitação: {limitacao}")



def _situacao_tucurui() -> pd.DataFrame:
    return pd.DataFrame(
        {
            "Campo": [
                "CO_CURSO",
                "Oferta",
                "Localização na planilha de Conceito Enade",
                "Localização nos microdados",
                "Inscritos",
                "Participantes",
                "Presentes",
                "Percentual no padrão de proficiência",
                "Conceito Enade",
                "Classificação analítica",
            ],
            "Resultado": [
                "1627581",
                "Tucuruí",
                "Não localizada",
                "Não localizada",
                "Não disponível",
                "Não disponível",
                "Não disponível",
                "Não disponível",
                "Não disponível",
                "Oferta não localizada nas fontes; não classificada como Conceito 1",
            ],
        }
    )

def _markdown(dados: dict[str, pd.DataFrame], tabelas: dict[str, pd.DataFrame]) -> str:
    base = dados["base"]
    cursos = dados["cursos"]
    linhas = [
        "# DESEMPENHO, COMPOSIÇÃO DISCENTE E PROCESSO FORMATIVO NAS OFERTAS DE LICENCIATURA EM FÍSICA DA UFPA NO ENADE 2025",
        "",
        "## RESUMO",
        "",
        resumo_executivo(base),
        "",
        "**Palavras-chave:** Enade; licenciatura em Física; UFPA; microdados; avaliação da educação superior.",
        "",
        "## ABSTRACT",
        "",
        "This technical-scientific report analyzes Physics teacher education programs offered by the Federal University of Pará in the 2025 Enade for Teacher Education Programs. The course (CO_CURSO) is the main unit of analysis. Thematic files were processed separately and aggregated at course level, without reconstructing individual records. Performance, participation, socioeconomic composition, academic trajectory, formative process evaluation, recommendation, and broad and comparable benchmarks were examined. Associations are ecological and do not support individual or causal inference.",
        "",
        "**Keywords:** Enade; Physics teacher education; UFPA; microdata; higher education assessment.",
        "",
        "# 1 INTRODUÇÃO",
        "",
        "O relatório examina as ofertas de Física da UFPA com ênfase nos cursos classificados com Conceito Enade 1. A pergunta central é: quais características de desempenho, composição discente, trajetória acadêmica e avaliação do processo formativo diferenciam essas ofertas das demais ofertas da mesma área na UFPA, no Pará, na Região Norte e no Brasil?",
        "",
        "# 2 REFERENCIAL INSTITUCIONAL E METODOLÓGICO",
        "",
        "O Conceito Enade é tratado como classificação externa do curso. Os 28 arquivos temáticos não permitem reconhecer o mesmo estudante entre temas. A posição da linha não é uma chave. Cada arquivo foi tratado isoladamente, agregado por CO_CURSO e somente então relacionado às demais tabelas agregadas.",
        "",
        "# 3 METODOLOGIA",
        "",
        "A unidade principal é CO_CURSO. Foram construídos grupos exclusivos: A) UFPA conceito 1; B) demais ofertas da UFPA com conceito superior; C) outras IES do Pará; D) restante da Região Norte; E) restante do Brasil. Pará, Norte e Brasil completos foram empregados apenas como benchmarks descritivos sobrepostos.",
        "",
        "# 4 PANORAMA DA LICENCIATURA EM FÍSICA",
        "",
        panorama_texto(cursos),
        "",
        "## Tabela 1 – Ofertas da UFPA",
        "",
        tabelas["ofertas"].to_markdown(index=False, floatfmt=".2f"),
        "",
        "## 4.1 Situação da oferta de Tucuruí",
        "",
        tucurui_texto(),
        "",
        tabelas["tucurui"].to_markdown(index=False),
        "",
        "# 5 RESULTADOS",
        "",
        "## 5.1 Desempenho",
        "",
        desempenho_texto(base),
        "",
        tabelas["participacao"].to_markdown(index=False, floatfmt=".2f"),
        "",
        tabelas["desempenho"].to_markdown(index=False, floatfmt=".2f"),
        "",
        "## 5.2 Perfil demográfico e socioeconômico",
        "",
        socio_texto(dados["socio"]),
        "",
        tabelas["socio"].to_markdown(index=False, floatfmt=".2f"),
        "",
        "## 5.3 Trajetória e condições acadêmicas",
        "",
        "Os indicadores de trabalho, bolsas, auxílios, horas de estudo, turno e intenção de exercer o magistério foram agregados por curso. A leitura contextual deve ser combinada ao N válido e às ausências de cada item.",
        "",
        "## 5.4 Processo formativo",
        "",
        "As dimensões candidatas foram mantidas como exploratórias. Não foi criada uma média única de processo formativo. A inclusão no relatório depende de coerência teórica, orientação da escala, cobertura e consistência interna.",
        "",
        tabelas["dimensoes"].to_markdown(index=False, floatfmt=".3f"),
        "",
        "## 5.5 Recomendação",
        "",
        "Os itens QE_I68 e QE_I69 foram apresentados conforme sua função de recomendação. Eles não foram renomeados como satisfação geral. A percepção de dificuldade foi analisada no nível agregado do curso.",
        "",
        tabelas["recomendacao"].to_markdown(index=False, floatfmt=".2f"),
        "",
        "## 5.6 Benchmark comparável",
        "",
        "A sensibilidade foi examinada com diferentes critérios de porte e composição institucional. Resultados de conjuntos com poucos comparáveis devem ser tratados como descritivos.",
        "",
        tabelas["territorio"].to_markdown(index=False, floatfmt=".2f"),
        "",
        tabelas["benchmark"].to_markdown(index=False, floatfmt=".2f"),
        "",
        "## 5.7 Associações ecológicas",
        "",
        associacoes_texto(dados["associacoes"]),
        "",
        tabelas["associacoes"].to_markdown(index=False, floatfmt=".4f"),
        "",
        "# 6 DISCUSSÃO",
        "",
        "O padrão principal é a distância de desempenho entre a oferta presencial de Belém, de conceito superior, e as ofertas conceito 1. A magnitude varia entre ofertas e benchmarks. Modalidade, localização, porte, participação e composição discente devem ser tratados como dimensões concorrentes, e não como explicações previamente determinadas. As diferenças agregadas são compatíveis com hipóteses institucionais, mas exigem investigação documental e qualitativa.",
        "",
        "# 7 CONCLUSÃO",
        "",
        "As ofertas conceito 1 da UFPA apresentam desempenho agregado inferior à oferta interna de conceito superior e, em vários critérios, aos benchmarks comparáveis. Há heterogeneidade entre os próprios cursos conceito 1 em participação, perfil socioeconômico, recomendação e processo formativo. Os resultados sustentam priorização de monitoramento por oferta, mas não permitem atribuir causalidade individual.",
        "",
        "# REFERÊNCIAS",
        "",
        "As referências bibliográficas completas são mantidas no DOCX e no módulo versionado `src/relatorios/referencias.py`.",
        "",
        "# APÊNDICES",
        "",
        "Foram exportadas tabelas derivadas em `relatorios/fisica/apendices/`, preservando CO_CURSO e os indicadores usados no relatório.",
        "",
        "## Aprofundamentos sugeridos",
        "",
    ]
    for titulo, texto in aprofundamentos():
        linhas.extend([f"### {titulo}", "", texto, ""])
    return "\n".join(linhas)


def gerar_relatorio(
    base_projeto: Path,
    saida_docx: Path,
    saida_md: Path,
    metadados: dict[str, str] | None = None,
) -> dict[str, Path | None]:
    metadados = metadados or {}
    dados = _ler(base_projeto)
    tabelas = {
        "ofertas": ofertas_ufpa(dados["cursos"]),
        "participacao": participacao(dados["presenca"]),
        "desempenho": desempenho_ufpa(dados["base"]),
        "territorio": comparacao_territorial(dados["territorio"]),
        "socio": socioeconomico(dados["socio"]),
        "recomendacao": recomendacao_dificuldade(dados["base"]),
        "dimensoes": dimensoes_processo(dados["dimensoes"]),
        "associacoes": associacoes_selecionadas(dados["associacoes"]),
        "benchmark": benchmark_sensibilidade(dados["sensibilidade"]),
        "tucurui": _situacao_tucurui(),
    }
    exportar_apendices(base_projeto, tabelas)
    saida_md.parent.mkdir(parents=True, exist_ok=True)
    saida_md.write_text(_markdown(dados, tabelas), encoding="utf-8")

    doc = Document()
    configurar_documento(doc)
    configurar_cabecalho_rodape(doc)
    _capa(doc, metadados)

    nova_secao(doc)
    _h(doc, "FICHA DE IDENTIFICAÇÃO TÉCNICA", 1)
    _p(doc, "Instituição: Universidade Federal do Pará.")
    _p(doc, "Área: Licenciatura em Física (CO_GRUPO=1402).")
    _p(doc, "Unidade principal de análise: curso (CO_CURSO).")
    _p(doc, f"Data de geração: {datetime.now().strftime('%d/%m/%Y %H:%M')}.")
    _p(doc, "Fontes: microdados do Enade das Licenciaturas 2025 e planilha externa de Conceito Enade.")

    _h(doc, "RESUMO", 1)
    _p(doc, resumo_executivo(dados["base"]), "Resumo")
    _p(doc, "Palavras-chave: Enade; licenciatura em Física; UFPA; microdados; avaliação da educação superior.", "Resumo")
    _h(doc, "ABSTRACT", 1)
    _p(doc, "This report analyzes Physics teacher education programs offered by the Federal University of Pará in the 2025 Enade for Teacher Education Programs. The course, identified by CO_CURSO, is the main unit of analysis. Thematic files were processed separately and aggregated at course level, without reconstructing individual records. Performance, participation, socioeconomic composition, academic trajectory, formative process evaluation, recommendation, and broad and comparable benchmarks were examined. Associations are ecological and do not support individual or causal inference.", "Resumo")
    _p(doc, "Keywords: Enade; Physics teacher education; UFPA; microdata; higher education assessment.", "Resumo")

    _h(doc, "LISTA DE FIGURAS", 1)
    for i, titulo in enumerate([
        "Ofertas de Física da UFPA", "Taxa de presença", "NT_GER por oferta",
        "Comparação territorial de NT_GER", "NT_OBJ por oferta", "NT_DIS por oferta",
        "Conceito e dificuldade", "Processo formativo", "Recomendação do curso",
        "Recomendação da instituição", "Perfil socioeconômico", "Benchmark comparável",
        "Síntese socioeconômica e desempenho",
    ], 1):
        _p(doc, f"Figura {i} – {titulo}", "Fonte")
    _h(doc, "LISTA DE TABELAS", 1)
    for i, titulo in enumerate([
        "Ofertas da UFPA", "Participação", "Desempenho", "Comparação territorial",
        "Perfil socioeconômico", "Processo formativo", "Dificuldade e recomendação",
        "Sensibilidade dos benchmarks", "Associações ecológicas",
    ], 1):
        _p(doc, f"Tabela {i} – {titulo}", "Fonte")
    _h(doc, "LISTA DE ABREVIATURAS E SIGLAS", 1)
    for sigla in [
        "ABNT – Associação Brasileira de Normas Técnicas", "Enade – Exame Nacional de Desempenho dos Estudantes",
        "IES – Instituição de Educação Superior", "Inep – Instituto Nacional de Estudos e Pesquisas Educacionais Anísio Teixeira",
        "NT_DIS – Nota da prova discursiva", "NT_GER – Nota geral", "NT_OBJ – Nota da prova objetiva",
        "UFPA – Universidade Federal do Pará",
    ]:
        _p(doc, sigla, "Fonte")
    _h(doc, "SUMÁRIO", 1)
    for item in [
        "1 INTRODUÇÃO", "2 REFERENCIAL INSTITUCIONAL E METODOLÓGICO", "3 METODOLOGIA",
        "4 PANORAMA DA LICENCIATURA EM FÍSICA", "5 RESULTADOS", "5.1 Desempenho",
        "5.2 Perfil demográfico e socioeconômico", "5.3 Trajetória e condições acadêmicas",
        "5.4 Processo formativo", "5.5 Recomendação", "5.6 Benchmark comparável",
        "5.7 Associações ecológicas", "6 DISCUSSÃO", "7 CONCLUSÃO", "REFERÊNCIAS", "APÊNDICES",
    ]:
        _p(doc, item, "Fonte")
    _p(doc, "Nota: atualize o sumário e os campos ao abrir o documento no Word ou LibreOffice.", "Fonte")

    nova_secao(doc)
    _h(doc, "1 INTRODUÇÃO", 1)
    _p(doc, "O Enade produz evidências sobre desempenho e contexto de formação no ensino superior. Este relatório concentra-se nas ofertas de licenciatura em Física da UFPA, com ênfase naquelas classificadas com Conceito Enade 1 e no contraste com a oferta interna de conceito superior e com cursos da mesma área em diferentes territórios.")
    _p(doc, "A pergunta central é: quais características de desempenho, composição discente, trajetória acadêmica e avaliação do processo formativo diferenciam as ofertas da UFPA com Conceito Enade 1 das demais ofertas da mesma área na UFPA, no Pará, na Região Norte e no Brasil?")
    _p(doc, "O objetivo geral é descrever e contrastar essas ofertas sem atribuir causalidade individual. Os objetivos específicos são auditar participação e desempenho, caracterizar composição discente, examinar processo formativo e recomendação, construir benchmarks comparáveis e sintetizar associações ecológicas no nível do curso.")

    _h(doc, "2 REFERENCIAL INSTITUCIONAL E METODOLÓGICO", 1)
    _h(doc, "2.1 Enade e Conceito Enade", 2)
    _p(doc, "O Conceito Enade é utilizado como classificação externa do curso. As variáveis contínuas dos microdados descrevem desempenho e composição, enquanto a planilha externa preserva o conceito e a situação de divulgação. Ausência de conceito não foi tratada como conceito 1.")
    _h(doc, "2.2 Unidade de análise e anonimização", 2)
    _p(doc, "Os arquivos temáticos não contêm identificador público que permita reconhecer o mesmo estudante entre desempenho, perfil e percepção. A posição da linha não é uma chave. Cada arquivo foi tratado separadamente, agregado por CO_CURSO e somente então relacionado às demais tabelas agregadas.")
    _h(doc, "2.3 Inferência ecológica", 2)
    _p(doc, "Relações entre nota média, renda, trabalho, auxílio, infraestrutura e recomendação são examinadas no nível do curso. Uma associação entre cursos não autoriza concluir que a mesma relação ocorre entre estudantes. Essa restrição é explicitada em todas as interpretações.")

    _h(doc, "3 METODOLOGIA", 1)
    _p(doc, "Foram selecionados cursos com CO_GRUPO=1402. A UFPA foi identificada por CO_IES=569. A unidade principal é CO_CURSO. Os arquivos temáticos foram submetidos a tratamento de ausentes, agregação por curso e validação de unicidade antes da junção um-para-um das tabelas agregadas.")
    _p(doc, "Os grupos exclusivos são: A) UFPA com conceito 1; B) demais ofertas da UFPA com conceito superior; C) outras IES do Pará; D) restante da Região Norte, excluindo Pará; E) restante do Brasil, excluindo Norte. Pará, Norte e Brasil completos foram mantidos como benchmarks descritivos, não como grupos independentes em testes.")
    _p(doc, "Foram calculadas médias, medianas, desvios-padrão, quartis, intervalos interquartis, percentis e intervalos de confiança quando disponíveis. Os benchmarks comparáveis combinaram modalidade, categoria administrativa, organização acadêmica e porte. As correlações de Spearman foram calculadas entre indicadores agregados por curso, com N e p exploratório registrados.")

    _h(doc, "4 PANORAMA DA LICENCIATURA EM FÍSICA", 1)
    _p(doc, panorama_texto(dados["cursos"]))
    adicionar_tabela(doc, "Tabela 1 – Ofertas de Física da UFPA e situação do Conceito Enade", tabelas["ofertas"], "Elaboração própria com base nas fontes do projeto.")
    _h(doc, "4.1 Situação da oferta de Tucuruí", 2)
    _p(doc, tucurui_texto())
    adicionar_tabela(
        doc,
        "Quadro 1 – Disponibilidade de informações da oferta de Física da UFPA em Tucuruí",
        tabelas["tucurui"],
        "Elaboração própria com base na verificação da planilha de Conceito Enade 2025 e dos microdados utilizados no projeto.",
    )

    _h(doc, "5 RESULTADOS", 1)
    _h(doc, "5.1 Desempenho", 2)
    _p(doc, desempenho_texto(dados["base"]))
    adicionar_tabela(doc, "Tabela 2 – Inscritos, participantes e taxa de presença", tabelas["participacao"], "Elaboração própria; N válido e auditoria indicados na tabela.")
    adicionar_tabela(doc, "Tabela 3 – Estatísticas de NT_GER, NT_OBJ e NT_DIS", tabelas["desempenho"], "Elaboração própria; escala de 0 a 100; somente resultados válidos.")

    figdir = base_projeto / "figuras" / "fisica"
    figuras = [
        ("01_ofertas_ufpa.png", "Ofertas de Física da UFPA", "O painel localiza as ofertas validadas e distingue modalidade e conceito.", "A visualização não representa a oferta de Tucuruí, não localizada analiticamente."),
        ("validada_02_taxa_presenca.png", "Taxa de presença por oferta", "As taxas variam entre as ofertas e devem ser lidas com inscritos, participantes e N de nota válida.", "Baixo N aumenta a instabilidade percentual."),
        ("validada_03_nt_ger_ofertas.png", "Nota geral por oferta", "A oferta presencial de Belém ocupa a posição superior no contraste interno.", "A média não substitui a distribuição e não identifica causa."),
        ("04_nt_ger_comparativo.png", "Comparação territorial de NT_GER", "As referências territoriais situam as ofertas da UFPA em relação a conjuntos amplos.", "Pará, Norte e Brasil completos se sobrepõem e são apenas benchmarks descritivos."),
        ("05_nt_obj_por_oferta.png", "Nota objetiva por oferta", "A prova objetiva preserva diferenças entre ofertas, com N válido indicado.", "NT_OBJ possui relação mecânica com NT_GER."),
        ("06_nt_dis_por_oferta.png", "Nota discursiva por oferta", "A prova discursiva apresenta cobertura e dispersão próprias.", "Diferenças de cobertura podem afetar comparações."),
    ]
    for numero, (nome, titulo, leitura, limitacao) in enumerate(figuras, 1):
        _adicionar_figura_com_leitura(doc, figdir / nome, numero, titulo, leitura, limitacao)

    _h(doc, "5.2 Perfil demográfico e socioeconômico", 2)
    _p(doc, socio_texto(dados["socio"]))
    adicionar_tabela(doc, "Tabela 4 – Perfil socioeconômico agregado das ofertas da UFPA", tabelas["socio"], "Elaboração própria; percentuais calculados entre respostas válidas.")
    _adicionar_figura_com_leitura(doc, figdir / "11_perfil_socioeconomico.png", 7, "Perfil socioeconômico por oferta", "O conjunto mostra heterogeneidade de renda, trabalho, ação afirmativa, auxílios e estudo.", "Os indicadores são agregados e sujeitos a ausências específicas por item.")
    _adicionar_figura_com_leitura(doc, figdir / "validada_13_sintese_socioeconomica.png", 8, "Síntese socioeconômica e desempenho", "A figura reúne desempenho e indicadores de composição para leitura contextual por oferta.", "A disposição conjunta não implica relação causal entre perfil individual e nota.")

    _h(doc, "5.3 Trajetória e condições acadêmicas", 2)
    _p(doc, "As ofertas diferem em trabalho, carga horária, bolsas, auxílio permanência, horas de estudo e intenção de exercer o magistério. Esses indicadores ajudam a caracterizar condições acadêmicas e disponibilidade de tempo, mas não devem ser usados isoladamente para explicar desempenho. A análise preserva o percentual válido e as ausências de cada variável.")

    _h(doc, "5.4 Processo formativo", 2)
    _p(doc, "Os itens QE_I20–QE_I66 foram agrupados em dimensões candidatas. As medidas são exploratórias e não constituem uma nota única de qualidade. A decisão metodológica registrada na Sprint 5 foi preservar a leitura dimensional e validar redação, direção da escala e coerência teórica antes de interpretações substantivas.")
    adicionar_tabela(doc, "Tabela 5 – Diagnóstico das dimensões do processo formativo", tabelas["dimensoes"], "Elaboração própria; scores exploratórios no nível do curso.")
    _adicionar_figura_com_leitura(doc, figdir / "08_processo_formativo.png", 9, "Percepção do processo formativo", "O heatmap permite localizar itens e dimensões com maior contraste entre ofertas.", "As respostas são autorreferidas e a consistência dimensional deve ser considerada.")

    _h(doc, "5.5 Recomendação", 2)
    _p(doc, "A percepção de dificuldade foi agregada por curso. QE_I68 e QE_I69 foram mantidos como recomendação do curso e da instituição, respectivamente, sem substituição automática pelo termo satisfação. O percentual 9–10 é apresentado entre respostas válidas.")
    adicionar_tabela(doc, "Tabela 6 – Percepção de dificuldade e recomendação", tabelas["recomendacao"], "Elaboração própria; N válido indicado por item.")
    _adicionar_figura_com_leitura(doc, figdir / "validada_07_conceito_dificuldade.png", 10, "Conceito Enade e percepção de dificuldade", "A figura compara a composição agregada de dificuldade entre ofertas com conceitos distintos.", "Com apenas cinco ofertas da UFPA, não se estima correlação interna robusta.")
    _adicionar_figura_com_leitura(doc, figdir / "09_recomendacao_curso.png", 11, "Recomendação do curso", "A recomendação varia entre as ofertas e complementa os indicadores de desempenho.", "Não constitui avaliação total da qualidade do curso.")
    _adicionar_figura_com_leitura(doc, figdir / "10_recomendacao_instituicao.png", 12, "Recomendação da instituição", "O indicador contextualiza a relação dos concluintes com a instituição.", "A escala deve ser interpretada conforme o rótulo oficial do questionário.")

    _h(doc, "5.6 Benchmark comparável", 2)
    _p(doc, "Os benchmarks amplos oferecem posição territorial; os comparáveis restringem modalidade, categoria, organização acadêmica e porte. A sensibilidade mostra que a magnitude do contraste depende do critério, embora a direção possa permanecer estável. Conjuntos com poucos cursos comparáveis são apresentados como evidência descritiva.")
    adicionar_tabela(doc, "Tabela 7 – Comparação territorial validada", tabelas["territorio"], "Elaboração própria; médias no nível do curso, com versão ponderada por participantes.")
    adicionar_tabela(doc, "Tabela 8 – Sensibilidade dos benchmarks comparáveis", tabelas["benchmark"], "Elaboração própria; cada oferta da UFPA avaliada em quatro critérios.")
    _adicionar_figura_com_leitura(doc, figdir / "12_benchmark_comparavel.png", 13, "Contraste com benchmark comparável", "A figura compara cada oferta com cursos estruturalmente semelhantes.", "O resultado depende da disponibilidade de comparáveis e de características observadas.")

    _h(doc, "5.7 Associações ecológicas", 2)
    _p(doc, associacoes_texto(dados["associacoes"]))
    adicionar_tabela(doc, "Tabela 9 – Associações ecológicas selecionadas", tabelas["associacoes"], "Elaboração própria; Spearman no nível do curso; p-valores exploratórios.")

    _h(doc, "6 DISCUSSÃO", 1)
    _p(doc, "O contraste interno mostra que Belém Presencial, de conceito superior, apresenta desempenho agregado mais elevado que as quatro ofertas conceito 1. Entretanto, as ofertas conceito 1 não constituem um bloco homogêneo: participação, modalidade, porte, composição socioeconômica, recomendação e indicadores formativos variam entre elas.")
    _p(doc, "Os benchmarks comparáveis reduzem parte das diferenças estruturais, mas não eliminam confundimento. A modalidade EaD, a localização no interior ou o baixo N não foram assumidos previamente como explicações. Eles foram tratados como dimensões a verificar empiricamente, e os resultados devem orientar diagnóstico institucional complementar.")
    _p(doc, "A leitura conjunta sugere priorizar análise por oferta. O desempenho deve ser acompanhado de participação e cobertura; o perfil socioeconômico deve ser interpretado como composição; o processo formativo e a recomendação devem manter seus rótulos e limites. A principal hipótese é que múltiplos mecanismos institucionais e contextuais se combinam, hipótese que exige evidências documentais e qualitativas.")

    _h(doc, "7 CONCLUSÃO", 1)
    _p(doc, "As ofertas de Física da UFPA com Conceito Enade 1 diferenciam-se da oferta interna com conceito superior principalmente pelo desempenho agregado e pela posição em benchmarks amplos e comparáveis. Também apresentam heterogeneidade de participação, composição discente, processo formativo e recomendação. Não foi identificado um único fator capaz de explicar o padrão.")
    _p(doc, "Os resultados justificam monitoramento institucional por CO_CURSO, com atenção ao N válido, à modalidade, ao porte e às condições acadêmicas. Relações entre indicadores socioeconômicos e nota permanecem ecológicas. A oferta de Tucuruí deve continuar no cadastro institucional como situação não localizada ou sem conceito, sem ser tratada como conceito 1.")
    _h(doc, "7.1 Aprofundamentos sugeridos", 2)
    for titulo, texto in aprofundamentos():
        _h(doc, titulo, 3)
        _p(doc, texto)

    _h(doc, "REFERÊNCIAS", 1)
    adicionar_referencias(doc)
    _h(doc, "APÊNDICES", 1)
    _p(doc, "Os apêndices em CSV contêm as tabelas derivadas utilizadas no relatório, preservando CO_CURSO quando aplicável. Eles são gerados automaticamente em relatorios/fisica/apendices.")

    saida_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(saida_docx)
    validar_relatorio(saida_docx, saida_md)

    pdf: Path | None = None
    libreoffice = shutil.which("libreoffice") or shutil.which("soffice")
    if libreoffice:
        subprocess.run(
            [libreoffice, "--headless", "--convert-to", "pdf", "--outdir", str(saida_docx.parent), str(saida_docx)],
            check=False,
            capture_output=True,
            text=True,
        )
        candidato = saida_docx.with_suffix(".pdf")
        if candidato.exists() and candidato.stat().st_size > 0:
            pdf = candidato
    return {"docx": saida_docx, "markdown": saida_md, "pdf": pdf}
