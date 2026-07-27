from __future__ import annotations

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


FONTE_PADRAO = "Arial"


def configurar_documento(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Cm(3)
    sec.left_margin = Cm(3)
    sec.right_margin = Cm(2)
    sec.bottom_margin = Cm(2)

    normal = doc.styles["Normal"]
    normal.font.name = FONTE_PADRAO
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.first_line_indent = Cm(1.25)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for nome, tamanho in [("Title", 14), ("Heading 1", 12), ("Heading 2", 12), ("Heading 3", 12)]:
        estilo = doc.styles[nome]
        estilo.font.name = FONTE_PADRAO
        estilo.font.size = Pt(tamanho)
        estilo.font.bold = True
        estilo.paragraph_format.space_before = Pt(12)
        estilo.paragraph_format.space_after = Pt(6)
        estilo.paragraph_format.keep_with_next = True
        estilo.paragraph_format.first_line_indent = Cm(0)
        estilo.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if "Fonte" not in [s.name for s in doc.styles]:
        fonte = doc.styles.add_style("Fonte", WD_STYLE_TYPE.PARAGRAPH)
        fonte.font.name = FONTE_PADRAO
        fonte.font.size = Pt(10)
        fonte.paragraph_format.line_spacing = 1.0
        fonte.paragraph_format.first_line_indent = Cm(0)
        fonte.paragraph_format.space_after = Pt(6)

    if "Resumo" not in [s.name for s in doc.styles]:
        resumo = doc.styles.add_style("Resumo", WD_STYLE_TYPE.PARAGRAPH)
        resumo.font.name = FONTE_PADRAO
        resumo.font.size = Pt(12)
        resumo.paragraph_format.line_spacing = 1.0
        resumo.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        resumo.paragraph_format.first_line_indent = Cm(0)


def adicionar_numero_pagina(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])


def configurar_cabecalho_rodape(doc: Document) -> None:
    for sec in doc.sections:
        adicionar_numero_pagina(sec.header.paragraphs[0])
        rodape = sec.footer.paragraphs[0]
        rodape.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = rodape.add_run("Fonte: elaboração própria com dados do Inep (2025), quando não indicado em contrário.")
        r.font.name = FONTE_PADRAO
        r.font.size = Pt(8)


def nova_secao(doc: Document) -> None:
    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    sec.top_margin = Cm(3)
    sec.left_margin = Cm(3)
    sec.right_margin = Cm(2)
    sec.bottom_margin = Cm(2)
