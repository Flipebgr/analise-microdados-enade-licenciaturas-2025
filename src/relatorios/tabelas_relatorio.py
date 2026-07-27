from __future__ import annotations

import math
from typing import Iterable

import pandas as pd
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def _fmt(valor) -> str:
    if pd.isna(valor):
        return "-"
    if isinstance(valor, float):
        if math.isclose(valor, round(valor), abs_tol=1e-10):
            return str(int(round(valor)))
        return f"{valor:.2f}".replace(".", ",")
    return str(valor)


def adicionar_tabela(doc, titulo: str, df: pd.DataFrame, fonte: str, colunas: Iterable[str] | None = None):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = 0
    run = p.add_run(titulo)
    run.bold = True
    run.font.size = Pt(10)

    dados = df[list(colunas)] if colunas else df
    table = doc.add_table(rows=1, cols=len(dados.columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for i, coluna in enumerate(dados.columns):
        cell = table.rows[0].cells[i]
        cell.text = str(coluna)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for par in cell.paragraphs:
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in par.runs:
                r.bold = True
                r.font.size = Pt(8)

    for _, row in dados.iterrows():
        cells = table.add_row().cells
        for i, valor in enumerate(row):
            cells[i].text = _fmt(valor)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for par in cells[i].paragraphs:
                par.paragraph_format.first_line_indent = 0
                par.paragraph_format.line_spacing = 1.0
                for r in par.runs:
                    r.font.size = Pt(8)

    p_fonte = doc.add_paragraph(style="Fonte")
    p_fonte.add_run(f"Fonte: {fonte}")
    return table
