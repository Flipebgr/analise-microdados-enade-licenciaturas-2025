from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
import math

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

REGIOES = {1: "Norte", 2: "Nordeste", 3: "Sudeste", 4: "Sul", 5: "Centro-Oeste"}
CO_IES_UFPA = 569
SEED = 2025


@dataclass(frozen=True)
class CaminhosRelatorio:
    base: Path
    saida: Path
    figuras: Path
    tabelas: Path


def _num(serie: pd.Series) -> pd.Series:
    return pd.to_numeric(serie, errors="coerce")


def carregar_base(caminho: Path) -> pd.DataFrame:
    if not caminho.exists():
        raise FileNotFoundError(f"Base da Sprint 4 não localizada: {caminho}")
    df = pd.read_csv(caminho)
    obrigatorias = {
        "CO_CURSO", "CO_IES", "CO_REGIAO_CURSO", "CO_MODALIDADE", "CO_CATEGAD",
        "CO_ORGACAD", "ROTULO_OFERTA", "nt_ger_count", "nt_ger_mean", "nt_ger_std",
        "nt_ger_median", "CONCEITO_ENADE_NUM",
    }
    ausentes = sorted(obrigatorias.difference(df.columns))
    if ausentes:
        raise ValueError(f"Colunas obrigatórias ausentes: {ausentes}")
    for col in [
        "CO_CURSO", "CO_IES", "CO_REGIAO_CURSO", "CO_MODALIDADE", "CO_CATEGAD",
        "CO_ORGACAD", "nt_ger_count", "nt_ger_mean", "nt_ger_std", "nt_ger_median",
        "CONCEITO_ENADE_NUM",
    ]:
        df[col] = _num(df[col])
    df["REGIAO"] = df["CO_REGIAO_CURSO"].map(REGIOES)
    df = df[df["nt_ger_count"].fillna(0).gt(0) & df["nt_ger_mean"].notna()].copy()
    if df["CO_CURSO"].duplicated().any():
        raise ValueError("A base analítica deve conter uma linha por CO_CURSO.")
    return df


def _media_ponderada(df: pd.DataFrame) -> float:
    return float(np.average(df["nt_ger_mean"], weights=df["nt_ger_count"]))


def _variancia_agrupada(df: pd.DataFrame) -> float:
    n = df["nt_ger_count"].astype(float)
    medias = df["nt_ger_mean"].astype(float)
    desvios = df["nt_ger_std"].fillna(0).astype(float)
    total = float(n.sum())
    if total <= 1:
        return float("nan")
    media = float(np.average(medias, weights=n))
    soma_intra = float((((n - 1).clip(lower=0)) * desvios.pow(2)).sum())
    soma_entre = float((n * (medias - media).pow(2)).sum())
    return (soma_intra + soma_entre) / (total - 1)


def _bootstrap_media_ponderada(df: pd.DataFrame, repeticoes: int = 2000) -> tuple[float, float]:
    if len(df) < 2:
        media = _media_ponderada(df)
        return media, media
    rng = np.random.default_rng(SEED)
    valores = np.empty(repeticoes)
    indices = np.arange(len(df))
    for i in range(repeticoes):
        amostra = df.iloc[rng.choice(indices, size=len(indices), replace=True)]
        valores[i] = _media_ponderada(amostra)
    return float(np.quantile(valores, 0.025)), float(np.quantile(valores, 0.975))


def resumir_grupo(df: pd.DataFrame, grupo: str) -> dict[str, float | int | str]:
    if df.empty:
        return {
            "grupo": grupo, "n_cursos": 0, "n_participantes": 0,
            "media_ponderada": np.nan, "ic95_inf": np.nan, "ic95_sup": np.nan,
            "media_cursos": np.nan, "mediana_cursos": np.nan, "dp_cursos": np.nan,
            "p25_cursos": np.nan, "p75_cursos": np.nan, "min_cursos": np.nan,
            "max_cursos": np.nan, "dp_participantes_agrupado": np.nan,
        }
    inf, sup = _bootstrap_media_ponderada(df)
    var = _variancia_agrupada(df)
    return {
        "grupo": grupo,
        "n_cursos": int(df["CO_CURSO"].nunique()),
        "n_participantes": int(df["nt_ger_count"].sum()),
        "media_ponderada": _media_ponderada(df),
        "ic95_inf": inf,
        "ic95_sup": sup,
        "media_cursos": float(df["nt_ger_mean"].mean()),
        "mediana_cursos": float(df["nt_ger_mean"].median()),
        "dp_cursos": float(df["nt_ger_mean"].std(ddof=1)) if len(df) > 1 else 0.0,
        "p25_cursos": float(df["nt_ger_mean"].quantile(0.25)),
        "p75_cursos": float(df["nt_ger_mean"].quantile(0.75)),
        "min_cursos": float(df["nt_ger_mean"].min()),
        "max_cursos": float(df["nt_ger_mean"].max()),
        "dp_participantes_agrupado": math.sqrt(var) if pd.notna(var) and var >= 0 else np.nan,
    }


def construir_resumos(df: pd.DataFrame) -> pd.DataFrame:
    grupos: list[tuple[str, pd.DataFrame]] = [
        ("UFPA", df[df["CO_IES"].eq(CO_IES_UFPA)]),
        ("Norte sem UFPA", df[df["REGIAO"].eq("Norte") & ~df["CO_IES"].eq(CO_IES_UFPA)]),
    ]
    for regiao in ["Norte", "Nordeste", "Sudeste", "Sul", "Centro-Oeste"]:
        grupos.append((regiao, df[df["REGIAO"].eq(regiao)]))
    grupos.extend([
        ("Brasil geral", df),
        ("Brasil sem UFPA", df[~df["CO_IES"].eq(CO_IES_UFPA)]),
        ("Brasil sem Norte", df[~df["REGIAO"].eq("Norte")]),
    ])
    return pd.DataFrame([resumir_grupo(parte, nome) for nome, parte in grupos])


def _cohen_d_resumos(a: pd.Series, b: pd.Series) -> float:
    n1, n2 = float(a["n_participantes"]), float(b["n_participantes"])
    s1, s2 = float(a["dp_participantes_agrupado"]), float(b["dp_participantes_agrupado"])
    if n1 <= 1 or n2 <= 1 or not np.isfinite(s1) or not np.isfinite(s2):
        return float("nan")
    pooled = math.sqrt(((n1 - 1) * s1**2 + (n2 - 1) * s2**2) / (n1 + n2 - 2))
    return (float(a["media_ponderada"]) - float(b["media_ponderada"])) / pooled if pooled else np.nan


def construir_contrastes(resumos: pd.DataFrame) -> pd.DataFrame:
    idx = resumos.set_index("grupo")
    comparacoes = [
        ("UFPA", "Norte sem UFPA"),
        ("UFPA", "Nordeste"),
        ("UFPA", "Sudeste"),
        ("UFPA", "Sul"),
        ("UFPA", "Centro-Oeste"),
        ("UFPA", "Brasil sem UFPA"),
        ("Norte", "Brasil sem Norte"),
    ]
    linhas = []
    for referencia, comparador in comparacoes:
        a, b = idx.loc[referencia], idx.loc[comparador]
        linhas.append({
            "referencia": referencia,
            "comparador": comparador,
            "media_referencia": a["media_ponderada"],
            "media_comparador": b["media_ponderada"],
            "diferenca": a["media_ponderada"] - b["media_ponderada"],
            "cohen_d": _cohen_d_resumos(a, b),
            "n_cursos_referencia": int(a["n_cursos"]),
            "n_cursos_comparador": int(b["n_cursos"]),
            "n_participantes_referencia": int(a["n_participantes"]),
            "n_participantes_comparador": int(b["n_participantes"]),
        })
    return pd.DataFrame(linhas)


def construir_ofertas_ufpa(df: pd.DataFrame, resumos: pd.DataFrame) -> pd.DataFrame:
    ufpa = df[df["CO_IES"].eq(CO_IES_UFPA)].copy()
    norte = df[df["REGIAO"].eq("Norte") & ~df["CO_IES"].eq(CO_IES_UFPA)]
    brasil = df[~df["CO_IES"].eq(CO_IES_UFPA)]
    media_norte = _media_ponderada(norte)
    media_brasil = _media_ponderada(brasil)
    medias_norte = norte["nt_ger_mean"].dropna()
    medias_brasil = brasil["nt_ger_mean"].dropna()
    ufpa["dif_norte_sem_ufpa"] = ufpa["nt_ger_mean"] - media_norte
    ufpa["dif_brasil_sem_ufpa"] = ufpa["nt_ger_mean"] - media_brasil
    ufpa["percentil_norte_sem_ufpa"] = ufpa["nt_ger_mean"].apply(
        lambda x: float((medias_norte.le(x).mean()) * 100)
    )
    ufpa["percentil_brasil_sem_ufpa"] = ufpa["nt_ger_mean"].apply(
        lambda x: float((medias_brasil.le(x).mean()) * 100)
    )
    cols = [
        "CO_CURSO", "ROTULO_OFERTA", "CONCEITO_ENADE_NUM", "nt_ger_count",
        "nt_ger_mean", "nt_ger_median", "nt_ger_std", "nt_ger_ic95_inf",
        "nt_ger_ic95_sup", "dif_norte_sem_ufpa", "dif_brasil_sem_ufpa",
        "percentil_norte_sem_ufpa", "percentil_brasil_sem_ufpa",
    ]
    return ufpa[[c for c in cols if c in ufpa.columns]].sort_values("nt_ger_mean", ascending=False)


def _resumo_recorte(df: pd.DataFrame, recorte: str, filtro: pd.Series) -> list[dict]:
    sub = df[filtro].copy()
    linhas = []
    for grupo, parte in [
        ("UFPA", sub[sub["CO_IES"].eq(CO_IES_UFPA)]),
        ("Norte sem UFPA", sub[sub["REGIAO"].eq("Norte") & ~sub["CO_IES"].eq(CO_IES_UFPA)]),
        ("Brasil sem UFPA", sub[~sub["CO_IES"].eq(CO_IES_UFPA)]),
    ]:
        r = resumir_grupo(parte, grupo)
        r["recorte"] = recorte
        linhas.append(r)
    return linhas


def construir_sensibilidade(df: pd.DataFrame) -> pd.DataFrame:
    linhas: list[dict] = []
    linhas += _resumo_recorte(df, "Todos", pd.Series(True, index=df.index))
    linhas += _resumo_recorte(df, "Presencial", df["CO_MODALIDADE"].eq(1))
    linhas += _resumo_recorte(df, "EaD", df["CO_MODALIDADE"].eq(0))
    linhas += _resumo_recorte(df, "N válido >= 10", df["nt_ger_count"].ge(10))
    linhas += _resumo_recorte(df, "N válido >= 20", df["nt_ger_count"].ge(20))
    linhas += _resumo_recorte(df, "Universidades federais", df["CO_CATEGAD"].eq(1))
    return pd.DataFrame(linhas)[[
        "recorte", "grupo", "n_cursos", "n_participantes", "media_ponderada",
        "media_cursos", "mediana_cursos", "ic95_inf", "ic95_sup",
    ]]


def _salvar_figura(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    plt.tight_layout()
    plt.savefig(path, dpi=220, bbox_inches="tight")
    plt.close()


def gerar_figuras(df: pd.DataFrame, resumos: pd.DataFrame, ofertas: pd.DataFrame, sens: pd.DataFrame, pasta: Path) -> list[Path]:
    pasta.mkdir(parents=True, exist_ok=True)
    geradas: list[Path] = []

    # 1. Ofertas UFPA com referências exclusivas.
    fig, ax = plt.subplots(figsize=(11, 6.2))
    plot = ofertas.sort_values("nt_ger_mean")
    ax.barh(plot["ROTULO_OFERTA"], plot["nt_ger_mean"])
    for i, v in enumerate(plot["nt_ger_mean"]):
        ax.text(v + 0.35, i, f"{v:.1f} (N={int(plot.iloc[i]['nt_ger_count'])})", va="center", fontsize=9)
    refs = resumos.set_index("grupo")["media_ponderada"]
    for nome, estilo in [("UFPA", "-"), ("Norte sem UFPA", "--"), ("Brasil sem UFPA", ":")]:
        ax.axvline(refs[nome], linestyle=estilo, linewidth=1.6, label=f"{nome}: {refs[nome]:.1f}")
    ax.set_title("Nota geral das ofertas de Física da UFPA")
    ax.set_xlabel("NT_GER média (0 a 100)")
    ax.set_ylabel("")
    ax.legend(loc="lower right")
    p = pasta / "01_nt_ger_ufpa_referencias.png"
    _salvar_figura(p)
    geradas.append(p)

    # 2. UFPA e regiões.
    ordem = ["UFPA", "Norte sem UFPA", "Nordeste", "Centro-Oeste", "Sudeste", "Sul", "Brasil geral"]
    r = resumos.set_index("grupo").loc[ordem].reset_index()
    fig, ax = plt.subplots(figsize=(11, 6.2))
    y = np.arange(len(r))
    xerr = np.vstack([r["media_ponderada"] - r["ic95_inf"], r["ic95_sup"] - r["media_ponderada"]])
    ax.errorbar(r["media_ponderada"], y, xerr=xerr, fmt="o", capsize=4)
    ax.set_yticks(y, r["grupo"])
    ax.invert_yaxis()
    for i, row in r.iterrows():
        ax.text(row["media_ponderada"] + 0.35, i, f"{row['media_ponderada']:.1f} | {int(row['n_cursos'])} cursos", va="center", fontsize=9)
    ax.set_title("UFPA e regiões brasileiras: média ponderada de NT_GER")
    ax.set_xlabel("Média ponderada pelos participantes; IC bootstrap por curso")
    p = pasta / "02_ufpa_regioes_brasil.png"
    _salvar_figura(p)
    geradas.append(p)

    # 3. Distribuição de médias de curso.
    fig, ax = plt.subplots(figsize=(11, 6.5))
    regioes = ["Norte", "Nordeste", "Centro-Oeste", "Sudeste", "Sul"]
    dados = [df.loc[df["REGIAO"].eq(reg), "nt_ger_mean"].dropna().values for reg in regioes]
    ax.boxplot(dados, tick_labels=regioes, showfliers=False)
    for i, reg in enumerate(regioes, 1):
        vals = df.loc[df["REGIAO"].eq(reg), "nt_ger_mean"].dropna()
        jitter = np.random.default_rng(SEED + i).normal(i, 0.045, len(vals))
        ax.scatter(jitter, vals, s=13, alpha=0.45)
    ufpa_df = df[df["CO_IES"].eq(CO_IES_UFPA)]
    ax.scatter(np.ones(len(ufpa_df)), ufpa_df["nt_ger_mean"], marker="D", s=55, label="Ofertas da UFPA")
    ax.set_title("Distribuição das médias de NT_GER dos cursos de Física")
    ax.set_ylabel("Média do curso")
    ax.legend()
    p = pasta / "03_distribuicao_cursos_regiao.png"
    _salvar_figura(p)
    geradas.append(p)

    # 4. Diferença regional para UFPA.
    ufpa_media = float(resumos.set_index("grupo").loc["UFPA", "media_ponderada"])
    reg = resumos[resumos["grupo"].isin(["Norte sem UFPA", "Nordeste", "Centro-Oeste", "Sudeste", "Sul", "Brasil sem UFPA"])].copy()
    reg["diferenca_para_ufpa"] = reg["media_ponderada"] - ufpa_media
    reg = reg.sort_values("diferenca_para_ufpa")
    fig, ax = plt.subplots(figsize=(10.5, 5.8))
    ax.barh(reg["grupo"], reg["diferenca_para_ufpa"])
    ax.axvline(0, linewidth=1)
    for i, v in enumerate(reg["diferenca_para_ufpa"]):
        ax.text(v + (0.15 if v >= 0 else -0.15), i, f"{v:+.1f}", va="center", ha="left" if v >= 0 else "right")
    ax.set_title("Diferença de NT_GER em relação à média agregada da UFPA")
    ax.set_xlabel("Média do grupo - média da UFPA")
    p = pasta / "04_diferencas_para_ufpa.png"
    _salvar_figura(p)
    geradas.append(p)

    # 5. Percentis das ofertas.
    fig, ax = plt.subplots(figsize=(10.5, 6))
    o = ofertas.sort_values("percentil_brasil_sem_ufpa")
    y = np.arange(len(o))
    ax.scatter(o["percentil_norte_sem_ufpa"], y, label="Percentil no Norte sem UFPA", marker="o", s=60)
    ax.scatter(o["percentil_brasil_sem_ufpa"], y, label="Percentil no Brasil sem UFPA", marker="s", s=60)
    ax.set_yticks(y, o["ROTULO_OFERTA"])
    ax.set_xlim(0, 100)
    ax.set_xlabel("Percentil entre médias de curso")
    ax.set_title("Posição relativa das ofertas de Física da UFPA")
    ax.legend()
    p = pasta / "05_percentis_ofertas_ufpa.png"
    _salvar_figura(p)
    geradas.append(p)

    # 6. Sensibilidade.
    pivot = sens.pivot(index="recorte", columns="grupo", values="media_ponderada")
    recortes = [r for r in ["Todos", "Presencial", "EaD", "N válido >= 10", "N válido >= 20", "Universidades federais"] if r in pivot.index]
    fig, ax = plt.subplots(figsize=(11, 6.5))
    x = np.arange(len(recortes))
    largura = 0.25
    for i, grupo in enumerate(["UFPA", "Norte sem UFPA", "Brasil sem UFPA"]):
        vals = pivot.reindex(recortes)[grupo]
        ax.bar(x + (i - 1) * largura, vals, largura, label=grupo)
    ax.set_xticks(x, recortes, rotation=20, ha="right")
    ax.set_ylabel("Média ponderada de NT_GER")
    ax.set_title("Análise de sensibilidade dos contrastes")
    ax.legend()
    p = pasta / "06_sensibilidade_recortes.png"
    _salvar_figura(p)
    geradas.append(p)
    return geradas


def _fmt(v: float | int | str, casas: int = 1) -> str:
    if isinstance(v, str):
        return v
    if pd.isna(v):
        return "Não disponível"
    if isinstance(v, (int, np.integer)):
        return f"{int(v)}"
    return f"{float(v):.{casas}f}".replace(".", ",")


def _tabela_docx(doc: Document, df: pd.DataFrame, titulo: str, casas: int = 1) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = 0
    r = p.add_run(titulo)
    r.bold = True
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    for i, col in enumerate(df.columns):
        table.rows[0].cells[i].text = str(col)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = _fmt(val, casas)
    fonte = doc.add_paragraph("Fonte: elaboração própria com base nos microdados do Enade das Licenciaturas 2025.")
    fonte.paragraph_format.first_line_indent = 0
    fonte.runs[0].font.size = Pt(9)


def _adicionar_figura(doc: Document, path: Path, titulo: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = 0
    p.add_run(titulo).bold = True
    doc.add_picture(str(path), width=Inches(6.4))
    cap = doc.paragraphs[-1]
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fonte = doc.add_paragraph("Fonte: elaboração própria com base nos microdados do Enade das Licenciaturas 2025.")
    fonte.paragraph_format.first_line_indent = 0
    fonte.runs[0].font.size = Pt(9)


def _configurar_doc(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Inches(1.18)
    sec.left_margin = Inches(1.18)
    sec.right_margin = Inches(0.79)
    sec.bottom_margin = Inches(0.79)
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].paragraph_format.line_spacing = 1.5
    styles["Normal"].paragraph_format.first_line_indent = Inches(0.49)
    for name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        styles[name].font.name = "Arial"


def _texto_resultados(resumos: pd.DataFrame, contrastes: pd.DataFrame, ofertas: pd.DataFrame) -> dict[str, str]:
    r = resumos.set_index("grupo")
    c = contrastes.set_index(["referencia", "comparador"])
    melhor = ofertas.iloc[0]
    pior = ofertas.iloc[-1]
    return {
        "resumo": (
            f"O relatório compara a nota geral (NT_GER) de {int(r.loc['UFPA', 'n_participantes'])} participantes válidos "
            f"distribuídos em {int(r.loc['UFPA', 'n_cursos'])} ofertas de Física da UFPA com os cursos da Região Norte, "
            "das demais regiões e do Brasil. A análise combina médias ponderadas pelos participantes e estatísticas entre cursos. "
            "A unidade principal permanece CO_CURSO; não há reconstrução de vínculos individuais entre arquivos temáticos."
        ),
        "ufpa": (
            f"A média ponderada da UFPA foi {_fmt(r.loc['UFPA', 'media_ponderada'])}. Entre as ofertas, "
            f"{melhor['ROTULO_OFERTA']} apresentou a maior média ({_fmt(melhor['nt_ger_mean'])}; N={int(melhor['nt_ger_count'])}), "
            f"enquanto {pior['ROTULO_OFERTA']} apresentou a menor ({_fmt(pior['nt_ger_mean'])}; N={int(pior['nt_ger_count'])}). "
            "A amplitude interna mostra que uma média institucional isolada não representa adequadamente a heterogeneidade das ofertas."
        ),
        "norte": (
            f"A comparação exclusiva entre UFPA e Norte sem UFPA resultou em diferença de "
            f"{_fmt(c.loc[('UFPA', 'Norte sem UFPA'), 'diferenca'])} ponto(s) na NT_GER e tamanho de efeito "
            f"d={_fmt(c.loc[('UFPA', 'Norte sem UFPA'), 'cohen_d'], 2)}. A Região Norte completa inclui a própria UFPA e, "
            "por isso, é apresentada apenas como benchmark descritivo sobreposto."
        ),
        "brasil": (
            f"Frente ao Brasil sem UFPA, a diferença da média ponderada da UFPA foi "
            f"{_fmt(c.loc[('UFPA', 'Brasil sem UFPA'), 'diferenca'])} ponto(s), com d="
            f"{_fmt(c.loc[('UFPA', 'Brasil sem UFPA'), 'cohen_d'], 2)}. O Brasil geral é mantido como referência descritiva; "
            "os contrastes inferenciais usam grupos exclusivos."
        ),
    }


def gerar_markdown(resumos: pd.DataFrame, contrastes: pd.DataFrame, ofertas: pd.DataFrame, sens: pd.DataFrame, figs: list[Path]) -> str:
    t = _texto_resultados(resumos, contrastes, ofertas)
    resumo_cols = ["grupo", "n_cursos", "n_participantes", "media_ponderada", "ic95_inf", "ic95_sup", "media_cursos", "mediana_cursos", "dp_cursos"]
    oferta_cols = ["CO_CURSO", "ROTULO_OFERTA", "CONCEITO_ENADE_NUM", "nt_ger_count", "nt_ger_mean", "nt_ger_median", "dif_norte_sem_ufpa", "dif_brasil_sem_ufpa", "percentil_norte_sem_ufpa", "percentil_brasil_sem_ufpa"]
    linhas = [
        "# Desempenho dos cursos de Física da UFPA no contexto regional e nacional",
        "",
        "## Resumo",
        "",
        t["resumo"],
        "",
        "**Palavras-chave:** Enade; Física; UFPA; Região Norte; comparação regional; NT_GER.",
        "",
        "# 1 Introdução",
        "",
        "O objetivo é posicionar o desempenho agregado das ofertas de Física da UFPA diante da Região Norte, das demais regiões brasileiras e do Brasil geral. A pergunta central é: como a NT_GER das ofertas da UFPA se posiciona nesses referenciais, considerando a heterogeneidade de modalidade, organização acadêmica, categoria administrativa e porte?",
        "",
        "# 2 Metodologia",
        "",
        "A base foi construída previamente por curso. O desempenho foi agregado no arquivo temático de notas por CO_CURSO e relacionado ao cadastro regional somente após a obtenção de uma linha por curso. A média ponderada usa nt_ger_count como peso. A média e a mediana entre cursos atribuem o mesmo peso a cada oferta. Os intervalos de 95% foram estimados por reamostragem de cursos com semente fixa 2025. Esses intervalos descrevem a incerteza entre cursos e não corrigem integralmente a dependência institucional.",
        "",
        "# 3 Cobertura e resultados gerais",
        "",
        resumos[resumo_cols].to_markdown(index=False, floatfmt=".2f"),
        "",
        "# 4 Ofertas da UFPA",
        "",
        t["ufpa"],
        "",
        ofertas[oferta_cols].to_markdown(index=False, floatfmt=".2f"),
        "",
        "# 5 UFPA e Região Norte",
        "",
        t["norte"],
        "",
        "# 6 UFPA, demais regiões e Brasil",
        "",
        t["brasil"],
        "",
        contrastes.to_markdown(index=False, floatfmt=".3f"),
        "",
        "# 7 Análise de sensibilidade",
        "",
        sens.to_markdown(index=False, floatfmt=".2f"),
        "",
        "# 8 Discussão",
        "",
        "A posição da UFPA deve ser interpretada em conjunto com a forte heterogeneidade entre suas ofertas. Diferenças regionais podem refletir simultaneamente composição institucional, modalidade, porte, perfil discente e condições territoriais. O desenho é descritivo e ecológico; não identifica efeito causal da região ou da instituição sobre a nota individual.",
        "",
        "# 9 Conclusão",
        "",
        "O relatório fornece três perspectivas complementares: desempenho de cada oferta da UFPA, média institucional da UFPA e referências regionais e nacionais exclusivas. A comparação com Norte sem UFPA e Brasil sem UFPA evita que a própria instituição componha o benchmark usado para avaliá-la. As conclusões devem priorizar a heterogeneidade entre ofertas e a estabilidade dos resultados nos recortes de sensibilidade.",
        "",
        "# Figuras",
        "",
    ]
    for i, fig in enumerate(figs, 1):
        linhas.extend([f"Figura {i} - {fig.stem.replace('_', ' ').title()}", "", f"`{fig.as_posix()}`", ""])
    linhas.extend([
        "# Limitações",
        "",
        "- O relatório usa estatísticas agregadas por curso; não há vinculação individual entre região e nota entre arquivos.",
        "- A média ponderada é reproduzida a partir de médias e N por curso; a mediana regional individual não é estimada.",
        "- Os intervalos bootstrap reamostram cursos e são descritivos.",
        "- Comparações amplas não substituem benchmark comparável por modalidade, categoria, organização e porte.",
        "- Tucuruí, CO_CURSO 1627581, não entra nos cálculos por não ter sido localizado nas fontes analíticas usadas.",
    ])
    return "\n".join(linhas)


def gerar_docx(resumos: pd.DataFrame, contrastes: pd.DataFrame, ofertas: pd.DataFrame, sens: pd.DataFrame, figs: list[Path], saida: Path) -> None:
    t = _texto_resultados(resumos, contrastes, ofertas)
    doc = Document()
    _configurar_doc(doc)
    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = 0

    run = p.add_run("UNIVERSIDADE FEDERAL DO PARÁ")
    run.bold = True
    run.font.size = Pt(12)

    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = 0

    run = p.add_run(
        "DESEMPENHO DOS CURSOS DE FÍSICA DA UFPA NO CONTEXTO "
        "REGIONAL E NACIONAL\n"
        "COMPARAÇÃO DA NOTA GERAL NO ENADE DAS LICENCIATURAS 2025"
    )
    run.bold = True
    run.font.size = Pt(14)

    for _ in range(8):
        doc.add_paragraph()

    p = doc.add_paragraph("Belém\n2026")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = 0

    doc.add_page_break()

    doc.add_heading("RESUMO", level=1)
    doc.add_paragraph(t["resumo"])
    doc.add_paragraph("Palavras-chave: Enade; Física; UFPA; Região Norte; comparação regional; NT_GER.")
    doc.add_heading("1 INTRODUÇÃO", level=1)
    doc.add_paragraph("Este relatório complementar posiciona as ofertas de Física da UFPA diante da Região Norte, das demais regiões brasileiras e do Brasil geral. O foco é a nota geral (NT_GER), apresentada por oferta e de forma agregada para a instituição.")
    doc.add_heading("2 METODOLOGIA", level=1)
    doc.add_paragraph("A unidade principal é CO_CURSO. O arquivo de desempenho foi tratado e agregado por curso antes da junção um-para-um com o cadastro regional. Não foram vinculados registros individuais entre arquivos. A média ponderada utiliza o número de participantes com NT_GER válida. A média e a mediana entre cursos atribuem peso igual a cada oferta. Os intervalos de 95% foram obtidos por bootstrap de cursos, com semente 2025.")
    doc.add_paragraph("Para evitar sobreposição, os contrastes principais usam Norte sem UFPA, Brasil sem UFPA e Brasil sem Norte. Norte completo e Brasil geral são benchmarks descritivos.")

    doc.add_heading("3 COBERTURA E RESULTADOS GERAIS", level=1)
    tab = resumos[["grupo", "n_cursos", "n_participantes", "media_ponderada", "media_cursos", "mediana_cursos"]].copy()
    tab.columns = ["Grupo", "Cursos", "Participantes", "Média ponderada", "Média dos cursos", "Mediana dos cursos"]
    _tabela_docx(doc, tab, "Tabela 1 - Síntese de NT_GER por grupo")
    _adicionar_figura(doc, figs[1], "Figura 1 - UFPA e regiões brasileiras")

    doc.add_heading("4 OFERTAS DA UFPA", level=1)
    doc.add_paragraph(t["ufpa"])
    o = ofertas[["CO_CURSO", "ROTULO_OFERTA", "CONCEITO_ENADE_NUM", "nt_ger_count", "nt_ger_mean", "dif_norte_sem_ufpa", "dif_brasil_sem_ufpa"]].copy()
    o.columns = ["CO_CURSO", "Oferta", "Conceito", "N válido", "Média", "Dif. Norte", "Dif. Brasil"]
    _tabela_docx(doc, o, "Tabela 2 - Desempenho das ofertas de Física da UFPA")
    _adicionar_figura(doc, figs[0], "Figura 2 - Ofertas da UFPA e referências exclusivas")
    _adicionar_figura(doc, figs[4], "Figura 3 - Posição percentílica das ofertas da UFPA")

    doc.add_heading("5 UFPA E REGIÃO NORTE", level=1)
    doc.add_paragraph(t["norte"])
    _adicionar_figura(doc, figs[2], "Figura 4 - Distribuição das médias dos cursos por região")

    doc.add_heading("6 UFPA, DEMAIS REGIÕES E BRASIL", level=1)
    doc.add_paragraph(t["brasil"])
    c = contrastes[["referencia", "comparador", "diferenca", "cohen_d", "n_cursos_referencia", "n_cursos_comparador"]].copy()
    c.columns = ["Referência", "Comparador", "Diferença", "Cohen d", "Cursos ref.", "Cursos comp."]
    _tabela_docx(doc, c, "Tabela 3 - Contrastes exclusivos de NT_GER", casas=2)
    _adicionar_figura(doc, figs[3], "Figura 5 - Diferenças em relação à UFPA")

    doc.add_heading("7 ANÁLISE DE SENSIBILIDADE", level=1)
    s = sens[["recorte", "grupo", "n_cursos", "n_participantes", "media_ponderada"]].copy()
    s.columns = ["Recorte", "Grupo", "Cursos", "Participantes", "Média ponderada"]
    _tabela_docx(doc, s, "Tabela 4 - Sensibilidade por modalidade, porte e categoria")
    _adicionar_figura(doc, figs[5], "Figura 6 - Sensibilidade dos contrastes")

    doc.add_heading("8 DISCUSSÃO", level=1)
    doc.add_paragraph("A comparação mostra que a média institucional deve ser lida juntamente com a dispersão entre as cinco ofertas da UFPA. A localização regional não é uma exposição isolada: ela se combina com modalidade, categoria administrativa, organização acadêmica, porte e composição discente. Assim, diferenças observadas são padrões agregados compatíveis com hipóteses institucionais, não efeitos causais.")
    doc.add_heading("9 CONCLUSÃO", level=1)
    doc.add_paragraph("O posicionamento da UFPA é apresentado em três escalas: oferta, instituição e território. A exclusão da própria UFPA dos benchmarks do Norte e do Brasil melhora a independência dos contrastes. A leitura final deve destacar tanto a posição regional e nacional quanto a heterogeneidade interna entre Belém Presencial, Belém EaD, Abaetetuba, Ananindeua e Salinópolis.")
    doc.add_heading("LIMITAÇÕES", level=1)
    for item in [
        "Uso de estatísticas agregadas por curso, sem vinculação individual entre arquivos.",
        "Impossibilidade de estimar a mediana individual regional a partir da base agregada.",
        "Intervalos bootstrap descritivos por curso.",
        "Necessidade de benchmark comparável para interpretações institucionais mais estritas.",
        "Tucuruí (CO_CURSO 1627581) não integra os cálculos, pois não foi localizado nas fontes analíticas.",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    saida.parent.mkdir(parents=True, exist_ok=True)
    doc.save(saida)


def executar(base_projeto: Path) -> dict[str, Path]:
    caminhos = CaminhosRelatorio(
        base=base_projeto / "dados_processados" / "fisica" / "base_analitica_cursos.csv",
        saida=base_projeto / "relatorios" / "fisica",
        figuras=base_projeto / "figuras" / "fisica" / "regional",
        tabelas=base_projeto / "relatorios" / "fisica" / "apendices" / "regional",
    )
    df = carregar_base(caminhos.base)
    resumos = construir_resumos(df)
    contrastes = construir_contrastes(resumos)
    ofertas = construir_ofertas_ufpa(df, resumos)
    sens = construir_sensibilidade(df)
    caminhos.tabelas.mkdir(parents=True, exist_ok=True)
    resumos.to_csv(caminhos.tabelas / "resumo_regional.csv", index=False, encoding="utf-8-sig")
    contrastes.to_csv(caminhos.tabelas / "contrastes_regionais.csv", index=False, encoding="utf-8-sig")
    ofertas.to_csv(caminhos.tabelas / "ofertas_ufpa_posicao.csv", index=False, encoding="utf-8-sig")
    sens.to_csv(caminhos.tabelas / "sensibilidade_regional.csv", index=False, encoding="utf-8-sig")
    figs = gerar_figuras(df, resumos, ofertas, sens, caminhos.figuras)
    md = caminhos.saida / "relatorio_comparacao_regional_fisica_2025.md"
    docx = caminhos.saida / "relatorio_comparacao_regional_fisica_2025.docx"
    md.write_text(gerar_markdown(resumos, contrastes, ofertas, sens, figs), encoding="utf-8")
    gerar_docx(resumos, contrastes, ofertas, sens, figs, docx)
    return {"markdown": md, "docx": docx, "figuras": caminhos.figuras, "tabelas": caminhos.tabelas}
