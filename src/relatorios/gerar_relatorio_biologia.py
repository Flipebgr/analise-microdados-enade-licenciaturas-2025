from __future__ import annotations

from pathlib import Path

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from src.biologia import BIOLOGIA, CO_CURSO_SOURE
from src.biologia.rotulos_questionario import ROTULOS_QE, rotulo_item
from src.relatorios.conversao_pdf import converter_docx_para_pdf
from src.relatorios.figuras_relatorio import adicionar_figura
from src.relatorios.formatacao_abnt import configurar_cabecalho_rodape, configurar_documento
from src.relatorios.referencias import adicionar_referencias
from src.relatorios.resultado_relatorio import ResultadoRelatorio
from src.relatorios.tabelas_relatorio import adicionar_tabela

FONTE_DADOS = (
    "Elaboração própria com base nos microdados do Enade das Licenciaturas 2025 "
    "e na planilha de Conceito Enade."
)
FONTE_QUESTIONARIO = (
    "Inep, Dicionário de Variáveis e Questionário do Estudante – "
    "Enade das Licenciaturas 2025."
)


def _ler_csv(caminho: Path, *, obrigatorio: bool = True) -> pd.DataFrame:
    if not caminho.exists():
        if obrigatorio:
            raise FileNotFoundError(f"Produto analítico ausente: {caminho}")
        return pd.DataFrame()
    return pd.read_csv(caminho)


def carregar_produtos(base: Path) -> dict[str, pd.DataFrame]:
    pasta = base / "dados_processados" / "biologia"
    arquivos = {
        "cursos": "cursos_biologia.csv",
        "base": "base_analitica_cursos.csv",
        "comparacoes": "comparacoes_regionais_validadas_sprint11.csv",
        "benchmark_soure": "benchmark_soure_cursos.csv",
        "benchmark_resumo": "benchmark_soure_resumo.csv",
        "sensibilidade": "sensibilidade_benchmark_soure.csv",
        "percentis": "percentis_soure.csv",
        "comparacao_focal": "comparacao_focal_soure.csv",
        "perfil_diferencial": "perfil_diferencial_soure.csv",
        "perfil_focal": "perfil_focal_soure_validado.csv",
        "processo_itens": "processo_formativo_soure_itens_validado.csv",
        "processo_bruto": "itens_processo_formativo.csv",
        "dimensoes": "dimensoes_processo_exploratorias.csv",
        "recomendacao_dist": "distribuicao_recomendacao.csv",
        "desempenho_individual": "desempenho_individual_soure_descritivas.csv",
        "correlacoes_individuais": "desempenho_individual_soure_correlacoes.csv",
        "associacoes": "associacoes_ecologicas_sprint11.csv",
        "outliers": "diagnostico_outliers_sprint11.csv",
        "auditoria": "auditoria_desempenho_sprint11.csv",
    }
    return {chave: _ler_csv(pasta / nome) for chave, nome in arquivos.items()}


def obter_soure(base: pd.DataFrame) -> pd.Series:
    mask = pd.to_numeric(base["CO_CURSO"], errors="coerce").eq(CO_CURSO_SOURE)
    soure = base.loc[mask]
    if len(soure) != 1:
        raise ValueError(
            f"Esperada exatamente uma oferta de Soure ({CO_CURSO_SOURE}); "
            f"encontradas {len(soure)}."
        )
    return soure.iloc[0]


def tabela_ofertas_ufpa(base: pd.DataFrame) -> pd.DataFrame:
    colunas = [
        "CO_CURSO",
        "ROTULO_OFERTA",
        "CONCEITO_ENADE_NUM",
        "INSCRITOS_NUM",
        "PARTICIPANTES_NUM",
        "TAXA_PARTICIPACAO_OFICIAL",
        "nt_ger_mean",
        "nt_obj_mean",
        "nt_dis_mean",
        "nt_ger_percentil_brasil",
    ]
    presentes = [c for c in colunas if c in base.columns]
    return (
        base.loc[base["CO_IES"].eq(BIOLOGIA.co_ies_focal), presentes]
        .sort_values("ROTULO_OFERTA")
        .reset_index(drop=True)
    )


def tabela_regional_nt_ger(comparacoes: pd.DataFrame) -> pd.DataFrame:
    colunas = [
        "RECORTE",
        "N_CURSOS",
        "N_PARTICIPANTES",
        "MEDIA_CURSOS",
        "MEDIA_PONDERADA_PARTICIPANTES",
        "MEDIANA_CURSOS",
        "DP_CURSOS",
        "P25",
        "P75",
    ]
    base = comparacoes.loc[comparacoes["INDICADOR"].eq("nt_ger_mean")].copy()
    presentes = [c for c in colunas if c in base.columns]
    return base[presentes].reset_index(drop=True)


def tabela_soure(base: pd.DataFrame) -> pd.DataFrame:
    soure = obter_soure(base)
    colunas = [
        "CO_CURSO",
        "ROTULO_OFERTA",
        "CONCEITO_ENADE_NUM",
        "INSCRITOS_NUM",
        "PARTICIPANTES_NUM",
        "TAXA_PARTICIPACAO_OFICIAL",
        "taxa_presenca_microdados",
        "nt_ger_mean",
        "nt_obj_mean",
        "nt_dis_mean",
        "nt_ger_percentil_brasil",
        "nt_ger_percentil_norte",
        "nt_ger_percentil_para",
    ]
    return pd.DataFrame([{c: soure.get(c) for c in colunas if c in base.columns}])


def tabela_perfil_soure(base: pd.DataFrame) -> pd.DataFrame:
    soure = obter_soure(base)
    colunas = [
        "ROTULO_OFERTA",
        "idade_media",
        "sexo_feminino_pct",
        "renda_ate_3sm_pct",
        "trabalha_pct",
        "acao_afirmativa_pct",
        "auxilio_permanencia_pct",
        "bolsa_academica_pct",
        "estudo_4h_ou_mais_pct",
        "escolaridade_mae_superior_pct",
        "escolaridade_pai_superior_pct",
        "qe_i68_media",
        "qe_i69_media",
        "qe_i70_interesse_pct",
    ]
    return pd.DataFrame([{c: soure.get(c) for c in colunas if c in base.columns}])


def tabela_benchmark_principal(sensibilidade: pd.DataFrame) -> pd.DataFrame:
    alvo = sensibilidade.loc[
        sensibilidade["CENARIO"].eq("estrutura_porte_0_5_2_0")
    ].copy()
    if alvo.empty:
        return sensibilidade.head(0).copy()
    colunas = [
        "CENARIO",
        "N_CURSOS",
        "nt_ger_mean_ALVO",
        "nt_ger_mean_MEDIA_BENCHMARK",
        "nt_ger_mean_DIFERENCA",
        "nt_obj_mean_ALVO",
        "nt_obj_mean_MEDIA_BENCHMARK",
        "nt_obj_mean_DIFERENCA",
        "nt_dis_mean_ALVO",
        "nt_dis_mean_MEDIA_BENCHMARK",
        "nt_dis_mean_DIFERENCA",
        "taxa_presenca_microdados_ALVO",
        "taxa_presenca_microdados_MEDIA_BENCHMARK",
        "taxa_presenca_microdados_DIFERENCA",
    ]
    presentes = [c for c in colunas if c in alvo.columns]
    return alvo[presentes].reset_index(drop=True)


def tabela_itens_processo(
    comparacao: pd.DataFrame,
    *,
    referencia: str = "Benchmark comparável",
    n: int = 10,
) -> pd.DataFrame:
    trabalho = comparacao.loc[comparacao["REFERENCIA"].eq(referencia)].copy()
    if trabalho.empty:
        return trabalho
    trabalho["ABS_DIF"] = pd.to_numeric(
        trabalho["DIFERENCA_SOURE_REFERENCIA"], errors="coerce"
    ).abs()
    trabalho["ROTULO_OFICIAL"] = trabalho["ITEM"].map(rotulo_item)
    colunas = [
        "ITEM",
        "ROTULO_OFICIAL",
        "MEDIA_SOURE",
        "N_VALIDO_SOURE",
        "MEDIA_REFERENCIA",
        "N_CURSOS_REFERENCIA",
        "DIFERENCA_SOURE_REFERENCIA",
    ]
    return (
        trabalho.sort_values("ABS_DIF", ascending=False)
        .head(n)[colunas]
        .reset_index(drop=True)
    )


def resumo_desempenho_geral(base: pd.DataFrame) -> str:
    ofertas = tabela_ofertas_ufpa(base)
    conceitos = sorted(
        int(x)
        for x in pd.to_numeric(
            ofertas["CONCEITO_ENADE_NUM"], errors="coerce"
        ).dropna().unique()
    )
    return (
        f"Foram localizadas {len(ofertas)} ofertas da UFPA. "
        f"Os Conceitos Enade observados entre elas são {conceitos}; "
        "não existe oferta da UFPA com Conceito Enade 1. Por isso, a análise não "
        "reproduz artificialmente o contraste das áreas anteriores e trata Soure "
        "como caso focal, mantendo as demais ofertas da UFPA e referências "
        "territoriais e estruturais como contrastes descritivos."
    )


def resumo_soure(base: pd.DataFrame, sensibilidade: pd.DataFrame) -> str:
    soure = obter_soure(base)
    principal = tabela_benchmark_principal(sensibilidade)
    texto = (
        f"Soure (CO_CURSO={CO_CURSO_SOURE}) apresentou média de NT_GER "
        f"{soure.get('nt_ger_mean', np.nan):.2f}, NT_OBJ "
        f"{soure.get('nt_obj_mean', np.nan):.2f} e NT_DIS "
        f"{soure.get('nt_dis_mean', np.nan):.2f}. "
        f"Seu percentil de NT_GER foi {soure.get('nt_ger_percentil_brasil', np.nan):.1f} "
        "no Brasil, "
        f"{soure.get('nt_ger_percentil_norte', np.nan):.1f} na Região Norte e "
        f"{soure.get('nt_ger_percentil_para', np.nan):.1f} no Pará."
    )
    if not principal.empty:
        linha = principal.iloc[0]
        dif_ger = linha.get("nt_ger_mean_DIFERENCA", np.nan)
        dif_obj = linha.get("nt_obj_mean_DIFERENCA", np.nan)
        dif_dis = linha.get("nt_dis_mean_DIFERENCA", np.nan)
        texto += (
            " No cenário estrutural principal, as diferenças Soure menos benchmark "
            f"foram {dif_ger:.2f} em NT_GER, {dif_obj:.2f} em NT_OBJ e "
            f"{dif_dis:.2f} em NT_DIS."
        )
        if pd.notna(dif_obj) and pd.notna(dif_dis) and abs(dif_obj) > abs(dif_dis):
            texto += (
                " A distância é mais pronunciada no componente objetivo do que "
                "no componente discursivo; esse padrão é descritivo e não causal."
            )
    return texto


def resumo_regional(comparacoes: pd.DataFrame) -> str:
    tab = tabela_regional_nt_ger(comparacoes).set_index("RECORTE")
    desejados = ["UFPA agregada", "Região Norte sem UFPA", "Brasil geral"]
    faltantes = [r for r in desejados if r not in tab.index]
    if faltantes:
        return (
            "Nem todos os recortes regionais obrigatórios estavam disponíveis "
            "para a síntese textual."
        )
    partes = []
    for recorte in desejados:
        linha = tab.loc[recorte]
        partes.append(
            f"{recorte}: média simples {linha['MEDIA_CURSOS']:.2f} e média "
            "ponderada por participantes "
            f"{linha['MEDIA_PONDERADA_PARTICIPANTES']:.2f}"
        )
    return "; ".join(partes) + "."


def resumo_associacoes(associacoes: pd.DataFrame) -> str:
    validas = associacoes.dropna(subset=["SPEARMAN_RHO"]).copy()
    if validas.empty:
        return "Não houve associações ecológicas calculáveis."
    idx = validas["SPEARMAN_RHO"].abs().idxmax()
    linha = validas.loc[idx]
    return (
        "A associação ecológica de maior magnitude absoluta entre as examinadas foi "
        f"entre {linha['INDICADOR_X']} e NT_GER médio "
        f"(Spearman rho={linha['SPEARMAN_RHO']:.3f}; "
        f"N={int(linha['N_CURSOS'])} cursos). Trata-se de associação entre cursos, "
        "não entre estudantes, e não sustenta inferência causal individual."
    )


def resumo_processo(itens: pd.DataFrame) -> str:
    tabela = tabela_itens_processo(itens, n=4)
    if tabela.empty:
        return "Não houve itens de processo formativo comparáveis para a síntese."
    negativos = tabela.sort_values("DIFERENCA_SOURE_REFERENCIA").head(2)
    positivos = tabela.sort_values(
        "DIFERENCA_SOURE_REFERENCIA", ascending=False
    ).head(2)
    neg = "; ".join(
        f"{r.ITEM} ({r.ROTULO_OFICIAL}): {r.DIFERENCA_SOURE_REFERENCIA:+.2f}"
        for r in negativos.itertuples()
    )
    pos = "; ".join(
        f"{r.ITEM} ({r.ROTULO_OFICIAL}): {r.DIFERENCA_SOURE_REFERENCIA:+.2f}"
        for r in positivos.itertuples()
    )
    return (
        "Na comparação com o benchmark, entre os itens de maior diferença absoluta, "
        f"os contrastes negativos incluem {neg}. Os positivos incluem {pos}. "
        "Os rótulos usados são os textos oficiais do Inep; mesmo assim, os itens "
        "permanecem analisados separadamente e não são condensados em índice único."
    )


def _p(doc: Document, texto: str, estilo: str | None = None):
    p = doc.add_paragraph(style=estilo)
    p.add_run(texto)
    return p


def _capa(doc: Document) -> None:
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = 0
    r = p.add_run("UNIVERSIDADE FEDERAL DO PARÁ")
    r.bold = True
    r.font.size = Pt(12)
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = 0
    r = p.add_run(
        "CIÊNCIAS BIOLÓGICAS NO ENADE 2025: PANORAMA DA UFPA E "
        "ESTUDO FOCAL DA OFERTA DE SOURE"
    )
    r.bold = True
    r.font.size = Pt(14)
    for _ in range(7):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = 0
    p.add_run("Autor(a): [preencher]")
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = 0
    p.add_run("Belém\n2026")


def _figura(
    doc: Document,
    base: Path,
    nome: str,
    numero: int,
    titulo: str,
    leitura: str,
    limitacao: str,
) -> None:
    adicionar_figura(
        doc,
        base / "figuras" / "biologia" / nome,
        f"Figura {numero} – {titulo}",
        FONTE_DADOS,
    )
    _p(doc, f"Descrição e interpretação: {leitura}")
    _p(doc, f"Limitação: {limitacao}")


def _markdown(dados: dict[str, pd.DataFrame]) -> str:
    base = dados["base"]
    ofertas = tabela_ofertas_ufpa(base)
    soure = tabela_soure(base)
    regional = tabela_regional_nt_ger(dados["comparacoes"])
    perfil = tabela_perfil_soure(base)
    benchmark = tabela_benchmark_principal(dados["sensibilidade"])
    processo = tabela_itens_processo(dados["processo_itens"], n=10)

    linhas = [
        "# CIÊNCIAS BIOLÓGICAS NO ENADE 2025: PANORAMA DA UFPA E ESTUDO FOCAL DA OFERTA DE SOURE",
        "",
        "## RESUMO",
        "",
        f"Este relatório analisa {len(base)} cursos de Ciências Biológicas no Enade 2025, "
        f"com {len(ofertas)} ofertas localizadas da UFPA e aprofundamento da oferta de "
        f"Soure (CO_CURSO={CO_CURSO_SOURE}). A unidade principal é CO_CURSO. Arquivos "
        "temáticos foram processados separadamente, agregados por curso e somente então "
        "relacionados por junções one-to-one. Não existe oferta da UFPA com Conceito Enade 1 "
        "na área, razão pela qual o estudo não cria artificialmente esse grupo. São examinados "
        "desempenho, composição discente, trajetória acadêmica, processo formativo, recomendação, "
        "comparações regionais, benchmark estrutural de Soure e associações ecológicas. "
        "A interpretação é descritiva e não causal.",
        "",
        "**Palavras-chave:** Enade; Ciências Biológicas; UFPA; Soure; formação de professores; microdados.",
        "",
        "## ABSTRACT",
        "",
        "This technical-scientific report analyzes Biological Sciences teacher education programs "
        "in the 2025 Enade, focusing on UFPA and especially on the Soure offer. CO_CURSO is the "
        "main unit of analysis. Thematic files were processed separately, aggregated at course "
        "level and only then combined one-to-one. Since UFPA has no Biological Sciences offer "
        "with Enade Concept 1, no artificial Concept-1 group is created. Performance, student "
        "composition, academic trajectory, formative process, recommendation, regional benchmarks, "
        "a structural benchmark for Soure and ecological associations are examined. Results are "
        "descriptive and non-causal.",
        "",
        "**Keywords:** Enade; Biological Sciences; UFPA; Soure; teacher education; microdata.",
        "",
        "# 1 INTRODUÇÃO",
        "",
        "A pergunta central é: quais características de desempenho, composição discente, trajetória "
        "acadêmica, condições socioeconômicas e avaliação do processo formativo caracterizam a "
        "oferta de Ciências Biológicas da UFPA em Soure, e como ela se posiciona em relação às "
        "demais ofertas da UFPA e a cursos comparáveis no Pará, na Região Norte e no Brasil?",
        "",
        "# 2 REFERENCIAL INSTITUCIONAL E METODOLÓGICO",
        "",
        "O Conceito Enade é tratado como classificação externa do curso. Ausência de conceito não é "
        "recodificada. Como não existe oferta da UFPA com Conceito Enade 1 em Ciências Biológicas, "
        "o contraste principal é focal: Soure versus demais ofertas da UFPA e referências "
        "territoriais e estruturalmente comparáveis.",
        "",
        "# 3 METODOLOGIA",
        "",
        "A unidade principal é CO_CURSO. Cada arquivo temático foi tratado separadamente, com "
        "tratamento de ausentes e agregação por curso antes das junções one-to-one. Não foi usada "
        "posição de linha nem identificador artificial. Análises individuais foram restritas a "
        "variáveis do mesmo arquivo temático. Relações entre indicadores de arquivos diferentes "
        "foram examinadas apenas no nível agregado do curso, com ressalva de falácia ecológica.",
        "",
        "# 4 PANORAMA DE CIÊNCIAS BIOLÓGICAS",
        "",
        resumo_desempenho_geral(base),
        "",
        "## Tabela 1 – Ofertas da UFPA",
        "",
        ofertas.to_markdown(index=False, floatfmt=".3f"),
        "",
        "# 5 RESULTADOS",
        "",
        "## 5.1 Desempenho",
        "",
        "O panorama geral combina posição relativa das ofertas da UFPA e distribuição de NT_GER, "
        "NT_OBJ e NT_DIS, mantendo N e dispersão. Componentes da nota são relacionados entre si "
        "apenas dentro do arquivo de desempenho.",
        "",
        "## 5.2 Perfil demográfico e socioeconômico",
        "",
        "Indicadores demográficos e socioeconômicos são agregados por CO_CURSO, com percentuais "
        "calculados sobre respostas válidas e registro de ausências. Relações com desempenho "
        "entre arquivos distintos são ecológicas.",
        "",
        "## 5.3 Trajetória e condições acadêmicas",
        "",
        "A trajetória é examinada por turno, tempo desde o ingresso, trabalho, bolsas, auxílios, "
        "horas de estudo e demais variáveis disponíveis, sempre em nível de curso.",
        "",
        "## 5.4 Processo formativo",
        "",
        "QE_I20–QE_I66 foram examinados item a item. Os textos oficiais do Inep foram vinculados "
        "aos códigos no relatório final. A escala válida de concordância vai de 1 a 6; respostas "
        "7 (não sei responder) e 8 (não se aplica) não são tratadas como concordância. Não foi "
        "criado índice único.",
        "",
        resumo_processo(dados["processo_itens"]),
        "",
        processo.to_markdown(index=False, floatfmt=".3f"),
        "",
        "## 5.5 Recomendação",
        "",
        "QE_I68 corresponde oficialmente à recomendação do curso e QE_I69 à recomendação da IES. "
        "QE_I70 mede interesse em participar da Prova Nacional Docente 2025. Esses itens não são "
        "renomeados genericamente como satisfação.",
        "",
        "## 5.6 Benchmark comparável",
        "",
        "O benchmark focal de Soure é recalculado sob filtros progressivos de modalidade, categoria "
        "administrativa, organização acadêmica e porte. O cenário estrutural principal usa cursos "
        "externos à UFPA com mesma modalidade, categoria e organização e porte entre 0,5x e 2x "
        "o número de participantes de Soure.",
        "",
        benchmark.to_markdown(index=False, floatfmt=".3f"),
        "",
        "## 5.7 Associações ecológicas",
        "",
        resumo_associacoes(dados["associacoes"]),
        "",
        "## 5.8 Comparação regional e nacional",
        "",
        resumo_regional(dados["comparacoes"]),
        "",
        regional.to_markdown(index=False, floatfmt=".3f"),
        "",
        "## 5.9 Estudo focal da oferta de Soure",
        "",
        "### 5.9.1 Participação e desempenho",
        "",
        resumo_soure(base, dados["sensibilidade"]),
        "",
        soure.to_markdown(index=False, floatfmt=".3f"),
        "",
        "### 5.9.2 Perfil discente",
        "",
        "O perfil de Soure é apresentado separadamente e contrastado com UFPA sem Soure, Pará, "
        "Norte, Brasil e benchmark comparável. Diferenças de composição são padrões agregados.",
        "",
        perfil.to_markdown(index=False, floatfmt=".3f"),
        "",
        "### 5.9.3 Trajetória acadêmica",
        "",
        "Os indicadores de trajetória de Soure são comparados em nível de curso. Não se vinculam "
        "notas individuais a respostas de questionários provenientes de arquivos distintos.",
        "",
        "### 5.9.4 Processo formativo",
        "",
        resumo_processo(dados["processo_itens"]),
        "",
        "### 5.9.5 Recomendação",
        "",
        "A recomendação de curso e IES é lida pelos textos oficiais de QE_I68 e QE_I69, enquanto "
        "QE_I70 é tratado como interesse na Prova Nacional Docente. O padrão é comparado com "
        "referências sem converter esses itens em uma medida genérica de satisfação.",
        "",
        "### 5.9.6 Benchmark comparável",
        "",
        "A robustez é avaliada por cinco cenários de seleção. Diferenças persistentes sob cenários "
        "mais restritivos fortalecem a descrição de um contraste estrutural, mas não identificam "
        "mecanismos causais.",
        "",
        "### 5.9.7 Síntese dos pontos distintivos",
        "",
        "A síntese integra desempenho, participação, perfil discente, trajetória, processo formativo "
        "e recomendação. O padrão central a ser destacado é a coexistência de desempenho objetivo "
        "relativamente mais distante do benchmark com componente discursivo mais próximo e "
        "indicadores de recomendação que não reproduzem necessariamente o mesmo gradiente.",
        "",
        "# 6 DISCUSSÃO",
        "",
        "Os resultados mostram que Soure deve ser interpretada como caso institucional específico, "
        "não como sinônimo de baixo desempenho nem como evidência de efeito de localização. "
        "Diferenças em participação, composição discente, trajetória, formação e recomendação "
        "podem coexistir, mas o desenho observacional e ecológico não permite atribuição causal. "
        "A comparação com cursos estruturalmente semelhantes reduz parte da heterogeneidade "
        "observável, sem equivaler a pareamento causal.",
        "",
        "# 7 CONCLUSÃO",
        "",
        "O panorama de Ciências Biológicas da UFPA é globalmente distinto das áreas previamente "
        "estudadas porque não há oferta com Conceito Enade 1. A análise focal identifica Soure "
        "como uma oferta que merece leitura multidimensional: sua posição relativa e seu contraste "
        "com benchmarks são mais pronunciados em NT_GER e NT_OBJ do que em NT_DIS, enquanto "
        "perfil, processo formativo e recomendação apresentam padrões próprios. As evidências "
        "servem para orientar investigação institucional e priorização de aprofundamentos, não "
        "para inferência causal sobre estudantes, campus ou modalidade.",
        "",
        "# REFERÊNCIAS",
        "",
        "As referências bibliográficas e normativas são inseridas integralmente na versão DOCX pelo "
        "módulo compartilhado de referências do projeto.",
        "",
        "# APÊNDICES",
        "",
        "## APÊNDICE A – REGRAS DE INTEGRIDADE",
        "",
        "Não foram realizadas junções individuais entre arquivos temáticos. Todas as integrações "
        "ocorreram após agregação por CO_CURSO e validação one-to-one. Ausência de conceito não "
        "equivale a Conceito 1. Análises individuais usam somente variáveis do mesmo arquivo.",
        "",
        "## APÊNDICE B – APROFUNDAMENTOS SUGERIDOS",
        "",
        "1. Conteúdo objetivo e discursivo em Soure — pergunta: quais subcomponentes do desempenho "
        "explicam descritivamente a distância observada? Variáveis: NT_GER, NT_OBJ, NT_DIS, "
        "QT_ACERTOS e PROFICIENCIA. Método: distribuição individual no mesmo arquivo, percentis e "
        "sensibilidade. Limitação: relações parcialmente mecânicas entre indicadores.",
        "2. Participação e estabilidade — pergunta: a posição relativa de Soure se mantém em cursos "
        "com taxas de participação semelhantes? Variáveis: inscritos, participantes e presença. "
        "Método: estratificação e sensibilidade ecológica. Limitação: participação pode refletir "
        "fatores não observados.",
        "3. Processo formativo item a item — pergunta: quais itens oficiais de QE_I20–QE_I66 "
        "concentram contrastes robustos? Método: diferenças de médias por item, N válido e análise "
        "de sensibilidade. Limitação: escala ordinal e respostas autorreferidas.",
        "4. Perfil socioeconômico e trajetória — pergunta: quais características agregadas distinguem "
        "Soure de cursos comparáveis? Variáveis: renda, trabalho, ação afirmativa, auxílios, "
        "escolaridade parental, estudo e trajetória. Método: diferenças padronizadas e associações "
        "ecológicas. Limitação: falácia ecológica.",
        "5. Recomendação e formação — pergunta: a recomendação do curso e da IES acompanha os itens "
        "formativos em cursos comparáveis? Variáveis: QE_I68, QE_I69, QE_I70 e QE_I20–QE_I66. "
        "Método: associações agregadas por curso. Limitação: não representa relação individual entre "
        "respostas de arquivos distintos.",
        "",
        "## APÊNDICE C – RÓTULOS OFICIAIS DOS ITENS QE_I20–QE_I66",
        "",
    ]

    for codigo in [f"QE_I{i}" for i in range(20, 67)]:
        linhas.append(f"- **{codigo}** — {ROTULOS_QE[codigo]}")
    linhas.extend(
        [
            "",
            f"Fonte dos rótulos: {FONTE_QUESTIONARIO}",
        ]
    )
    return "\n".join(linhas)


def gerar_relatorio(
    base_projeto: Path,
    saida_docx: Path,
    saida_md: Path,
) -> dict:
    dados = carregar_produtos(base_projeto)
    base = dados["base"]
    ofertas = tabela_ofertas_ufpa(base)
    regional = tabela_regional_nt_ger(dados["comparacoes"])
    soure = tabela_soure(base)
    perfil = tabela_perfil_soure(base)
    benchmark = tabela_benchmark_principal(dados["sensibilidade"])
    processo = tabela_itens_processo(dados["processo_itens"], n=10)

    doc = Document()
    configurar_documento(doc)
    _capa(doc)
    doc.add_page_break()

    doc.add_heading("RESUMO", level=1)
    _p(
        doc,
        f"Este relatório analisa {len(base)} cursos de Ciências Biológicas no Enade 2025, "
        f"com {len(ofertas)} ofertas localizadas da UFPA e aprofundamento da oferta de Soure "
        f"(CO_CURSO={CO_CURSO_SOURE}). A unidade principal é CO_CURSO. Arquivos temáticos foram "
        "processados separadamente, agregados por curso e somente então relacionados por junções "
        "one-to-one. Não existe oferta da UFPA com Conceito Enade 1 na área. São examinados "
        "desempenho, composição discente, trajetória acadêmica, processo formativo, recomendação, "
        "comparações regionais, benchmark estrutural de Soure e associações ecológicas. "
        "A interpretação é descritiva e não causal.",
        "Resumo",
    )
    _p(
        doc,
        "Palavras-chave: Enade; Ciências Biológicas; UFPA; Soure; formação de professores; microdados.",
    )

    doc.add_heading("ABSTRACT", level=1)
    _p(
        doc,
        "This technical-scientific report analyzes Biological Sciences teacher education programs "
        "in the 2025 Enade, focusing on UFPA and especially on the Soure offer. CO_CURSO is the "
        "main unit of analysis. Thematic files were processed separately, aggregated at course "
        "level and only then combined one-to-one. Since UFPA has no Biological Sciences offer "
        "with Enade Concept 1, no artificial Concept-1 group is created. Results are descriptive "
        "and non-causal.",
        "Resumo",
    )
    _p(
        doc,
        "Keywords: Enade; Biological Sciences; UFPA; Soure; teacher education; microdata.",
    )

    doc.add_heading("1 INTRODUÇÃO", level=1)
    _p(
        doc,
        "O relatório responde à seguinte pergunta: quais características de desempenho, composição "
        "discente, trajetória acadêmica, condições socioeconômicas e avaliação do processo formativo "
        "caracterizam a oferta de Ciências Biológicas da UFPA em Soure, e como ela se posiciona em "
        "relação às demais ofertas da UFPA e a cursos comparáveis no Pará, na Região Norte e no Brasil?",
    )

    doc.add_heading("2 REFERENCIAL INSTITUCIONAL E METODOLÓGICO", level=1)
    _p(
        doc,
        "O Conceito Enade é tratado como classificação externa do curso. A ausência de conceito "
        "permanece distinta de Conceito 1. Como não existe oferta da UFPA com Conceito Enade 1 em "
        "Ciências Biológicas, o desenho focaliza Soure e utiliza as demais ofertas da UFPA e "
        "referências territoriais e estruturais como contrastes descritivos.",
    )

    doc.add_heading("3 METODOLOGIA", level=1)
    _p(
        doc,
        "A unidade de análise é CO_CURSO. Cada arquivo temático foi tratado separadamente, com "
        "tratamento de ausentes e agregação por curso antes de qualquer junção. Não foi usada "
        "posição de linha nem criado identificador artificial. As junções entre tabelas agregadas "
        "são validadas como one-to-one. Análises individuais usam somente variáveis do mesmo "
        "arquivo temático; relações entre arquivos distintos são ecológicas.",
    )
    _p(
        doc,
        "O caso focal é Soure (CO_CURSO 104640). As referências incluem UFPA sem Soure, outras IES "
        "do Pará, Norte sem Pará, Brasil sem Norte e benchmark estrutural comparável. O benchmark "
        "é submetido a cinco cenários de sensibilidade por modalidade, categoria administrativa, "
        "organização acadêmica e porte.",
    )

    doc.add_heading("4 PANORAMA DE CIÊNCIAS BIOLÓGICAS", level=1)
    _p(doc, resumo_desempenho_geral(base))
    adicionar_tabela(doc, "Tabela 1 – Ofertas de Ciências Biológicas da UFPA", ofertas, FONTE_DADOS)
    _figura(
        doc,
        base_projeto,
        "01_painel_ofertas_ufpa.png",
        1,
        "Ofertas de Ciências Biológicas da UFPA",
        "O painel reúne as ofertas da UFPA, seus conceitos, porte e proficiência.",
        "A leitura é descritiva e não controla diferenças estruturais entre ofertas.",
    )

    doc.add_heading("5 RESULTADOS", level=1)

    doc.add_heading("5.1 Desempenho", level=2)
    _p(doc, resumo_desempenho_geral(base))
    _figura(
        doc,
        base_projeto,
        "02_posicao_relativa_nt_ger.png",
        2,
        "Posição relativa das ofertas em NT_GER",
        "As ofertas da UFPA são situadas na distribuição nacional da própria área.",
        "Percentis não eliminam diferenças de composição, porte e participação.",
    )
    _figura(
        doc,
        base_projeto,
        "03_distribuicao_nt_ger_foco_soure.png",
        3,
        "Distribuição de NT_GER com foco em Soure",
        "A distribuição individual usa somente o arquivo temático de desempenho.",
        "Cursos com N menor apresentam maior instabilidade amostral.",
    )
    _figura(
        doc,
        base_projeto,
        "04_distribuicao_nt_obj_foco_soure.png",
        4,
        "Distribuição de NT_OBJ com foco em Soure",
        "O componente objetivo é examinado separadamente da nota geral.",
        "NT_OBJ e NT_GER possuem relação mecânica e não são evidências independentes.",
    )
    _figura(
        doc,
        base_projeto,
        "05_distribuicao_nt_dis_foco_soure.png",
        5,
        "Distribuição de NT_DIS com foco em Soure",
        "O componente discursivo é preservado em sua escala específica.",
        "Diferenças de correção e N válido devem ser consideradas.",
    )

    doc.add_heading("5.2 Perfil demográfico e socioeconômico", level=2)
    _p(
        doc,
        "Os percentuais utilizam denominadores válidos por item. Renda, trabalho, ação afirmativa, "
        "auxílios, bolsas, escolaridade parental e horas de estudo são tratados como composição "
        "agregada do curso.",
    )
    _figura(
        doc,
        base_projeto,
        "06_perfil_socioeconomico_foco_soure.png",
        6,
        "Perfil socioeconômico com foco em Soure",
        "O gráfico situa o perfil agregado de Soure frente às referências.",
        "Não se infere associação individual entre perfil e desempenho.",
    )

    doc.add_heading("5.3 Trajetória e condições acadêmicas", level=2)
    _p(
        doc,
        "Turno, tempo desde o ingresso, trabalho, bolsas, auxílios e horas de estudo são "
        "analisados no nível do curso. O relatório preserva N válido e ausências e evita "
        "relacionar individualmente respostas de arquivos temáticos diferentes.",
    )

    doc.add_heading("5.4 Processo formativo", level=2)
    _p(
        doc,
        "Os 47 itens QE_I20–QE_I66 foram vinculados aos textos oficiais do Inep. A escala válida "
        "de concordância vai de 1 a 6; categorias de não resposta substantiva não são incorporadas "
        "como concordância. Não foi formado índice único. Qualquer agrupamento dimensional permanece "
        "exploratório até validação teórica.",
    )
    _p(doc, resumo_processo(dados["processo_itens"]))
    adicionar_tabela(
        doc,
        "Tabela 2 – Itens de processo formativo com maiores diferenças absolutas em Soure",
        processo,
        FONTE_DADOS,
    )
    _figura(
        doc,
        base_projeto,
        "07_processo_formativo_foco_soure.png",
        7,
        "Processo formativo com foco em Soure",
        "A figura mantém os itens formativos individualizados.",
        "Médias de itens distintos não constituem automaticamente uma escala validada.",
    )

    doc.add_heading("5.5 Recomendação", level=2)
    _p(
        doc,
        f"{rotulo_item('QE_I68')} {rotulo_item('QE_I69')} "
        f"QE_I70 corresponde a: {rotulo_item('QE_I70')} "
        "Os itens são apresentados por seus significados oficiais, sem serem convertidos "
        "automaticamente em satisfação geral.",
    )
    _figura(
        doc,
        base_projeto,
        "12_recomendacao_foco_soure.png",
        8,
        "Recomendação e interesse associados à oferta focal",
        "Os itens oficiais de recomendação e interesse são comparados entre referências.",
        "Respostas autorreferidas não equivalem a uma medida causal de qualidade.",
    )

    doc.add_heading("5.6 Benchmark comparável", level=2)
    _p(
        doc,
        "O benchmark de Soure restringe progressivamente modalidade, categoria administrativa, "
        "organização acadêmica e porte, excluindo a UFPA da referência. O cenário estrutural "
        "principal usa faixa de 0,5x a 2x participantes.",
    )
    if not benchmark.empty:
        adicionar_tabela(
            doc,
            "Tabela 3 – Cenário estrutural principal do benchmark de Soure",
            benchmark,
            FONTE_DADOS,
        )
    _figura(
        doc,
        base_projeto,
        "11_benchmark_soure.png",
        9,
        "Soure e benchmark comparável",
        "A figura contrasta Soure com cursos estruturalmente semelhantes.",
        "Sem desenho causal, similaridade observável não equivale a contrafactual.",
    )

    doc.add_heading("5.7 Associações ecológicas", level=2)
    _p(doc, resumo_associacoes(dados["associacoes"]))

    doc.add_heading("5.8 Comparação regional e nacional", level=2)
    _p(doc, resumo_regional(dados["comparacoes"]))
    adicionar_tabela(
        doc,
        "Tabela 4 – Comparação regional e nacional de NT_GER",
        regional,
        FONTE_DADOS,
    )
    _figura(
        doc,
        base_projeto,
        "08_comparacao_regional_nacional.png",
        10,
        "Comparação regional e nacional",
        "Médias simples e ponderadas por participantes são apresentadas separadamente.",
        "Recortes nacionais e regionais sobrepostos são benchmarks descritivos, não grupos independentes.",
    )

    doc.add_heading("5.9 Estudo focal da oferta de Soure", level=2)

    doc.add_heading("5.9.1 Participação e desempenho", level=3)
    _p(doc, resumo_soure(base, dados["sensibilidade"]))
    adicionar_tabela(doc, "Tabela 5 – Síntese da oferta de Soure", soure, FONTE_DADOS)
    _figura(
        doc,
        base_projeto,
        "09_desempenho_ofertas_ufpa.png",
        11,
        "Desempenho das ofertas da UFPA",
        "Soure é posicionada frente às demais ofertas da instituição.",
        "Diferenças internas à UFPA permanecem descritivas.",
    )
    _figura(
        doc,
        base_projeto,
        "10_percentis_soure.png",
        12,
        "Percentis de Soure",
        "Os percentis situam NT_GER de Soure no Pará, Norte e Brasil.",
        "Percentis dependem da distribuição de cursos de cada referência.",
    )
    _figura(
        doc,
        base_projeto,
        "validada_14_participacao_ufpa.png",
        13,
        "Auditoria de participação das ofertas da UFPA",
        "Participação oficial e presença observada nos microdados são confrontadas.",
        "A concordância entre fontes não elimina possível viés de não participação.",
    )
    _figura(
        doc,
        base_projeto,
        "validada_15_componentes_soure.png",
        14,
        "Componentes de desempenho de Soure",
        "A validação separa NT_GER, NT_OBJ e NT_DIS no contraste focal.",
        "Os componentes têm escalas e relações mecânicas próprias.",
    )

    doc.add_heading("5.9.2 Perfil discente", level=3)
    _p(
        doc,
        "O perfil focal considera idade, sexo, renda, trabalho, ação afirmativa, auxílios, bolsas, "
        "horas de estudo e escolaridade parental quando disponíveis.",
    )
    adicionar_tabela(doc, "Tabela 6 – Perfil agregado de Soure", perfil, FONTE_DADOS)
    _figura(
        doc,
        base_projeto,
        "13_perfil_diferencial_soure.png",
        15,
        "Perfil diferencial de Soure",
        "Diferenças e distâncias padronizadas mostram em quais indicadores Soure se afasta das referências.",
        "Distância padronizada entre cursos não é efeito causal.",
    )
    _figura(
        doc,
        base_projeto,
        "validada_18_perfil_focal_soure.png",
        16,
        "Validação do perfil focal",
        "A figura revisita indicadores selecionados sob a validação da Sprint 11.",
        "O número de respostas válidas varia por indicador.",
    )

    doc.add_heading("5.9.3 Trajetória acadêmica", level=3)
    _p(
        doc,
        "A trajetória de Soure é interpretada em conjunto com seus contrastes agregados. Turno, "
        "tempo desde ingresso, trabalho e estudo são descritos sem cruzamento individual com notas "
        "de outros arquivos.",
    )

    doc.add_heading("5.9.4 Processo formativo", level=3)
    _p(doc, resumo_processo(dados["processo_itens"]))
    _figura(
        doc,
        base_projeto,
        "validada_17_processo_itens_soure.png",
        17,
        "Itens de processo formativo com maiores diferenças",
        "Os códigos são apresentados com interpretação vinculada ao texto oficial do questionário.",
        "A seleção por diferença absoluta é exploratória e sensível ao N válido.",
    )

    doc.add_heading("5.9.5 Recomendação", level=3)
    _p(
        doc,
        "A recomendação do curso e da IES é interpretada pelos itens oficiais QE_I68 e QE_I69. "
        "O interesse na Prova Nacional Docente é mantido como constructo distinto em QE_I70.",
    )
    _figura(
        doc,
        base_projeto,
        "validada_19_recomendacao_soure.png",
        18,
        "Validação dos itens de recomendação de Soure",
        "A figura contrasta os itens de recomendação no estudo focal.",
        "Recomendação autorreferida não é sinônimo automático de avaliação global da formação.",
    )

    doc.add_heading("5.9.6 Benchmark comparável", level=3)
    _p(
        doc,
        "A análise de sensibilidade testa cinco cenários progressivamente restritivos. O objetivo "
        "é verificar a estabilidade descritiva do contraste, não produzir estimativa causal.",
    )
    _figura(
        doc,
        base_projeto,
        "validada_16_sensibilidade_benchmark_soure.png",
        19,
        "Sensibilidade do benchmark de Soure",
        "A figura mostra como as diferenças mudam sob critérios alternativos de comparabilidade.",
        "Cenários mais restritivos podem reduzir N e aumentar instabilidade.",
    )

    doc.add_heading("5.9.7 Síntese dos pontos distintivos", level=3)
    _p(
        doc,
        resumo_soure(base, dados["sensibilidade"])
        + " O conjunto deve ser lido junto ao perfil, ao processo formativo e à recomendação; "
        "nenhuma dimensão isolada explica causalmente o resultado observado.",
    )

    doc.add_heading("6 DISCUSSÃO", level=1)
    _p(
        doc,
        "Soure apresenta uma configuração multidimensional que não pode ser reduzida ao Conceito Enade. "
        "O contraste com cursos comparáveis é mais pronunciado em NT_GER e NT_OBJ do que em NT_DIS, "
        "enquanto participação, perfil socioeconômico, processo formativo e recomendação apresentam "
        "padrões próprios. O desenho não permite concluir que localização, perfil discente, infraestrutura "
        "ou qualquer indicador isolado cause as diferenças observadas.",
    )

    doc.add_heading("7 CONCLUSÃO", level=1)
    _p(
        doc,
        "Ciências Biológicas da UFPA apresenta conceitos relativamente altos e não contém oferta com "
        "Conceito Enade 1. O estudo focal de Soure mostra que sua posição deve ser compreendida por meio "
        "de múltiplas dimensões. A evidência mais consistente da validação é a maior distância no "
        "componente objetivo frente ao benchmark estrutural, sem contraste equivalente no componente "
        "discursivo. Os resultados devem orientar aprofundamentos institucionais, não inferência causal.",
    )

    doc.add_heading("REFERÊNCIAS", level=1)
    adicionar_referencias(doc)

    doc.add_heading("APÊNDICES", level=1)
    doc.add_heading("APÊNDICE A – REGRAS DE INTEGRIDADE", level=2)
    _p(
        doc,
        "Não foram realizadas junções individuais entre arquivos temáticos. Todas as integrações ocorreram "
        "após agregação por CO_CURSO e validação one-to-one. Ausência de conceito não equivale a Conceito 1. "
        "Análises individuais usam somente variáveis do mesmo arquivo temático.",
    )

    doc.add_heading("APÊNDICE B – APROFUNDAMENTOS SUGERIDOS", level=2)
    aprofundamentos = [
        (
            "Conteúdo objetivo e discursivo",
            "Quais subcomponentes do desempenho concentram a distância de Soure?",
            "NT_GER, NT_OBJ, NT_DIS, QT_ACERTOS e PROFICIENCIA",
            "distribuições individuais no mesmo arquivo, percentis e sensibilidade",
            "relações parcialmente mecânicas",
        ),
        (
            "Participação e estabilidade",
            "A posição de Soure se mantém em cursos com participação semelhante?",
            "inscritos, participantes, presença e desempenho agregado",
            "estratificação e sensibilidade ecológica",
            "participação pode refletir fatores não observados",
        ),
        (
            "Processo formativo item a item",
            "Quais itens oficiais apresentam contrastes robustos?",
            "QE_I20–QE_I66",
            "diferenças por item, N válido e sensibilidade",
            "escala ordinal e respostas autorreferidas",
        ),
        (
            "Perfil e trajetória",
            "Quais características agregadas distinguem Soure de cursos comparáveis?",
            "renda, trabalho, ação afirmativa, auxílios, escolaridade parental e trajetória",
            "diferenças padronizadas e associações ecológicas",
            "falácia ecológica",
        ),
        (
            "Recomendação e formação",
            "A recomendação acompanha os itens formativos em cursos comparáveis?",
            "QE_I68, QE_I69, QE_I70 e QE_I20–QE_I66",
            "associações agregadas por curso",
            "não representa relação individual entre arquivos distintos",
        ),
    ]
    adicionar_tabela(
        doc,
        "Tabela A1 – Aprofundamentos sugeridos",
        pd.DataFrame(
            aprofundamentos,
            columns=["Tema", "Pergunta", "Variáveis", "Método", "Limitação"],
        ),
        "Elaboração própria.",
    )

    doc.add_heading(
        "APÊNDICE C – RÓTULOS OFICIAIS DOS ITENS QE_I20–QE_I66",
        level=2,
    )
    mapa = pd.DataFrame(
        [(codigo, ROTULOS_QE[codigo]) for codigo in [f"QE_I{i}" for i in range(20, 67)]],
        columns=["ITEM", "TEXTO_OFICIAL"],
    )
    adicionar_tabela(
        doc,
        "Tabela A2 – Mapeamento oficial QE_I20–QE_I66",
        mapa,
        FONTE_QUESTIONARIO,
    )

    configurar_cabecalho_rodape(doc)
    saida_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(saida_docx)

    saida_md.parent.mkdir(parents=True, exist_ok=True)
    saida_md.write_text(_markdown(dados), encoding="utf-8")

    conversao = converter_docx_para_pdf(saida_docx, saida_docx.parent)
    return ResultadoRelatorio(
        docx=saida_docx,
        markdown=saida_md,
        conversao_pdf=conversao,
    ).como_dict()
