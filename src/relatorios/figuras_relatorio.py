from __future__ import annotations

from pathlib import Path

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


def adicionar_figura(doc, caminho: Path, titulo: str, fonte: str, largura_cm: float = 15.5) -> None:
    if not caminho.exists():
        p = doc.add_paragraph()
        p.add_run(f"[Figura não localizada: {caminho.name}]").bold = True
        return
    p_titulo = doc.add_paragraph()
    p_titulo.paragraph_format.first_line_indent = 0
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_titulo.add_run(titulo)
    r.bold = True
    r.font.size = Pt(10)

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = 0
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(caminho), width=Cm(largura_cm))

    pf = doc.add_paragraph(style="Fonte")
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.add_run(f"Fonte: {fonte}")
