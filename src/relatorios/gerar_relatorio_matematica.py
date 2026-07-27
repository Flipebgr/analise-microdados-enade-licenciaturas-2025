from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from src.relatorios.figuras_relatorio import adicionar_figura
from src.relatorios.formatacao_abnt import configurar_cabecalho_rodape, configurar_documento, nova_secao
from src.relatorios.referencias import adicionar_referencias
from src.relatorios.tabelas_relatorio import adicionar_tabela


def _p(doc, texto: str, estilo: str | None = None):
    p = doc.add_paragraph(style=estilo)
    p.add_run(texto)
    return p


def _h(doc, texto: str, nivel: int = 1):
    return doc.add_heading(texto, level=nivel)


def _carregar(base: Path):
    dp = base / "dados_processados" / "matematica"
    cursos = pd.read_csv(dp / "cursos_matematica.csv")
    analitica = pd.read_csv(dp / "base_analitica_cursos.csv")
    bench = pd.read_csv(dp / "benchmarks_amplos.csv")
    sens = pd.read_csv(dp / "sensibilidade_benchmarks.csv")
    dims = pd.read_csv(dp / "diagnostico_dimensoes_processo.csv")
    return cursos, analitica, bench, sens, dims


def _tabela_ofertas(cursos: pd.DataFrame) -> pd.DataFrame:
    ufpa = cursos[cursos["CO_IES"] == 569].copy()
    return ufpa[["CO_CURSO", "ROTULO_OFERTA", "MODALIDADE", "INSCRITOS_NUM", "PARTICIPANTES_NUM", "PCT_PADRAO_PROFICIENCIA_NUM", "CONCEITO_ENADE"]].rename(columns={
        "CO_CURSO": "Código", "ROTULO_OFERTA": "Oferta", "MODALIDADE": "Modalidade", "INSCRITOS_NUM": "Inscritos", "PARTICIPANTES_NUM": "Participantes", "PCT_PADRAO_PROFICIENCIA_NUM": "% proficiência", "CONCEITO_ENADE": "Conceito",
    })


def _tabela_grupos(cursos: pd.DataFrame) -> pd.DataFrame:
    g = cursos.groupby(["GRUPO_CODIGO", "GRUPO"], dropna=False).agg(Cursos=("CO_CURSO", "nunique"), Participantes=("PARTICIPANTES_NUM", "sum")).reset_index()
    return g.rename(columns={"GRUPO_CODIGO": "Grupo", "GRUPO": "Definição"})


def _tabela_desempenho(analitica: pd.DataFrame) -> pd.DataFrame:
    g = analitica.groupby(["GRUPO_CODIGO", "GRUPO"]).agg(
        Cursos=("CO_CURSO", "nunique"),
        Média=("nt_ger_mean", "mean"),
        Mediana=("nt_ger_mean", "median"),
        DP=("nt_ger_mean", "std"),
        Participantes=("PARTICIPANTES_NUM", "sum"),
    ).reset_index()
    return g.rename(columns={"GRUPO_CODIGO": "Grupo", "GRUPO": "Definição"})


def _tabela_sintese(analitica: pd.DataFrame) -> pd.DataFrame:
    a = analitica[analitica["GRUPO_CODIGO"] == "A"]
    b = analitica[analitica["GRUPO_CODIGO"] == "B"]
    def m(df, col):
        return float(df[col].mean()) if col in df and not df.empty else float("nan")
    linhas = [
        ["Desempenho", m(a, "nt_ger_mean"), m(b, "nt_ger_mean"), "Diferença persistente em favor da oferta interna com conceito superior."],
        ["Participação", m(a, "taxa_presenca_microdados") * 100, m(b, "taxa_presenca_microdados") * 100, "A interpretação deve distinguir inscritos, participantes oficiais e notas válidas."],
        ["Renda até 3 SM", m(a, "renda_ate_3sm_pct") * 100, m(b, "renda_ate_3sm_pct") * 100, "Associação agregada; não permite inferência individual."],
        ["Trabalho", m(a, "trabalha_pct") * 100, m(b, "trabalha_pct") * 100, "Indicador de composição discente, não explicação causal."],
        ["Recomendação do curso (9-10)", m(a, "qe_i68_nota_9_10_pct") * 100, m(b, "qe_i68_nota_9_10_pct") * 100, "Deve ser interpretada como recomendação, não como satisfação geral."],
    ]
    return pd.DataFrame(linhas, columns=["Dimensão", "UFPA conceito 1", "UFPA conceito superior", "Leitura cautelosa"])


def gerar_relatorio(base: Path, saida_docx: Path, saida_md: Path, metadados: dict | None = None) -> None:
    metadados = metadados or {}
    cursos, analitica, _, sens, dims = _carregar(base)
    figdir = base / "figuras" / "matematica"

    doc = Document()
    configurar_documento(doc)
    configurar_cabecalho_rodape(doc)

    # Capa
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = 0
    r = p.add_run(metadados.get("instituicao", "UNIVERSIDADE FEDERAL DO PARÁ"))
    r.bold = True
    r.font.size = Pt(12)
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = 0
    r = p.add_run(
        "DESEMPENHO, COMPOSIÇÃO DISCENTE E PROCESSO FORMATIVO NOS CURSOS "
        "DE LICENCIATURA EM MATEMÁTICA DA UFPA NO ENADE 2025"
    )
    r.bold = True
    r.font.size = Pt(14)
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = 0
    p.add_run(metadados.get("autor", "Autor(a): [preencher]"))
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = 0
    p.add_run(f"{metadados.get('cidade', 'Belém')}\n2026")

    nova_secao(doc)
    _h(doc, "RESUMO", 1)
    _p(doc, "Este relatório analisa as ofertas de licenciatura em Matemática da Universidade Federal do Pará no Enade das Licenciaturas 2025, com ênfase nas sete ofertas enquadradas no Conceito Enade 1. A unidade principal é o curso, identificado por CO_CURSO. Os arquivos temáticos foram tratados separadamente e agregados por curso, sem reconstrução de registros individuais. Foram examinados desempenho, participação, composição demográfica e socioeconômica, trajetória acadêmica, avaliação do processo formativo, recomendação e benchmarks territorialmente amplos e institucionalmente comparáveis. Os resultados mostram que as ofertas da UFPA com conceito 1 apresentam desempenho médio inferior à oferta presencial de Belém, de conceito 3, e aos benchmarks comparáveis. As diferenças permanecem em direção sob múltiplos critérios de porte, embora sua magnitude varie. As associações entre desempenho e indicadores de composição são ecológicas e não permitem inferência causal ou individual.", "Resumo")
    _p(doc, "Palavras-chave: Enade; licenciatura em Matemática; UFPA; microdados; avaliação da educação superior.", "Resumo")
    _h(doc, "ABSTRACT", 1)
    _p(doc, "This report analyzes Mathematics teacher education programs offered by the Federal University of Pará in the 2025 Enade for Teacher Education Programs, emphasizing the seven programs classified in Enade Concept 1. The course, identified by CO_CURSO, is the main unit of analysis. Thematic files were processed separately and aggregated at course level, without reconstructing individual records. Performance, participation, demographic and socioeconomic composition, academic trajectory, formative process evaluation, recommendation, and broad and comparable benchmarks were examined. Results indicate lower average performance among UFPA Concept 1 programs than the Belém in-person program, classified as Concept 3, and comparable benchmarks. Directional differences remain stable across size criteria, although their magnitude varies. Associations are ecological and do not support individual or causal inference.", "Resumo")
    _p(doc, "Keywords: Enade; Mathematics teacher education; UFPA; microdata; higher education assessment.", "Resumo")

    _h(doc, "SUMÁRIO", 1)
    itens_sumario = [
        "1 INTRODUÇÃO",
        "2 REFERENCIAL INSTITUCIONAL E METODOLÓGICO",
        "3 METODOLOGIA",
        "4 PANORAMA DA LICENCIATURA EM MATEMÁTICA",
        "5 RESULTADOS",
        "6 DISCUSSÃO",
        "7 CONCLUSÃO",
        "REFERÊNCIAS",
        "APÊNDICES",
    ]
    for item in itens_sumario:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = 0
        p.add_run(item)
    _p(doc, "Nota: no Word, atualize os campos e o sumário após abrir o documento.", "Fonte")

    nova_secao(doc)
    _h(doc, "1 INTRODUÇÃO", 1)
    _p(doc, "O Exame Nacional de Desempenho dos Estudantes integra o Sistema Nacional de Avaliação da Educação Superior e produz evidências sobre o desempenho dos concluintes e o contexto de formação. A edição de 2025 introduziu uma configuração específica para as licenciaturas, associando a avaliação teórica à Prova Nacional Docente e mantendo instrumentos de caracterização do estudante e de avaliação do processo formativo.")
    _p(doc, "Na UFPA, Matemática constitui um caso analítico relevante: foram localizadas oito ofertas, das quais sete receberam Conceito Enade 1 e uma, a oferta presencial de Belém, recebeu conceito 3. O contraste interno reduz parte da heterogeneidade institucional, mas não elimina diferenças de modalidade, localização, porte e composição discente.")
    _p(doc, "O objetivo geral é investigar quais características de desempenho, composição discente, trajetória acadêmica e avaliação do processo formativo diferenciam as ofertas da UFPA com Conceito Enade 1 das demais ofertas de Matemática na própria universidade, no Pará, na Região Norte e no Brasil.")

    _h(doc, "2 REFERENCIAL INSTITUCIONAL E METODOLÓGICO", 1)
    _h(doc, "2.1 Enade das Licenciaturas e indicadores de qualidade", 2)
    _p(doc, "O Enade afere o desempenho dos estudantes em relação aos conteúdos, habilidades e competências previstos para a formação. O Conceito Enade é um indicador do curso em escala de 1 a 5 e deve ser distinguido do IDD e do CPC. A análise utiliza o Conceito Enade como classificação externa do curso e as variáveis contínuas dos microdados para descrever distribuições e diferenças.")
    _h(doc, "2.2 Estrutura e anonimização dos microdados", 2)
    _p(doc, "Os 28 arquivos foram fragmentados por tema e não contêm identificador público que permita reconhecer o mesmo estudante entre desempenho, perfil e percepção. A posição das linhas não é uma chave. Por isso, cada arquivo foi tratado isoladamente, agregado por CO_CURSO e somente então relacionado às demais tabelas agregadas.")
    _h(doc, "2.3 Limites da inferência", 2)
    _p(doc, "As relações entre nota média do curso, composição socioeconômica e avaliação do processo formativo são ecológicas. Elas descrevem covariação entre cursos, mas não autorizam concluir que uma característica individual causou o desempenho de um estudante ou o conceito do curso.")

    _h(doc, "3 METODOLOGIA", 1)
    _h(doc, "3.1 Delineamento e fontes", 2)
    _p(doc, "Trata-se de estudo quantitativo, exploratório e comparativo, baseado nos microdados do Enade das Licenciaturas 2025 e na planilha oficial de Conceito Enade por curso. O recorte abrange todos os cursos de Matemática identificados pelo CO_GRUPO 702.")
    _h(doc, "3.2 Unidade e grupos de análise", 2)
    _p(doc, "A unidade principal é CO_CURSO. Foram definidos cinco grupos exclusivos: A, UFPA com conceito 1; B, demais ofertas da UFPA com conceito superior; C, outras IES do Pará; D, restante da Região Norte, excluindo o Pará; e E, restante do Brasil, excluindo o Norte.")
    adicionar_tabela(doc, "Tabela 1 - Grupos comparativos de Matemática", _tabela_grupos(cursos), "elaboração própria com dados do Inep (2025).")
    _h(doc, "3.3 Indicadores e tratamento", 2)
    _p(doc, "Foram calculadas medidas de tendência central, dispersão, percentis, taxas de presença, proporções válidas e diferenças em pontos percentuais. Os códigos especiais foram tratados variável a variável conforme o dicionário. Percentuais utilizam denominadores válidos explicitamente registrados.")
    _h(doc, "3.4 Benchmarks e sensibilidade", 2)
    _p(doc, "O benchmark amplo utiliza todos os cursos válidos no recorte territorial. O benchmark comparável restringe modalidade, categoria administrativa, organização acadêmica e porte. Foram testadas faixas de porte de ±25%, ±50% e razão de até duas vezes o número de participantes.")

    _h(doc, "4 PANORAMA DA LICENCIATURA EM MATEMÁTICA", 1)
    _p(doc, f"A base contém {cursos['CO_CURSO'].nunique()} cursos de Matemática. A UFPA possui oito ofertas localizadas: sete com Conceito Enade 1 e uma com conceito superior.")
    adicionar_tabela(doc, "Tabela 2 - Ofertas de Matemática da UFPA no Enade 2025", _tabela_ofertas(cursos), "planilha de Conceito Enade e microdados do Inep (2025).")
    adicionar_figura(doc, figdir / "01_painel_ofertas_ufpa.png", "Figura 1 - Painel das ofertas de Matemática da UFPA", "elaboração própria com dados do Inep (2025).")
    _p(doc, "A distribuição interna evidencia que o Conceito 1 não se restringe a uma única modalidade ou município. Há uma oferta EaD e seis ofertas presenciais com conceito 1, o que recomenda evitar explicações antecipadas baseadas exclusivamente na modalidade ou na localização interiorana.")

    _h(doc, "5 RESULTADOS", 1)
    _h(doc, "5.1 Desempenho e participação", 2)
    adicionar_tabela(doc, "Tabela 3 - Síntese de NT_GER por grupo comparativo", _tabela_desempenho(analitica), "elaboração própria com microdados do Inep (2025).")
    adicionar_figura(doc, figdir / "02_posicao_relativa_nt_ger.png", "Figura 2 - Posição relativa dos cursos de Matemática em NT_GER", "elaboração própria com microdados do Inep (2025).")
    adicionar_figura(doc, figdir / "03_boxplot_nt_ger_grupos.png", "Figura 3 - Distribuição de NT_GER por grupo comparativo", "elaboração própria com microdados do Inep (2025).")
    adicionar_figura(doc, figdir / "validada_04_participacao.png", "Figura 4 - Participantes oficiais e notas válidas nas ofertas da UFPA", "elaboração própria com dados do Inep (2025).")
    _p(doc, "As ofertas da UFPA com conceito 1 apresentam média de NT_GER inferior à oferta presencial de Belém e ao restante do Brasil. A diferença não se limita à média: a posição relativa e as distribuições indicam deslocamento global do desempenho. O número oficial de participantes coincide com o número de notas gerais válidas nas oito ofertas, embora o total de registros cadastrais seja maior.")

    _h(doc, "5.2 Perfil demográfico e socioeconômico", 2)
    adicionar_figura(doc, figdir / "05_indicadores_socioeconomicos.png", "Figura 5 - Indicadores socioeconômicos agregados por grupo", "elaboração própria com microdados do Inep (2025).")
    _p(doc, "Os indicadores de renda, trabalho, ação afirmativa e auxílio revelam heterogeneidade entre as ofertas da UFPA e os benchmarks. Essas diferenças caracterizam a composição dos cursos; não constituem evidência de que determinado perfil individual produziu menor desempenho.")

    _h(doc, "5.3 Trajetória e condições acadêmicas", 2)
    _p(doc, "A trajetória foi descrita por ano de conclusão do ensino médio, ano de ingresso, tempo desde o ingresso e turno. As diferenças devem ser interpretadas em conjunto com porte e modalidade. Cursos com trajetórias mais longas ou maior presença do turno noturno podem enfrentar condições formativas distintas, mas o desenho exploratório não isola efeitos causais.")

    _h(doc, "5.4 Processo formativo", 2)
    adicionar_figura(doc, figdir / "06_heatmap_processo_formativo.png", "Figura 6 - Concordância nos itens do processo formativo", "elaboração própria com microdados do Inep (2025).")
    adicionar_tabela(doc, "Tabela 4 - Diagnóstico preliminar das dimensões do processo formativo", dims[["dimensao", "n_itens", "n_casos_completos", "alpha_cronbach", "decisao"]].rename(columns={"dimensao":"Dimensão", "n_itens":"Itens", "n_casos_completos":"Casos completos", "alpha_cronbach":"Alfa", "decisao":"Decisão"}), "elaboração própria com microdados do Inep (2025).")
    _p(doc, "Os coeficientes de consistência interna são elevados, mas não bastam para validar dimensões. O agrupamento permanece preliminar e depende de conferência textual item a item. O relatório não utiliza uma média única de QE_I20 a QE_I66.")

    _h(doc, "5.5 Recomendação", 2)
    adicionar_figura(doc, figdir / "07_recomendacao.png", "Figura 7 - Recomendação do curso e da instituição", "elaboração própria com microdados do Inep (2025).")
    _p(doc, "QE_I68 e QE_I69 são apresentados com seus rótulos oficiais de recomendação. Diferenças nessas respostas não devem ser automaticamente descritas como satisfação geral, pois os itens possuem conteúdo específico.")

    _h(doc, "5.6 Benchmark comparável", 2)
    adicionar_figura(doc, figdir / "validada_08_benchmark_comparavel.png", "Figura 8 - Diferenças de NT_GER em relação aos benchmarks comparáveis", "elaboração própria com microdados do Inep (2025).")
    sens_tab = sens[["ROTULO_ALVO", "criterio", "n_comparaveis", "nt_ger_alvo", "media_benchmark", "diferenca_media"]].head(21).rename(columns={"ROTULO_ALVO":"Oferta", "criterio":"Critério", "n_comparaveis":"N comparáveis", "nt_ger_alvo":"NT_GER", "media_benchmark":"Média benchmark", "diferenca_media":"Diferença"})
    adicionar_tabela(doc, "Tabela 5 - Sensibilidade dos benchmarks comparáveis", sens_tab, "elaboração própria com microdados do Inep (2025).")
    _p(doc, "Todas as ofertas conceito 1 permaneceram abaixo da média de seus benchmarks nos três critérios. A magnitude variou, sobretudo quando o número de comparáveis foi pequeno. Abaetetuba, por exemplo, teve apenas um comparável no critério mais restritivo, exigindo cautela.")

    _h(doc, "5.7 Associações ecológicas exploratórias", 2)
    _p(doc, "As associações ecológicas entre indicadores agregados foram tratadas como diagnóstico complementar. O pequeno número de ofertas da UFPA impede uma matriz restrita ao grupo conceito 1. Análises futuras deverão utilizar o conjunto regional ou nacional, relatar N de cursos, examinar outliers e comparar estimativas ponderadas e não ponderadas.")

    _h(doc, "6 DISCUSSÃO", 1)
    _p(doc, "O achado mais estável é a diferença de desempenho entre as sete ofertas conceito 1 e a oferta presencial de Belém, bem como em relação aos benchmarks comparáveis. A persistência do sinal sob critérios alternativos fortalece a interpretação descritiva de que o resultado não decorre apenas da composição indiscriminada do benchmark.")
    _p(doc, "Não foi identificado um único atributo suficiente para explicar o padrão. A presença simultânea de ofertas presenciais e EaD entre os conceitos 1, além da heterogeneidade de porte e perfil, contraria explicações monocausais. Hipóteses plausíveis envolvem combinações de participação, composição discente, condições acadêmicas, oportunidades formativas e organização do curso, que precisam ser avaliadas com fontes institucionais complementares.")
    adicionar_tabela(doc, "Tabela 6 - Síntese interpretativa dos achados", _tabela_sintese(analitica), "elaboração própria com microdados do Inep (2025).")
    _p(doc, "O contraste interno é informativo, mas a oferta de Belém não constitui contrafactual causal. Ela difere das demais em localização, escala e possivelmente em organização acadêmica e perfil de estudantes. O benchmark comparável reduz algumas diferenças, sem eliminar variáveis não observadas.")

    _h(doc, "7 CONCLUSÃO", 1)
    _p(doc, "As ofertas de Matemática da UFPA com Conceito Enade 1 diferenciam-se principalmente pelo desempenho inferior em relação à oferta interna com conceito 3 e aos benchmarks comparáveis. A direção do contraste é robusta a diferentes regras de porte, embora a magnitude varie. Indicadores socioeconômicos, de trajetória e de processo formativo revelam heterogeneidade e oferecem hipóteses, mas não permitem atribuir causalidade individual ou institucional.")
    _p(doc, "Recomenda-se aprofundar: (a) comparação entre ofertas presenciais de universidades federais; (b) análise de participação e estabilidade em cursos pequenos; (c) validação teórica dos itens de processo formativo; (d) associações ecológicas nacionais com análise de outliers; e (e) incorporação de evidências institucionais sobre currículo, estágio, corpo docente e apoio estudantil.")

    _h(doc, "REFERÊNCIAS", 1)
    adicionar_referencias(doc)

    _h(doc, "APÊNDICE A - REGRAS DE INTEGRIDADE", 1)
    _p(doc, "Os arquivos temáticos foram tratados separadamente. Não foi utilizada posição de linha, não foram reconstruídos indivíduos e todas as junções analíticas foram realizadas após agregação por CO_CURSO.")
    _h(doc, "APÊNDICE B - CRITÉRIOS DE BENCHMARK", 1)
    _p(doc, "O benchmark comparável utilizou área, modalidade, categoria administrativa, organização acadêmica e faixas de porte. Os resultados foram recalculados sob ±25%, ±50% e razão de até duas vezes o número de participantes.")

    saida_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(saida_docx)

    # Markdown auditável
    md = ["# Relatório de Matemática - Enade 2025", "", "Documento integral gerado por `executar_sprint_03.py`.", "", "## Estrutura", "", "1. Introdução", "2. Referencial institucional e metodológico", "3. Metodologia", "4. Panorama", "5. Resultados", "6. Discussão", "7. Conclusão", "Referências", "Apêndices", "", f"Cursos de Matemática: {cursos['CO_CURSO'].nunique()}.", "Ofertas UFPA: 8; conceito 1: 7."]
    saida_md.write_text("\n".join(md), encoding="utf-8")
