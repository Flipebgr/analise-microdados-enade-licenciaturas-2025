from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from src.geografia import GEOGRAFIA
from src.geografia.rotulos_questionario import ROTULOS_QE, rotulo_item
from src.relatorios.conversao_pdf import converter_docx_para_pdf
from src.relatorios.figuras_relatorio import adicionar_figura
from src.relatorios.formatacao_abnt import (
    configurar_cabecalho_rodape,
    configurar_documento,
)
from src.relatorios.referencias import adicionar_referencias
from src.relatorios.resultado_relatorio import ResultadoRelatorio
from src.relatorios.tabelas_relatorio import adicionar_tabela

FONTE_DADOS = (
    "Elaboração própria com base nos microdados do Enade das Licenciaturas 2025 "
    "e na planilha de Conceito Enade."
)
FONTE_QUESTIONARIO = (
    "Inep, Dicionário de Variáveis e Questionário do Estudante – "
    "Enade das Licenciaturas 2025."
)
RECORTES_UFPA = ("UFPA — Conceito 3", "UFPA — Conceito 4")


def _ler_csv(caminho: Path) -> pd.DataFrame:
    if not caminho.exists():
        raise FileNotFoundError(
            f"Produto analítico ausente: {caminho}. "
            "Execute as Sprints 19 e 20 antes da Sprint 21."
        )
    return pd.read_csv(caminho, low_memory=False)


def carregar_produtos(base_projeto: Path) -> dict[str, pd.DataFrame]:
    pasta = base_projeto / "dados_processados" / "geografia"
    arquivos = {
        "cursos": "cursos_geografia.csv",
        "tabela_mestra": "tabela_mestra_ufpa.csv",
        "auditoria_fontes": "auditoria_fontes_ufpa.csv",
        "base": "base_analitica_cursos.csv",
        "comparacoes": "comparacoes_regionais_validadas_sprint20.csv",
        "sensibilidade": "sensibilidade_benchmarks_sprint20.csv",
        "membros_benchmark": "membros_benchmarks_sprint20.csv",
        "contraste": "contraste_interno_ufpa_validado_sprint20.csv",
        "perfil": "perfil_recortes_validado_sprint20.csv",
        "processo": "processo_formativo_grupos_validado_sprint20.csv",
        "recomendacao": "recomendacao_recortes_validada_sprint20.csv",
        "associacoes": "associacoes_ecologicas_sprint20.csv",
        "outliers": "diagnostico_outliers_sprint20.csv",
        "auditoria": "auditoria_desempenho_sprint20.csv",
        "recomendacao_dist": "distribuicao_recomendacao.csv",
        "processo_bruto": "itens_processo_formativo.csv",
    }
    return {chave: _ler_csv(pasta / nome) for chave, nome in arquivos.items()}


def _num(valor) -> float:
    return pd.to_numeric(pd.Series([valor]), errors="coerce").iloc[0]


def _fmt(valor, casas: int = 2) -> str:
    valor = _num(valor)
    if pd.isna(valor):
        return "não disponível"
    return f"{valor:.{casas}f}".replace(".", ",")


def _fmt_pct_proporcao(valor, casas: int = 1) -> str:
    valor = _num(valor)
    if pd.isna(valor):
        return "não disponível"
    return f"{100 * valor:.{casas}f}%".replace(".", ",")


def tabela_ofertas_ufpa(base: pd.DataFrame) -> pd.DataFrame:
    colunas = [
        "CO_CURSO",
        "ROTULO_OFERTA",
        "CONCEITO_ENADE_NUM",
        "INSCRITOS_NUM",
        "PARTICIPANTES_NUM",
        "TAXA_PARTICIPACAO_OFICIAL",
        "taxa_presenca_microdados",
        "nt_ger_mean",
        "nt_obj_mean",
        "nt_dis_mean",
        "nt_ger_percentil_brasil",
        "nt_ger_percentil_norte",
        "nt_ger_percentil_para",
    ]
    presentes = [c for c in colunas if c in base.columns]
    return (
        base.loc[base["CO_IES"].eq(GEOGRAFIA.co_ies_focal), presentes]
        .sort_values("ROTULO_OFERTA")
        .reset_index(drop=True)
    )


def tabela_regional_nt_ger(comparacoes: pd.DataFrame) -> pd.DataFrame:
    trabalho = comparacoes.loc[comparacoes["INDICADOR"].eq("nt_ger_mean")].copy()
    colunas = [
        "RECORTE",
        "N_CURSOS",
        "N_PARTICIPANTES",
        "MEDIA_CURSOS",
        "MEDIA_PONDERADA_PARTICIPANTES",
        "MEDIANA_CURSOS",
        "DP_CURSOS",
        "P25",
        "P75",
        "AMPLITUDE_IQR",
    ]
    presentes = [c for c in colunas if c in trabalho.columns]
    return trabalho[presentes].reset_index(drop=True)


def tabela_benchmark_principal(sensibilidade: pd.DataFrame) -> pd.DataFrame:
    principal = sensibilidade.loc[
        sensibilidade["CENARIO"].eq("estrutura_porte_0_5_2_0")
    ].copy()
    colunas = [
        "CO_CURSO_ALVO",
        "ROTULO_ALVO",
        "CONCEITO_ALVO",
        "N_CURSOS",
        "nt_ger_mean_ALVO",
        "nt_ger_mean_MEDIA_BENCHMARK",
        "nt_ger_mean_DIFERENCA",
        "nt_ger_mean_Z",
        "nt_obj_mean_DIFERENCA",
        "nt_dis_mean_DIFERENCA",
        "taxa_presenca_microdados_DIFERENCA",
    ]
    presentes = [c for c in colunas if c in principal.columns]
    return principal[presentes].sort_values("ROTULO_ALVO").reset_index(drop=True)


def tabela_perfil_interno(perfil: pd.DataFrame) -> pd.DataFrame:
    indicadores = [
        "sexo_feminino_pct",
        "idade_media",
        "mae_superior_pct",
        "pai_superior_pct",
        "renda_ate_3sm_pct",
        "trabalha_pct",
        "acao_afirmativa_pct",
        "auxilio_permanencia_pct",
        "bolsa_academica_pct",
        "estudo_4h_ou_mais_pct",
        "turno_noturno_pct",
        "anos_desde_ingresso_media",
        "qe_i68_media",
        "qe_i69_media",
        "qe_i70_interesse_pct",
    ]
    trabalho = perfil.loc[
        perfil["RECORTE_GEOGRAFIA"].isin(RECORTES_UFPA)
        & perfil["INDICADOR"].isin(indicadores)
    ].copy()
    colunas = [
        "RECORTE_GEOGRAFIA",
        "INDICADOR",
        "N_CURSOS",
        "MEDIA_CURSOS",
        "MEDIANA_CURSOS",
        "DP_CURSOS",
    ]
    return trabalho[colunas].reset_index(drop=True)


def tabela_processo_interno(processo: pd.DataFrame, n: int = 12) -> pd.DataFrame:
    trabalho = processo.copy()
    if trabalho.empty:
        return trabalho
    trabalho["ABS_DIF"] = pd.to_numeric(
        trabalho["DIFERENCA_C3_C4"], errors="coerce"
    ).abs()
    trabalho["ROTULO_OFICIAL"] = trabalho["ITEM"].map(rotulo_item)
    colunas = [
        "ITEM",
        "ROTULO_OFICIAL",
        "N_CURSOS_CONCEITO_3",
        "MEDIA_CONCEITO_3",
        "N_CURSOS_CONCEITO_4",
        "MEDIA_CONCEITO_4",
        "DIFERENCA_C3_C4",
        "MEDIA_OUTRAS_IES_PARA",
        "MEDIA_NORTE_SEM_PARA",
        "MEDIA_BRASIL_SEM_NORTE",
    ]
    return (
        trabalho.sort_values("ABS_DIF", ascending=False)
        .head(n)[colunas]
        .reset_index(drop=True)
    )


def resumo_panorama(base: pd.DataFrame) -> str:
    ufpa = tabela_ofertas_ufpa(base)
    conceitos = sorted(
        pd.to_numeric(ufpa["CONCEITO_ENADE_NUM"], errors="coerce")
        .dropna()
        .astype(int)
        .unique()
        .tolist()
    )
    return (
        f"O universo analítico reúne {len(base)} cursos de Geografia identificados por CO_CURSO. "
        f"Foram localizadas {len(ufpa)} ofertas da UFPA, com Conceitos Enade {conceitos}. "
        "Não existe oferta da UFPA com Conceito Enade 1 em Geografia; o Grupo A permanece vazio "
        "e não é reconstruído por aproximação. O contraste interno principal compara duas ofertas "
        "Conceito 3 e duas ofertas Conceito 4, sem tratar Conceito 3 como insuficiência e sem "
        "atribuir causalidade ao conceito."
    )


def resumo_desempenho_ufpa(base: pd.DataFrame, contraste: pd.DataFrame) -> str:
    ufpa = tabela_ofertas_ufpa(base)
    partes = []
    for linha in ufpa.itertuples():
        partes.append(
            f"{linha.ROTULO_OFERTA} (Conceito {int(linha.CONCEITO_ENADE_NUM)}): "
            f"NT_GER {_fmt(linha.nt_ger_mean)}, NT_OBJ {_fmt(linha.nt_obj_mean)} "
            f"e NT_DIS {_fmt(linha.nt_dis_mean)}"
        )
    texto = "; ".join(partes) + "."
    linha = contraste.loc[contraste["INDICADOR"].eq("nt_ger_mean")]
    if not linha.empty:
        r = linha.iloc[0]
        texto += (
            " Na média das ofertas, Conceito 3 menos Conceito 4 em NT_GER corresponde a "
            f"{_fmt(r.get('DIFERENCA_C3_C4'))} ponto(s), com tamanho padronizado "
            f"puramente descritivo de {_fmt(r.get('D_PADRONIZADO_DESCRITIVO'))}."
        )
    texto += (
        " O contraste envolve somente duas ofertas em cada estrato e não sustenta inferência causal."
    )
    return texto


def resumo_participacao(base: pd.DataFrame) -> str:
    ufpa = tabela_ofertas_ufpa(base)
    partes = []
    for linha in ufpa.itertuples():
        partes.append(
            f"{linha.ROTULO_OFERTA}: {int(linha.PARTICIPANTES_NUM)} participantes de "
            f"{int(linha.INSCRITOS_NUM)} inscritos "
            f"({_fmt_pct_proporcao(linha.TAXA_PARTICIPACAO_OFICIAL)})"
        )
    return "; ".join(partes) + "."


def resumo_regional(comparacoes: pd.DataFrame) -> str:
    tab = tabela_regional_nt_ger(comparacoes)
    recortes = ["UFPA agregada", "Região Norte sem UFPA", "Brasil geral"]
    partes: list[str] = []
    for recorte in recortes:
        sub = tab.loc[tab["RECORTE"].eq(recorte)]
        if sub.empty:
            continue
        linha = sub.iloc[0]
        partes.append(
            f"{recorte}: média simples {_fmt(linha.get('MEDIA_CURSOS'))} e média "
            f"ponderada por participantes {_fmt(linha.get('MEDIA_PONDERADA_PARTICIPANTES'))}"
        )
    if not partes:
        return "Os recortes de síntese regional não foram localizados."
    return "; ".join(partes) + "."


def resumo_benchmarks(sensibilidade: pd.DataFrame) -> str:
    principal = tabela_benchmark_principal(sensibilidade)
    if principal.empty:
        return "O benchmark estrutural principal não estava disponível."
    dif = pd.to_numeric(principal["nt_ger_mean_DIFERENCA"], errors="coerce")
    acima = int((dif > 0).sum())
    abaixo = int((dif < 0).sum())
    menor = principal.loc[dif.idxmin()]
    maior = principal.loc[dif.idxmax()]
    return (
        "No cenário estrutural principal — mesma modalidade, categoria administrativa, "
        "organização acadêmica e porte entre 0,5x e 2x participantes — "
        f"{acima} oferta(s) da UFPA apresentam NT_GER médio acima do benchmark e {abaixo} abaixo. "
        f"A menor diferença é observada em {menor['ROTULO_ALVO']} "
        f"({_fmt(menor['nt_ger_mean_DIFERENCA'])} ponto(s)) e a maior diferença no sentido oposto "
        f"em {maior['ROTULO_ALVO']} ({_fmt(maior['nt_ger_mean_DIFERENCA'])} ponto(s)). "
        "Foram calculados cinco cenários para cada uma das quatro ofertas, totalizando "
        "20 combinações oferta-cenário."
    )


def resumo_associacoes(associacoes: pd.DataFrame) -> str:
    validas = associacoes.dropna(subset=["SPEARMAN_RHO"]).copy()
    if validas.empty:
        return "Não houve associações ecológicas calculáveis."
    linha = validas.loc[validas["SPEARMAN_RHO"].abs().idxmax()]
    return (
        "Entre as associações ecológicas examinadas, a maior magnitude absoluta foi entre "
        f"{linha['INDICADOR_X']} e NT_GER médio "
        f"(Spearman rho={linha['SPEARMAN_RHO']:.3f}; N={int(linha['N_CURSOS'])} cursos). "
        "O coeficiente usa cursos como unidades, não estudantes; portanto, não representa "
        "associação individual e não sustenta inferência causal."
    )


def resumo_processo(processo: pd.DataFrame) -> str:
    tabela = tabela_processo_interno(processo, n=6)
    if tabela.empty:
        return "Não houve comparação item a item disponível para o processo formativo."
    neg = tabela.sort_values("DIFERENCA_C3_C4").head(2)
    pos = tabela.sort_values("DIFERENCA_C3_C4", ascending=False).head(2)
    texto_neg = "; ".join(
        f"{r.ITEM} ({r.ROTULO_OFICIAL}): {r.DIFERENCA_C3_C4:+.2f}"
        for r in neg.itertuples()
    )
    texto_pos = "; ".join(
        f"{r.ITEM} ({r.ROTULO_OFICIAL}): {r.DIFERENCA_C3_C4:+.2f}"
        for r in pos.itertuples()
    )
    return (
        "No contraste Conceito 3 menos Conceito 4, entre os itens com maiores diferenças "
        f"absolutas, aparecem diferenças negativas em {texto_neg} e positivas em {texto_pos}. "
        "Os textos oficiais do Inep são usados para interpretação, mas QE_I20–QE_I66 permanecem "
        "item a item; não é criado índice único sem validação teórica adicional."
    )


def resumo_recomendacao(recomendacao: pd.DataFrame) -> str:
    partes = []
    for indicador in ("qe_i68_media", "qe_i69_media", "qe_i70_interesse_pct"):
        for recorte in RECORTES_UFPA:
            sub = recomendacao.loc[
                recomendacao["RECORTE_GEOGRAFIA"].eq(recorte)
                & recomendacao["INDICADOR"].eq(indicador)
            ]
            if sub.empty:
                continue
            valor = sub.iloc[0]["MEDIA_CURSOS"]
            if indicador.endswith("_pct"):
                fmt = _fmt_pct_proporcao(valor)
            else:
                fmt = _fmt(valor)
            partes.append(f"{recorte}, {indicador}: {fmt}")
    return (
        "; ".join(partes)
        + ". QE_I68, QE_I69 e QE_I70 permanecem conceitualmente separados; "
        "não são reunidos automaticamente sob o termo satisfação."
    )


def _p(doc: Document, texto: str, estilo: str | None = None):
    par = doc.add_paragraph(style=estilo)
    par.add_run(texto)
    return par


def _capa(doc: Document) -> None:
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = 0
    r = p.add_run("UNIVERSIDADE FEDERAL DO PARÁ")
    r.bold = True
    r.font.size = Pt(12)
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = 0
    r = p.add_run(
        "GEOGRAFIA NO ENADE 2025: DESEMPENHO, PERFIL, PROCESSO FORMATIVO "
        "E CONTRASTES DAS OFERTAS DA UFPA"
    )
    r.bold = True
    r.font.size = Pt(14)
    for _ in range(7):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = 0
    p.add_run("Autor(a): [preencher]")
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = 0
    p.add_run("Belém\n2026")


def _sumario(doc: Document) -> None:
    doc.add_heading("SUMÁRIO", level=1)
    itens = [
        "1 INTRODUÇÃO",
        "2 REFERENCIAL INSTITUCIONAL E METODOLÓGICO",
        "3 METODOLOGIA",
        "4 PANORAMA DA LICENCIATURA EM GEOGRAFIA",
        "5 RESULTADOS",
        "5.1 Desempenho",
        "5.2 Perfil demográfico e socioeconômico",
        "5.3 Trajetória e condições acadêmicas",
        "5.4 Processo formativo",
        "5.5 Recomendação",
        "5.6 Benchmark comparável",
        "5.7 Associações ecológicas",
        "5.8 Comparações regionais e nacionais",
        "5.9 Contraste interno das ofertas da UFPA",
        "6 DISCUSSÃO",
        "7 CONCLUSÃO",
        "REFERÊNCIAS",
        "APÊNDICES",
    ]
    for item in itens:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = 0
        p.add_run(item)


def _figura(
    doc: Document,
    base_projeto: Path,
    nome: str,
    numero: int,
    titulo: str,
    leitura: str,
    hipotese: str,
    limitacao: str,
) -> None:
    adicionar_figura(
        doc,
        base_projeto / "figuras" / "geografia" / nome,
        f"Figura {numero} – {titulo}",
        FONTE_DADOS,
    )
    _p(doc, f"Descrição e interpretação: {leitura}")
    _p(doc, f"Hipótese analítica: {hipotese}")
    _p(doc, f"Limitação: {limitacao}")


def _markdown(dados: dict[str, pd.DataFrame]) -> str:
    base = dados["base"]
    ofertas = tabela_ofertas_ufpa(base)
    regional = tabela_regional_nt_ger(dados["comparacoes"])
    benchmark = tabela_benchmark_principal(dados["sensibilidade"])
    perfil = tabela_perfil_interno(dados["perfil"])
    processo = tabela_processo_interno(dados["processo"])

    linhas = [
        "# GEOGRAFIA NO ENADE 2025: DESEMPENHO, PERFIL, PROCESSO FORMATIVO E CONTRASTES DAS OFERTAS DA UFPA",
        "",
        "## RESUMO",
        "",
        f"Este relatório técnico-científico analisa {len(base)} cursos de Geografia presentes nos "
        f"microdados do Enade das Licenciaturas 2025, com {len(ofertas)} ofertas localizadas da UFPA. "
        "A unidade principal é CO_CURSO. Os arquivos temáticos foram processados separadamente, tratados "
        "quanto a ausências, agregados por curso e somente então relacionados por junções one-to-one. "
        "Não existe oferta da UFPA com Conceito Enade 1 em Geografia; o Grupo A permanece vazio. "
        "O contraste interno examina duas ofertas Conceito 3 e duas Conceito 4, sem tratar Conceito 3 "
        "como insuficiência. São analisados desempenho, participação, perfil discente, trajetória, "
        "processo formativo, recomendação, benchmarks estruturais, comparações regionais e associações "
        "ecológicas. As interpretações são descritivas e não causais.",
        "",
        "**Palavras-chave:** Enade; Geografia; UFPA; formação de professores; microdados; benchmark.",
        "",
        "## ABSTRACT",
        "",
        "This technical-scientific report analyzes Geography teacher education programs in the 2025 "
        "Enade, focusing on four UFPA offers. CO_CURSO is the main unit of analysis. Thematic files "
        "were processed separately, aggregated at course level and only then combined one-to-one. "
        "UFPA has no Geography offer with Enade Concept 1. The internal comparison contrasts two "
        "Concept-3 and two Concept-4 offers descriptively. Performance, participation, student profile, "
        "academic trajectory, formative process, recommendation, structural benchmarks, regional "
        "comparisons and ecological associations are examined. Results are descriptive and non-causal.",
        "",
        "**Keywords:** Enade; Geography; UFPA; teacher education; microdata; benchmark.",
        "",
        "## SUMÁRIO",
        "",
        "1 INTRODUÇÃO; 2 REFERENCIAL INSTITUCIONAL E METODOLÓGICO; 3 METODOLOGIA; "
        "4 PANORAMA DA LICENCIATURA EM GEOGRAFIA; 5 RESULTADOS; 6 DISCUSSÃO; "
        "7 CONCLUSÃO; REFERÊNCIAS; APÊNDICES.",
        "",
        "# 1 INTRODUÇÃO",
        "",
        "A pergunta central é: quais características de desempenho, composição discente, trajetória "
        "acadêmica e avaliação do processo formativo diferenciam as ofertas da UFPA em Geografia entre "
        "si e em relação às demais ofertas da mesma área na UFPA, no Pará, na Região Norte e no Brasil?",
        "",
        "# 2 REFERENCIAL INSTITUCIONAL E METODOLÓGICO",
        "",
        "O Conceito Enade é tratado como classificação externa do curso. Ausência de conceito não é "
        "recodificada. Em Geografia, não existe oferta UFPA com Conceito Enade 1. O contraste interno "
        "Conceito 3 versus Conceito 4 é descritivo e não representa um desenho causal nem transforma "
        "Conceito 3 em categoria de insuficiência.",
        "",
        "# 3 METODOLOGIA",
        "",
        "A unidade principal de análise é CO_CURSO. Não se usa posição de linha como chave, não se cria "
        "identificador artificial e não são realizadas junções individuais entre arquivos temáticos. "
        "O fluxo é arquivo temático → tratamento de ausentes → agregação por CO_CURSO → uma linha por "
        "curso → junções one-to-one → comparação entre cursos. Relações entre temas distintos são "
        "ecológicas e recebem ressalva explícita de falácia ecológica.",
        "",
        "Os recortes exclusivos usados no contraste descritivo são UFPA Conceito 3, UFPA Conceito 4, "
        "outras IES do Pará, Norte sem Pará e Brasil sem Norte. Pará, Norte e Brasil completos aparecem "
        "somente como benchmarks descritivos quando se sobrepõem. Para cada oferta UFPA foram calculados "
        "cinco cenários de benchmark, totalizando 20 combinações oferta-cenário.",
        "",
        "# 4 PANORAMA DA LICENCIATURA EM GEOGRAFIA",
        "",
        resumo_panorama(base),
        "",
        ofertas.to_markdown(index=False, floatfmt=".3f"),
        "",
        "# 5 RESULTADOS",
        "",
        "## 5.1 Desempenho",
        "",
        resumo_desempenho_ufpa(base, dados["contraste"]),
        "",
        "## 5.2 Perfil demográfico e socioeconômico",
        "",
        perfil.to_markdown(index=False, floatfmt=".3f"),
        "",
        "## 5.3 Trajetória e condições acadêmicas",
        "",
        "Turno, tempo desde o ingresso, trabalho, bolsas, auxílios e dedicação aos estudos são "
        "interpretados como características agregadas das ofertas. Relações com desempenho que envolvem "
        "arquivos distintos são examinadas somente no nível ecológico do curso.",
        "",
        "## 5.4 Processo formativo",
        "",
        resumo_processo(dados["processo"]),
        "",
        processo.to_markdown(index=False, floatfmt=".3f"),
        "",
        "## 5.5 Recomendação",
        "",
        resumo_recomendacao(dados["recomendacao"]),
        "",
        f"{rotulo_item('QE_I68')} {rotulo_item('QE_I69')} {rotulo_item('QE_I70')}",
        "",
        "## 5.6 Benchmark comparável",
        "",
        resumo_benchmarks(dados["sensibilidade"]),
        "",
        benchmark.to_markdown(index=False, floatfmt=".3f"),
        "",
        "## 5.7 Associações ecológicas",
        "",
        resumo_associacoes(dados["associacoes"]),
        "",
        dados["associacoes"].to_markdown(index=False, floatfmt=".4f"),
        "",
        "## 5.8 Comparações regionais e nacionais",
        "",
        resumo_regional(dados["comparacoes"]),
        "",
        regional.to_markdown(index=False, floatfmt=".3f"),
        "",
        "## 5.9 Contraste interno das ofertas da UFPA",
        "",
        "### 5.9.1 Participação e desempenho",
        "",
        resumo_participacao(base),
        "",
        resumo_desempenho_ufpa(base, dados["contraste"]),
        "",
        "### 5.9.2 Perfil discente",
        "",
        "As ofertas Conceito 3 e Conceito 4 são comparadas como conjuntos de dois cursos cada. "
        "As diferenças descrevem composição agregada e não relações individuais entre perfil e nota.",
        "",
        "### 5.9.3 Trajetória acadêmica",
        "",
        "Os indicadores de trajetória permanecem agregados por curso e são usados para identificar "
        "padrões e hipóteses, sem atribuição causal.",
        "",
        "### 5.9.4 Processo formativo",
        "",
        resumo_processo(dados["processo"]),
        "",
        "### 5.9.5 Recomendação",
        "",
        resumo_recomendacao(dados["recomendacao"]),
        "",
        "### 5.9.6 Benchmark comparável",
        "",
        resumo_benchmarks(dados["sensibilidade"]),
        "",
        "### 5.9.7 Síntese dos pontos distintivos",
        "",
        "No conjunto da UFPA, as ofertas Conceito 3 apresentam NT_GER e NT_OBJ médios inferiores às "
        "ofertas Conceito 4, enquanto NT_DIS é praticamente equivalente. Ao mesmo tempo, algumas "
        "características socioeconômicas e de recomendação seguem direção distinta, reforçando a "
        "necessidade de uma leitura multidimensional e não causal.",
        "",
        "# 6 DISCUSSÃO",
        "",
        "Os resultados indicam heterogeneidade interna relevante. O contraste médio em NT_GER se concentra "
        "sobretudo no componente objetivo, pois a diferença de NT_DIS entre os dois estratos é mínima. "
        "Isso não permite afirmar que qualquer característica de modalidade, localização, composição "
        "discente ou processo formativo cause o desempenho. O N institucional é de somente duas ofertas "
        "por estrato, e as comparações externas mostram que as quatro ofertas da UFPA ficam abaixo do "
        "benchmark estrutural principal em NT_GER, com magnitudes distintas.",
        "",
        "As associações ecológicas nacionais/regionalmente ampliadas servem para gerar hipóteses. Mesmo "
        "quando Spearman apresenta magnitude relevante, a unidade é o curso e permanece a possibilidade "
        "de confundimento, outliers e falácia ecológica.",
        "",
        "# 7 CONCLUSÃO",
        "",
        "Geografia na UFPA não apresenta oferta Conceito Enade 1 nas fontes de 2025. Entre as quatro "
        "ofertas localizadas, Belém e Ananindeua têm Conceito 4, enquanto Altamira e Cametá têm Conceito 3. "
        "O contraste descritivo aponta menor NT_GER e NT_OBJ médios no estrato Conceito 3, sem diferença "
        "material em NT_DIS. As quatro ofertas devem ser interpretadas individualmente diante dos seus "
        "Ns, perfis e benchmarks. O relatório não sustenta causalidade e não converte Conceito 3 em "
        "sinônimo de insuficiência.",
        "",
        "# REFERÊNCIAS",
        "",
        "As referências bibliográficas são as mesmas cadastradas no módulo compartilhado do projeto e "
        "correspondem às fontes institucionais e metodológicas efetivamente utilizadas.",
        "",
        "# APÊNDICE A – REGRAS DE INTEGRIDADE",
        "",
        "- CO_CURSO é a unidade principal.",
        "- Arquivos temáticos são agregados antes das junções.",
        "- Não há join individual entre desempenho e questionário.",
        "- Não há reconstrução de estudante entre arquivos.",
        "- Ausência de conceito não é Conceito 1.",
        "- O Grupo A permanece vazio em Geografia.",
        "- Conceito 3 não é tratado como insuficiência.",
        "- Relações entre temas diferentes são ecológicas.",
        "- Pará, Norte e Brasil completos não são grupos independentes em testes.",
        "",
        "# APÊNDICE B – APROFUNDAMENTOS SUGERIDOS",
        "",
        "1. Sensibilidade de porte — pergunta: os déficits frente aos benchmarks persistem sob faixas de "
        "porte mais estreitas? Variáveis: NT_GER, NT_OBJ, NT_DIS, participantes e atributos estruturais. "
        "Método: cenários progressivos e, quando viável, pareamento. Limitação: poucos comparáveis em "
        "alguns estratos.",
        "2. Componente objetivo — pergunta: quais domínios ou padrões de acerto ajudam a caracterizar a "
        "diferença observada em NT_OBJ? Variáveis: QT_ACERTOS e indicadores disponíveis no mesmo arquivo "
        "de desempenho. Método: análise intrarquivo por oferta. Limitação: relações mecânicas entre nota "
        "e acertos.",
        "3. Processo formativo — pergunta: quais itens QE_I20–QE_I66 mantêm contraste Conceito 3 × 4 "
        "quando consideradas respostas válidas e referências externas? Método: análise item a item, "
        "dimensionalidade e consistência interna antes de índices. Limitação: escala ordinal e resposta "
        "autorreferida.",
        "4. Permanência e trajetória — pergunta: como renda, trabalho, auxílios, bolsas, turno e tempo "
        "desde o ingresso se organizam ecologicamente entre cursos? Método: perfis padronizados e "
        "Spearman por curso. Limitação: falácia ecológica e confundimento residual.",
        "5. Recomendação — pergunta: QE_I68 e QE_I69 se associam ecologicamente a dimensões formativas "
        "específicas no universo nacional de Geografia? Método: agregação por curso, Spearman, dispersão "
        "e análise de outliers. Limitação: não representa relação individual entre respostas.",
        "",
        "# APÊNDICE C – RÓTULOS OFICIAIS DOS ITENS QE_I20–QE_I66",
        "",
    ]
    for codigo in [f"QE_I{i}" for i in range(20, 67)]:
        linhas.append(f"- **{codigo}** — {ROTULOS_QE[codigo]}")
    linhas.extend(["", f"Fonte dos rótulos: {FONTE_QUESTIONARIO}"])
    return "\n".join(linhas)


def gerar_relatorio(
    base_projeto: Path,
    saida_docx: Path,
    saida_md: Path,
) -> dict:
    dados = carregar_produtos(base_projeto)
    base = dados["base"]
    ofertas = tabela_ofertas_ufpa(base)
    regional = tabela_regional_nt_ger(dados["comparacoes"])
    benchmark = tabela_benchmark_principal(dados["sensibilidade"])
    perfil = tabela_perfil_interno(dados["perfil"])
    processo = tabela_processo_interno(dados["processo"])

    doc = Document()
    configurar_documento(doc)
    _capa(doc)
    doc.add_page_break()

    doc.add_heading("RESUMO", level=1)
    _p(
        doc,
        f"Este relatório técnico-científico analisa {len(base)} cursos de Geografia presentes nos "
        f"microdados do Enade das Licenciaturas 2025, com {len(ofertas)} ofertas da UFPA. A unidade "
        "principal é CO_CURSO. Os arquivos temáticos foram processados separadamente, agregados por curso "
        "e somente então relacionados por junções one-to-one. Não existe oferta da UFPA com Conceito "
        "Enade 1 em Geografia; o Grupo A permanece vazio. O contraste interno examina duas ofertas "
        "Conceito 3 e duas Conceito 4, sem tratar Conceito 3 como insuficiência. São examinados desempenho, "
        "participação, perfil discente, trajetória, processo formativo, recomendação, benchmarks, "
        "comparações regionais e associações ecológicas. A interpretação é descritiva e não causal.",
        "Resumo",
    )
    _p(
        doc,
        "Palavras-chave: Enade; Geografia; UFPA; formação de professores; microdados; benchmark.",
    )

    doc.add_heading("ABSTRACT", level=1)
    _p(
        doc,
        "This technical-scientific report analyzes Geography teacher education programs in the 2025 "
        "Enade, focusing on four UFPA offers. CO_CURSO is the main unit of analysis. Thematic files were "
        "processed separately, aggregated at course level and combined one-to-one. UFPA has no Geography "
        "offer with Enade Concept 1. The internal comparison contrasts two Concept-3 and two Concept-4 "
        "offers descriptively. Results are descriptive and non-causal.",
        "Resumo",
    )
    _p(
        doc,
        "Keywords: Enade; Geography; UFPA; teacher education; microdata; benchmark.",
    )

    doc.add_page_break()
    _sumario(doc)
    doc.add_page_break()

    doc.add_heading("1 INTRODUÇÃO", level=1)
    _p(
        doc,
        "Este relatório responde à pergunta: quais características de desempenho, composição discente, "
        "trajetória acadêmica e avaliação do processo formativo diferenciam as ofertas de Geografia da "
        "UFPA entre si e em relação às demais ofertas da mesma área no Pará, na Região Norte e no Brasil?",
    )

    doc.add_heading("2 REFERENCIAL INSTITUCIONAL E METODOLÓGICO", level=1)
    _p(
        doc,
        "O Conceito Enade é tratado como classificação externa do curso e não como variável causal. "
        "Ausência de conceito permanece distinta de Conceito 1. Em Geografia não existe oferta UFPA com "
        "Conceito Enade 1. O contraste entre duas ofertas Conceito 3 e duas Conceito 4 é descritivo; "
        "Conceito 3 não é interpretado como insuficiência.",
    )

    doc.add_heading("3 METODOLOGIA", level=1)
    _p(
        doc,
        "A unidade principal é CO_CURSO. Cada arquivo temático foi tratado separadamente, com tratamento "
        "de ausências e agregação por curso antes de qualquer junção. Não foi usada posição de linha, "
        "não foi criado identificador artificial e não foram realizadas junções individuais entre temas. "
        "As integrações das tabelas agregadas são one-to-one. Relações entre arquivos distintos são "
        "ecológicas e recebem ressalva de falácia ecológica.",
    )
    _p(
        doc,
        "Os recortes exclusivos usados no contraste descritivo são UFPA Conceito 3, UFPA Conceito 4, "
        "outras IES do Pará, Norte sem Pará e Brasil sem Norte. Para cada oferta UFPA são calculados cinco "
        "cenários de benchmark. O cenário principal preserva modalidade, categoria administrativa, "
        "organização acadêmica e porte de participantes entre 0,5x e 2x. Foram avaliadas 20 combinações "
        "oferta-cenário. Médias simples e ponderadas por participantes são apresentadas nos recortes "
        "territoriais.",
    )

    doc.add_heading("4 PANORAMA DA LICENCIATURA EM GEOGRAFIA", level=1)
    _p(doc, resumo_panorama(base))
    adicionar_tabela(doc, "Tabela 1 – Ofertas de Geografia da UFPA", ofertas, FONTE_DADOS)
    _figura(
        doc,
        base_projeto,
        "01_painel_ofertas_ufpa.png",
        1,
        "Ofertas de Geografia da UFPA",
        "O painel reúne as quatro ofertas, conceitos, porte e indicadores iniciais.",
        "Diferenças de porte, participação e composição podem coexistir com a heterogeneidade observada.",
        "O painel é descritivo e não controla diferenças estruturais entre ofertas.",
    )

    doc.add_heading("5 RESULTADOS", level=1)

    doc.add_heading("5.1 Desempenho", level=2)
    _p(doc, resumo_desempenho_ufpa(base, dados["contraste"]))
    for nome, numero, titulo, leitura, hipotese, limitacao in [
        (
            "02_posicao_relativa_nt_ger.png",
            2,
            "Posição relativa das ofertas em NT_GER",
            "As ofertas são situadas na distribuição nacional da própria área.",
            "A posição relativa ajuda a contextualizar diferenças entre ofertas.",
            "Percentis não eliminam diferenças de participação, porte ou composição.",
        ),
        (
            "03_distribuicao_nt_ger.png",
            3,
            "Distribuição de NT_GER",
            "A distribuição individual é construída apenas com variáveis do arquivo de desempenho.",
            "Diferenças de centro e dispersão podem indicar perfis de desempenho distintos.",
            "Cursos com N menor apresentam maior instabilidade amostral.",
        ),
        (
            "04_distribuicao_nt_obj.png",
            4,
            "Distribuição de NT_OBJ",
            "O componente objetivo é analisado separadamente da nota geral.",
            "O contraste interno pode estar mais concentrado no componente objetivo.",
            "NT_OBJ e NT_GER possuem relação mecânica e não são evidências independentes.",
        ),
        (
            "05_distribuicao_nt_dis.png",
            5,
            "Distribuição de NT_DIS",
            "O componente discursivo é preservado em sua escala própria.",
            "A proximidade entre estratos em NT_DIS contrasta com a diferença observada em NT_OBJ.",
            "N válido e particularidades de correção devem ser considerados.",
        ),
        (
            "09_desempenho_ofertas_ufpa.png",
            6,
            "Desempenho das ofertas da UFPA",
            "NT_GER, NT_OBJ e NT_DIS são comparados entre as quatro ofertas.",
            "A heterogeneidade interna não precisa seguir estritamente a ordenação do conceito.",
            "Há somente quatro ofertas da UFPA, o que restringe inferência entre estratos.",
        ),
        (
            "10_percentis_ofertas_ufpa.png",
            7,
            "Percentis das ofertas da UFPA",
            "Os percentis posicionam cada curso na distribuição de Geografia.",
            "Cursos de conceitos distintos podem apresentar sobreposição parcial em posições relativas.",
            "Percentis dependem do universo de cursos válidos.",
        ),
    ]:
        _figura(
            doc,
            base_projeto,
            nome,
            numero,
            titulo,
            leitura,
            hipotese,
            limitacao,
        )

    doc.add_heading("5.2 Perfil demográfico e socioeconômico", level=2)
    _p(
        doc,
        "Sexo, idade, escolaridade parental, renda, trabalho, ação afirmativa, bolsas, auxílios e "
        "horas de estudo são agregados por CO_CURSO. Percentuais são calculados sobre respostas válidas; "
        "ausências permanecem documentadas nos produtos analíticos.",
    )
    adicionar_tabela(
        doc,
        "Tabela 2 – Perfil interno da UFPA por estrato de conceito",
        perfil,
        FONTE_DADOS,
    )
    _figura(
        doc,
        base_projeto,
        "06_perfil_socioeconomico.png",
        8,
        "Perfil socioeconômico das ofertas",
        "A figura compara indicadores selecionados de composição discente.",
        "Diferenças de composição podem coexistir com diferenças de desempenho sem relação causal individual.",
        "Os indicadores são agregados por curso e estão sujeitos à falácia ecológica.",
    )
    _figura(
        doc,
        base_projeto,
        "validada_18_perfil_interno_ufpa.png",
        9,
        "Perfil interno validado da UFPA",
        "O contraste resume duas ofertas Conceito 3 e duas Conceito 4.",
        "Características de composição podem orientar aprofundamentos sobre permanência e trajetória.",
        "O N de cursos por estrato é dois.",
    )

    doc.add_heading("5.3 Trajetória e condições acadêmicas", level=2)
    _p(
        doc,
        "Turno, anos desde o ingresso, trabalho, bolsas, auxílios e dedicação aos estudos são "
        "interpretados como características agregadas das ofertas. Relações com desempenho, quando "
        "envolvem arquivos distintos, são examinadas somente no nível ecológico do curso.",
    )

    doc.add_heading("5.4 Processo formativo", level=2)
    _p(doc, resumo_processo(dados["processo"]))
    adicionar_tabela(
        doc,
        "Tabela 3 – Itens QE_I20–QE_I66 de maior diferença no contraste Conceito 3 × Conceito 4",
        processo,
        FONTE_QUESTIONARIO,
    )
    _figura(
        doc,
        base_projeto,
        "07_processo_formativo.png",
        10,
        "Itens do processo formativo",
        "QE_I20–QE_I66 são examinados item a item, preservando a escala oficial.",
        "Padrões por item podem sugerir dimensões para validação teórica posterior.",
        "Não é criado índice único sem validação de escala, dimensionalidade e consistência.",
    )
    _figura(
        doc,
        base_projeto,
        "validada_17_processo_formativo.png",
        11,
        "Processo formativo no contraste interno",
        "A figura destaca itens com maiores diferenças absolutas entre os estratos.",
        "Itens persistentes em diferentes referências merecem aprofundamento institucional.",
        "A escala é ordinal e as respostas são autorreferidas.",
    )

    doc.add_heading("5.5 Recomendação", level=2)
    _p(doc, resumo_recomendacao(dados["recomendacao"]))
    _p(
        doc,
        f"{rotulo_item('QE_I68')} {rotulo_item('QE_I69')} {rotulo_item('QE_I70')} "
        "Os três indicadores são tratados de acordo com seus rótulos oficiais.",
    )
    _figura(
        doc,
        base_projeto,
        "12_recomendacao.png",
        12,
        "Recomendação e interesse",
        "A figura preserva a distinção entre QE_I68, QE_I69 e QE_I70.",
        "Padrões de recomendação podem gerar hipóteses para investigação posterior.",
        "Não se interpreta recomendação como medida geral de satisfação.",
    )

    doc.add_heading("5.6 Benchmark comparável", level=2)
    _p(doc, resumo_benchmarks(dados["sensibilidade"]))
    adicionar_tabela(
        doc,
        "Tabela 4 – Benchmark estrutural principal das quatro ofertas da UFPA",
        benchmark,
        FONTE_DADOS,
    )
    _figura(
        doc,
        base_projeto,
        "11_benchmarks_ofertas_ufpa.png",
        13,
        "Benchmarks comparáveis por oferta",
        "Cada oferta UFPA é comparada a cursos externos com características estruturais selecionadas.",
        "A redução da heterogeneidade observável pode alterar a magnitude das diferenças brutas.",
        "O benchmark não controla confundidores não observados e não constitui pareamento causal.",
    )
    _figura(
        doc,
        base_projeto,
        "validada_16_benchmarks_por_oferta.png",
        14,
        "Sensibilidade dos benchmarks",
        "A figura destaca o cenário estrutural principal e a posição de cada oferta.",
        "Diferenças persistentes em cenários alternativos são mais robustas descritivamente.",
        "Alguns cenários podem conter poucos cursos, aumentando a incerteza.",
    )

    doc.add_heading("5.7 Associações ecológicas", level=2)
    _p(doc, resumo_associacoes(dados["associacoes"]))
    adicionar_tabela(
        doc,
        "Tabela 5 – Associações ecológicas com NT_GER médio",
        dados["associacoes"],
        FONTE_DADOS,
    )

    doc.add_heading("5.8 Comparações regionais e nacionais", level=2)
    _p(doc, resumo_regional(dados["comparacoes"]))
    adicionar_tabela(
        doc,
        "Tabela 6 – NT_GER por recortes regionais e nacionais",
        regional,
        FONTE_DADOS,
    )
    _figura(
        doc,
        base_projeto,
        "08_comparacao_regional_nacional.png",
        15,
        "Comparação regional e nacional",
        "As médias das ofertas e as médias ponderadas por participantes são apresentadas por recorte.",
        "Diferenças entre média simples e ponderada podem indicar influência do porte dos cursos.",
        "Recortes territoriais completos se sobrepõem e não são tratados como grupos independentes em testes.",
    )
    _figura(
        doc,
        base_projeto,
        "validada_19_comparacao_regional.png",
        16,
        "Comparação regional validada",
        "A versão validada acrescenta dispersão e auditoria dos recortes obrigatórios.",
        "A posição da UFPA deve ser interpretada no contexto da dispersão entre cursos.",
        "Médias agregadas podem ocultar heterogeneidade interna.",
    )

    doc.add_heading("5.9 Contraste interno das ofertas da UFPA", level=2)

    doc.add_heading("5.9.1 Participação e desempenho", level=3)
    _p(doc, resumo_participacao(base))
    _p(doc, resumo_desempenho_ufpa(base, dados["contraste"]))
    _figura(
        doc,
        base_projeto,
        "validada_14_participacao_ufpa.png",
        17,
        "Participação das ofertas da UFPA",
        "Inscritos, participantes e presença são comparados entre as quatro ofertas.",
        "Diferenças de participação podem afetar estabilidade e representatividade das estimativas.",
        "Participação não deve ser interpretada isoladamente como explicação do desempenho.",
    )
    _figura(
        doc,
        base_projeto,
        "validada_15_contraste_interno_ufpa.png",
        18,
        "Contraste interno Conceito 3 versus Conceito 4",
        "A figura reúne indicadores do contraste descritivo interno.",
        "O contraste em NT_GER e NT_OBJ pode ser mais forte que o observado em NT_DIS.",
        "Há apenas duas ofertas em cada estrato.",
    )
    _figura(
        doc,
        base_projeto,
        "13_contraste_interno_ufpa.png",
        19,
        "Síntese do contraste interno",
        "A figura integra indicadores selecionados das quatro ofertas da UFPA.",
        "A leitura multidimensional evita reduzir a comparação ao Conceito Enade.",
        "Indicadores de arquivos distintos são comparáveis somente no nível agregado do curso.",
    )

    doc.add_heading("5.9.2 Perfil discente", level=3)
    _p(
        doc,
        "As duas ofertas Conceito 3 são comparadas às duas Conceito 4 em indicadores de perfil. "
        "As diferenças representam composição agregada e não permitem inferir relações individuais "
        "entre características discentes e desempenho.",
    )

    doc.add_heading("5.9.3 Trajetória acadêmica", level=3)
    _p(
        doc,
        "Indicadores de turno, tempo desde o ingresso, trabalho, bolsas, auxílios e horas de estudo "
        "são usados para caracterizar as ofertas e gerar hipóteses no nível do curso.",
    )

    doc.add_heading("5.9.4 Processo formativo", level=3)
    _p(doc, resumo_processo(dados["processo"]))

    doc.add_heading("5.9.5 Recomendação", level=3)
    _p(doc, resumo_recomendacao(dados["recomendacao"]))

    doc.add_heading("5.9.6 Benchmark comparável", level=3)
    _p(doc, resumo_benchmarks(dados["sensibilidade"]))

    doc.add_heading("5.9.7 Síntese dos pontos distintivos", level=3)
    _p(
        doc,
        "No conjunto da UFPA, as ofertas Conceito 3 apresentam NT_GER e NT_OBJ médios inferiores às "
        "ofertas Conceito 4, enquanto NT_DIS é praticamente equivalente. O estrato Conceito 3 também "
        "apresenta, no agregado, maior proporção de renda até três salários mínimos, ação afirmativa, "
        "auxílio permanência e estudo de quatro horas ou mais, além de médias superiores em QE_I68 e "
        "QE_I69. Esses padrões coexistem no nível dos cursos e não devem ser interpretados como "
        "mecanismos causais individuais.",
    )

    doc.add_heading("6 DISCUSSÃO", level=1)
    _p(
        doc,
        "Os resultados evidenciam heterogeneidade interna relevante entre as quatro ofertas. O contraste "
        "médio em NT_GER acompanha uma diferença mais acentuada em NT_OBJ, ao passo que NT_DIS é muito "
        "próximo entre os estratos. Essa configuração sugere aprofundar componentes objetivos da prova, "
        "sem assumir que o Conceito Enade ou qualquer característica institucional tenha produzido o "
        "resultado. O N de cursos por estrato é dois.",
    )
    _p(
        doc,
        "No benchmark estrutural principal, as quatro ofertas da UFPA apresentam NT_GER médio abaixo "
        "dos respectivos grupos comparáveis, embora em magnitudes distintas. Belém está mais próxima "
        "de seu benchmark; Cametá apresenta a maior diferença negativa entre as quatro. A comparação "
        "regional mostra a UFPA agregada acima da Região Norte sem UFPA, mas abaixo do Brasil geral. "
        "Esses recortes têm objetivos descritivos diferentes e não podem ser tratados como grupos "
        "independentes quando se sobrepõem.",
    )
    _p(
        doc,
        "As associações ecológicas ampliam o contexto, mas não identificam mecanismos individuais. "
        "Correlação por curso pode refletir composição, porte, seletividade, condições institucionais "
        "ou outros confundidores não observados. Outliers nacionais foram preservados e sinalizados, "
        "em vez de removidos automaticamente.",
    )

    doc.add_heading("7 CONCLUSÃO", level=1)
    _p(
        doc,
        "Geografia na UFPA não apresenta oferta Conceito Enade 1 nas fontes de 2025. Entre as quatro "
        "ofertas localizadas, Belém e Ananindeua têm Conceito 4, enquanto Altamira e Cametá têm "
        "Conceito 3. O contraste descritivo aponta menor NT_GER e NT_OBJ médios no estrato Conceito 3, "
        "sem diferença material em NT_DIS. As quatro ofertas ficam abaixo dos respectivos benchmarks "
        "estruturais principais em NT_GER, com intensidades diferentes. Perfil, trajetória, processo "
        "formativo e recomendação adicionam dimensões que impedem uma leitura unidimensional do resultado. "
        "As evidências geram padrões e hipóteses, não causalidade individual ou institucional.",
    )

    doc.add_heading("REFERÊNCIAS", level=1)
    adicionar_referencias(doc)

    doc.add_heading("APÊNDICE A – REGRAS DE INTEGRIDADE", level=1)
    for texto in [
        "CO_CURSO é a unidade principal.",
        "Arquivos temáticos são agregados antes das junções.",
        "Não há join individual entre desempenho e questionário.",
        "Não há reconstrução de estudante entre arquivos.",
        "Ausência de conceito não é Conceito 1.",
        "O Grupo A permanece vazio em Geografia.",
        "Conceito 3 não é tratado como insuficiência.",
        "Relações entre temas diferentes são ecológicas.",
        "Pará, Norte e Brasil completos não são grupos independentes em testes.",
    ]:
        _p(doc, texto)

    doc.add_heading("APÊNDICE B – APROFUNDAMENTOS SUGERIDOS", level=1)
    aprofundamentos = [
        "1. Sensibilidade de porte — pergunta: os déficits frente aos benchmarks persistem sob faixas "
        "de porte mais estreitas? Variáveis: NT_GER, NT_OBJ, NT_DIS, participantes e atributos "
        "estruturais. Método: cenários progressivos e, quando viável, pareamento. Limitação: poucos "
        "comparáveis em alguns estratos.",
        "2. Componente objetivo — pergunta: quais padrões de acerto ajudam a caracterizar a diferença "
        "observada em NT_OBJ? Variáveis: QT_ACERTOS e indicadores disponíveis no mesmo arquivo de "
        "desempenho. Método: análise intrarquivo por oferta. Limitação: relações mecânicas entre nota "
        "e acertos.",
        "3. Processo formativo — pergunta: quais itens QE_I20–QE_I66 mantêm contraste Conceito 3 × 4 "
        "considerando respostas válidas e referências externas? Método: análise item a item, "
        "dimensionalidade e consistência interna antes de índices. Limitação: escala ordinal e "
        "respostas autorreferidas.",
        "4. Permanência e trajetória — pergunta: como renda, trabalho, auxílios, bolsas, turno e tempo "
        "desde o ingresso se organizam ecologicamente entre cursos? Método: perfis padronizados e "
        "Spearman por curso. Limitação: falácia ecológica e confundimento residual.",
        "5. Recomendação e processo formativo — pergunta: QE_I68 e QE_I69 se associam ecologicamente a "
        "dimensões formativas específicas no universo nacional de Geografia? Método: agregação por "
        "curso, Spearman, dispersão e outliers. Limitação: não representa relação individual.",
    ]
    for texto in aprofundamentos:
        _p(doc, texto)

    doc.add_heading("APÊNDICE C – RÓTULOS OFICIAIS DOS ITENS QE_I20–QE_I66", level=1)
    for codigo in [f"QE_I{i}" for i in range(20, 67)]:
        _p(doc, f"{codigo} – {ROTULOS_QE[codigo]}")
    _p(doc, f"Fonte dos rótulos: {FONTE_QUESTIONARIO}")

    configurar_cabecalho_rodape(doc)
    saida_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(saida_docx)

    saida_md.parent.mkdir(parents=True, exist_ok=True)
    saida_md.write_text(_markdown(dados), encoding="utf-8")

    conversao = converter_docx_para_pdf(saida_docx, saida_docx.parent)
    resultado = ResultadoRelatorio(
        docx=saida_docx,
        markdown=saida_md,
        conversao_pdf=conversao,
    )
    return resultado.como_dict()
